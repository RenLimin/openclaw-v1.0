#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 表格提取模块
=================
智能提取 Word 表格并转为结构化数据，支持合并单元格、跨页表格、表头识别

功能:
  - 提取 Word 表格转结构化数据（DataFrame/List/JSON）
  - 保留单元格格式（合并单元格识别）
  - 跨页表格自动合并
  - 表头识别
  - 表格跨行跨列检测

作者: Word-Analyzer Team
版本: v1.0
日期: 2026-04-24
"""

import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from collections import defaultdict

try:
    from docx import Document
    from docx.table import Table, _Cell, _Row
except ImportError:
    print("❌ 错误: 请安装 python-docx: pip install python-docx")
    exit(1)

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False


class TableExtractor:
    """Word 表格提取器"""

    def __init__(self, file_path: str, config_path: str = None):
        """
        初始化表格提取器

        Args:
            file_path: Word 文档路径 (.docx)
            config_path: 配置文件路径（可选）
        """
        self.file_path = Path(file_path)
        self.config = self._load_config(config_path)
        self.doc = None
        self._tables_cache = None

    def _load_config(self, config_path: str = None) -> Dict:
        """加载配置文件"""
        default_config = {
            'table_rules': {
                'header_detection': {
                    'max_header_rows': 3,
                    'header_keywords': ["序号", "编号", "名称", "项目", "内容", "说明", "备注", "日期", "金额", "数量"],
                    'header_bold': True
                },
                'cross_page_merge': {
                    'enabled': True,
                    'max_gap_lines': 2,
                    'detect_duplicate_headers': True
                },
                'merged_cells': {
                    'fill_mode': "repeat"  # repeat / empty / first_only
                }
            }
        }

        if config_path and HAS_YAML:
            cfg_path = Path(config_path)
            if cfg_path.exists():
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    user_config = yaml.safe_load(f)
                    # 合并配置
                    if 'table_rules' in user_config:
                        for key, value in user_config['table_rules'].items():
                            default_config['table_rules'][key].update(value)

        return default_config

    def load_document(self) -> bool:
        """加载 Word 文档"""
        try:
            self.doc = Document(str(self.file_path))
            return True
        except Exception as e:
            print(f"❌ 加载文档失败: {e}")
            return False

    def _detect_merged_cells(self, table: Table) -> Dict[Tuple[int, int], Tuple[int, int, int, int]]:
        """
        检测表格中的合并单元格

        Returns:
            Dict: {(row_idx, col_idx): (start_row, start_col, rowspan, colspan)}
        """
        merged_info = {}
        cell_mapping = {}

        # 遍历所有单元格，识别合并单元格
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 使用 _tc 属性获取 XML 元素
                cell_id = id(cell._tc)
                if cell_id not in cell_mapping:
                    cell_mapping[cell_id] = (row_idx, col_idx)
                else:
                    # 这是一个合并单元格的引用
                    start_row, start_col = cell_mapping[cell_id]
                    if (start_row, start_col) not in merged_info:
                        merged_info[(start_row, start_col)] = (start_row, start_col, 1, 1)
                    # 更新行列跨度
                    sr, sc, cr, cc = merged_info[(start_row, start_col)]
                    new_rowspan = max(cr, row_idx - sr + 1)
                    new_colspan = max(cc, col_idx - sc + 1)
                    merged_info[(start_row, start_col)] = (sr, sc, new_rowspan, new_colspan)

                    # 记录当前合并位置
                    merged_info[(row_idx, col_idx)] = (sr, sc, new_rowspan, new_colspan)

        return merged_info

    def _extract_cell_info(self, cell: _Cell, row_idx: int, col_idx: int,
                          merged_info: Dict) -> Dict[str, Any]:
        """提取单元格信息（含格式和合并信息）"""
        # 检查是否是合并单元格
        is_merged = (row_idx, col_idx) in merged_info
        merge_data = None
        if is_merged:
            sr, sc, rowspan, colspan = merged_info[(row_idx, col_idx)]
            is_merge_start = (sr == row_idx and sc == col_idx)
            merge_data = {
                'start_row': sr,
                'start_col': sc,
                'rowspan': rowspan,
                'colspan': colspan,
                'is_merge_start': is_merge_start
            }

        # 提取单元格样式（第一个段落的样式）
        bold = False
        font_size = 0
        font_name = ""
        if cell.paragraphs and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            bold = bool(run.bold)
            font_size = run.font.size.pt if run.font.size else 0
            font_name = run.font.name or ""

        return {
            'row': row_idx,
            'col': col_idx,
            'text': cell.text.strip(),
            'bold': bold,
            'font_size': font_size,
            'font_name': font_name,
            'is_merged': is_merged,
            'merge_info': merge_data
        }

    def _detect_header_rows(self, table_data: List[List[Dict]]) -> List[int]:
        """智能检测表头行"""
        header_rules = self.config['table_rules']['header_detection']
        max_header_rows = header_rules['max_header_rows']
        keywords = header_rules['header_keywords']
        require_bold = header_rules['header_bold']

        header_rows = []
        check_rows = min(max_header_rows, len(table_data))

        for row_idx in range(check_rows):
            row = table_data[row_idx]
            bold_count = sum(1 for cell in row if cell['bold'])
            keyword_count = sum(
                1 for cell in row
                for kw in keywords
                if kw in cell['text']
            )

            # 表头判断条件：有关键词 或 大部分单元格加粗
            if keyword_count > 0 or (require_bold and bold_count > len(row) * 0.5):
                header_rows.append(row_idx)

        return header_rows if header_rows else [0]  # 默认第一行是表头

    def extract_table(self, table: Table, table_index: int) -> Dict[str, Any]:
        """提取单个表格"""
        merged_info = self._detect_merged_cells(table)
        fill_mode = self.config['table_rules']['merged_cells']['fill_mode']

        # 提取所有单元格数据
        table_data = []
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                cell_info = self._extract_cell_info(cell, row_idx, col_idx, merged_info)

                # 处理合并单元格填充
                if cell_info['is_merged'] and not cell_info['merge_info']['is_merge_start']:
                    sr, sc = cell_info['merge_info']['start_row'], cell_info['merge_info']['start_col']
                    if fill_mode == 'repeat':
                        # 重复合并单元格内容
                        cell_info['text'] = table_data[sr][sc]['text']
                    elif fill_mode == 'empty':
                        cell_info['text'] = ''
                    # first_only 保持原样

                row_data.append(cell_info)
            table_data.append(row_data)

        # 检测表头
        header_rows = self._detect_header_rows(table_data)

        # 计算行列数
        rows_count = len(table_data)
        cols_count = len(table_data[0]) if rows_count > 0 else 0

        return {
            'index': table_index,
            'rows': rows_count,
            'columns': cols_count,
            'header_rows': header_rows,
            'merged_cells_count': len(set(
                (v[0], v[1]) for v in merged_info.values()
            )) if merged_info else 0,
            'cells': table_data
        }

    def _are_tables_continuous(self, table1: Dict, table2: Dict) -> bool:
        """判断两个表格是否是跨页拆分的连续表格"""
        if not self.config['table_rules']['cross_page_merge']['enabled']:
            return False

        # 列数必须相同
        if table1['columns'] != table2['columns']:
            return False

        # 检查表头是否重复（跨页表格通常重复表头）
        if self.config['table_rules']['cross_page_merge']['detect_duplicate_headers']:
            headers1 = [cell['text'] for cell in table1['cells'][0]]
            headers2 = [cell['text'] for cell in table2['cells'][0]]
            if headers1 == headers2:
                return True

        return False

    def _merge_cross_page_tables(self, tables: List[Dict]) -> List[Dict]:
        """合并跨页拆分的表格"""
        if len(tables) < 2:
            return tables

        merged_tables = []
        i = 0

        while i < len(tables):
            current_table = tables[i]

            # 尝试合并后续的连续表格
            j = i + 1
            while j < len(tables) and self._are_tables_continuous(current_table, tables[j]):
                # 跳过第二个表格的表头（与第一个重复）
                skip_rows = len(tables[j]['header_rows'])
                current_table['cells'].extend(tables[j]['cells'][skip_rows:])
                current_table['rows'] = len(current_table['cells'])
                current_table['merged_cells_count'] += tables[j]['merged_cells_count']
                current_table['is_merged_table'] = True
                j += 1

            merged_tables.append(current_table)
            i = j

        return merged_tables

    def extract_all(self, merge_cross_page: bool = True) -> List[Dict[str, Any]]:
        """提取所有表格"""
        if self._tables_cache is not None:
            return self._tables_cache

        if not self.doc:
            self.load_document()

        print(f"📋 开始提取表格...")

        tables = []
        for idx, table in enumerate(self.doc.tables):
            table_data = self.extract_table(table, idx)
            tables.append(table_data)
            print(f"  表格 {idx}: {table_data['rows']}行 x {table_data['columns']}列")

        if merge_cross_page:
            tables = self._merge_cross_page_tables(tables)
            print(f"  合并跨页表格后共 {len(tables)} 个表格")

        self._tables_cache = tables
        return tables

    def get_table(self, index: int, format: str = "list") -> Any:
        """
        获取指定索引的表格

        Args:
            index: 表格索引（从0开始）
            format: 输出格式: 'list', 'dict', 'dataframe', 'json'

        Returns:
            按指定格式返回的表格数据
        """
        tables = self.extract_all()

        if index < 0 or index >= len(tables):
            raise IndexError(f"表格索引超出范围: 共 {len(tables)} 个表格")

        table = tables[index]
        headers = [cell['text'] for cell in table['cells'][0]] if table['cells'] else []

        # 提取纯数据（跳过表头行）
        data_start_row = max(table['header_rows']) + 1 if table['header_rows'] else 1
        data_rows = []

        for row in table['cells'][data_start_row:]:
            data_rows.append([cell['text'] for cell in row])

        if format == "list":
            return {'headers': headers, 'data': data_rows}

        elif format == "dict":
            dict_data = []
            for row in data_rows:
                row_dict = {}
                for i, header in enumerate(headers):
                    if i < len(row):
                        row_dict[header] = row[i]
                    else:
                        row_dict[header] = ""
                dict_data.append(row_dict)
            return dict_data

        elif format == "dataframe":
            if not HAS_PANDAS:
                print("⚠️  pandas 未安装，返回 list 格式")
                return {'headers': headers, 'data': data_rows}
            return pd.DataFrame(data_rows, columns=headers)

        elif format == "json":
            dict_data = []
            for row in data_rows:
                row_dict = {}
                for i, header in enumerate(headers):
                    if i < len(row):
                        row_dict[header] = row[i]
                    else:
                        row_dict[header] = ""
                dict_data.append(row_dict)
            return json.dumps(dict_data, ensure_ascii=False, indent=2)

        else:
            raise ValueError(f"不支持的格式: {format}")

    def export_all(self, output_path: str, format: str = "json"):
        """导出所有表格"""
        tables = self.extract_all()

        # 简化输出，只保留关键信息
        export_data = []
        for table in tables:
            simple_table = {
                'index': table['index'],
                'rows': table['rows'],
                'columns': table['columns'],
                'headers': [cell['text'] for cell in table['cells'][0]] if table['cells'] else [],
                'data': [
                    [cell['text'] for cell in row]
                    for row in table['cells'][1:]  # 跳过第一行（表头）
                ]
            }
            export_data.append(simple_table)

        if format == "json":
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)

        elif format == "csv" and HAS_PANDAS:
            for i, table in enumerate(export_data):
                df = pd.DataFrame(table['data'], columns=table['headers'])
                csv_path = Path(output_path).with_stem(f"{Path(output_path).stem}_table{i}")
                df.to_csv(csv_path, index=False, encoding='utf-8-sig')
            print(f"💾 已导出 {len(export_data)} 个 CSV 文件")
            return

        print(f"💾 结果已导出: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Word 表格提取工具')
    parser.add_argument('--input', '-i', required=True, help='输入 Word 文档路径')
    parser.add_argument('--output', '-o', help='输出文件路径')
    parser.add_argument('--table-index', '-t', type=int, help='提取指定索引的表格')
    parser.add_argument('--format', '-f', choices=['json', 'csv', 'list', 'dict'],
                       default='json', help='输出格式 (默认: json)')
    parser.add_argument('--no-merge', action='store_true', help='不合并跨页表格')

    args = parser.parse_args()

    extractor = TableExtractor(args.input)

    if args.table_index is not None:
        table_data = extractor.get_table(args.table_index, format=args.format)
        if args.output:
            if args.format == "dataframe" and HAS_PANDAS:
                table_data.to_csv(args.output, index=False, encoding='utf-8-sig')
            else:
                with open(args.output, 'w', encoding='utf-8') as f:
                    json.dump(table_data, f, ensure_ascii=False, indent=2)
        else:
            print(json.dumps(table_data, ensure_ascii=False, indent=2))
    else:
        output = args.output or Path(args.input).with_suffix('.tables.json')
        extractor.export_all(output, format=args.format)


if __name__ == "__main__":
    main()
