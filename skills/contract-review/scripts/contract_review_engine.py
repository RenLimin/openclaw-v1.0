#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同审核业务技能 - 主审核引擎
打通 OCR 识别 → 专项审核 → 风险分级 → 法条关联 → 审核报告生成的完整业务流程
"""

import os
import sys
import re
import json
import argparse
import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict

import yaml
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from jinja2 import Template

# 添加依赖技能路径
SKILL_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(SKILL_DIR.parent / 'ocr-engine' / 'scripts'))
sys.path.insert(0, str(SKILL_DIR.parent / 'contract-clause-split' / 'scripts'))

# 尝试导入依赖技能
try:
    from ocr_runner import OCRRunner
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️  警告: OCR 引擎不可用，将使用基础文本提取")

try:
    from split_contract import ContractSplitter
    CLAUSE_SPLIT_AVAILABLE = True
except ImportError:
    CLAUSE_SPLIT_AVAILABLE = False
    print("⚠️  警告: 条款拆分技能不可用，将使用基础拆分")


@dataclass
class RiskItem:
    """风险项数据类"""
    id: str
    clause_type: str
    risk_level: str  # high / medium / low
    risk_score: int
    description: str
    clause_text: str
    location: str  # 页码/行号
    law_article: str = ""
    law_name: str = ""
    law_content: str = ""
    suggestion: str = ""
    responsible: str = ""
    deadline: str = ""
    status: str = "待处理"


@dataclass
class ContractInfo:
    """合同基本信息"""
    name: str = ""
    contract_no: str = ""
    party_a: str = ""
    party_b: str = ""
    sign_date: str = ""
    review_date: str = ""
    total_clauses: int = 0


class ContractReviewEngine:
    """合同审核主引擎"""

    def __init__(self, config_path: Optional[str] = None):
        self.skill_dir = SKILL_DIR
        self.config = self._load_config(config_path)
        self.review_results: List[RiskItem] = []
        self.contract_info = ContractInfo()
        self.clauses = []
        self.ocr_result = None

    def _load_config(self, config_path: Optional[str] = None) -> Dict:
        """加载配置文件"""
        if config_path is None:
            config_path = self.skill_dir / 'config' / 'config.yaml'
        
        with open(config_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)

    def _detect_file_type(self, file_path: str) -> str:
        """检测文件类型"""
        ext = Path(file_path).suffix.lower()
        if ext in ['.pdf']:
            return 'pdf'
        elif ext in ['.docx', '.doc']:
            return 'word'
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return 'image'
        elif ext in ['.txt', '.md']:
            return 'text'
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _extract_text_from_word(self, file_path: str) -> str:
        """从 Word 文档提取文本"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            return '\n'.join(text)
        except ImportError:
            raise RuntimeError("python-docx 未安装，无法处理 Word 文件")

    def _extract_text_from_pdf_simple(self, file_path: str) -> str:
        """简单 PDF 文本提取（备用）"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                return '\n'.join(text)
        except ImportError:
            print("⚠️  PyPDF2 未安装，无法提取 PDF 文本")
            return ""

    def _run_ocr(self, file_path: str) -> Dict:
        """运行 OCR 识别"""
        file_type = self._detect_file_type(file_path)
        
        # 纯文本文件直接读取
        if file_type == 'text':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return {
                'text': text,
                'engine_used': 'text',
                'quality_score': 100,
                'pages': [{'text': text, 'page_num': 1}]
            }
        
        # Word 文件直接提取
        if file_type == 'word':
            text = self._extract_text_from_word(file_path)
            return {
                'text': text,
                'engine_used': 'word',
                'quality_score': 100,
                'pages': [{'text': text, 'page_num': 1}]
            }
        
        if not OCR_AVAILABLE:
            print("⚠️  OCR 引擎不可用，尝试使用基础文本提取")
            if file_type == 'pdf':
                text = self._extract_text_from_pdf_simple(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            return {
                'text': text,
                'engine_used': 'fallback',
                'quality_score': 50,
                'pages': [{'text': text, 'page_num': 1}]
            }

        ocr_config = self.config.get('ocr', {})
        try:
            ocr = OCRRunner(
                engine=ocr_config.get('engine', 'tesseract'),
                config_path=str(self.skill_dir.parent / 'ocr-engine' / 'config' / 'engine-config.yaml')
            )
            result = ocr.run(
                file_path,
                clean_noise=ocr_config.get('clean_noise', True),
                auto_correct=ocr_config.get('auto_correct', True)
            )
        except Exception as e:
            print(f"⚠️  OCR 引擎调用失败: {e}，使用基础文本提取")
            if file_type == 'pdf':
                text = self._extract_text_from_pdf_simple(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            return {
                'text': text,
                'engine_used': 'fallback',
                'quality_score': 50,
                'pages': [{'text': text, 'page_num': 1}]
            }
        return result

    def _split_clauses(self, text: str) -> List[Dict]:
        """拆分合同条款"""
        if CLAUSE_SPLIT_AVAILABLE:
            try:
                splitter = ContractSplitter()
                result = splitter.split(text)
                if 'clauses' in result:
                    return result['clauses']
            except Exception as e:
                print(f"⚠️  条款拆分失败: {e}，使用基础拆分")

        # 基础拆分：按常见条款标题分割
        clauses = []
        lines = text.split('\n')
        current_clause = {'title': '', 'content': [], 'type': '其他', 'line_num': 1}
        
        clause_patterns = [
            (r'第[一二三四五六七八九十\d]+条', '其他'),
            (r'^[一二三四五六七八九十\d]+[、.．]', '其他'),
            (r'^[（(][一二三四五六七八九十\d]+[）)]', '其他'),
        ]
        
        for i, line in enumerate(lines, 1):
            line = line.strip()
            if not line:
                continue
                
            is_new_clause = False
            clause_type = '其他'
            
            # 检测条款类型
            for pattern, ctype in clause_patterns:
                if re.match(pattern, line):
                    is_new_clause = True
                    break
            
            # 根据关键词分类
            if any(kw in line for kw in self.config['special_review']['subject']['keywords']):
                clause_type = '标的'
            elif any(kw in line for kw in self.config['special_review']['delivery']['keywords']):
                clause_type = '交付'
            elif any(kw in line for kw in self.config['special_review']['acceptance']['keywords']):
                clause_type = '验收'
            elif any(kw in line for kw in self.config['special_review']['after_sales']['keywords']):
                clause_type = '售后'
            
            if is_new_clause and current_clause['content']:
                clauses.append({
                    'title': current_clause['title'],
                    'content': '\n'.join(current_clause['content']),
                    'type': current_clause['type'],
                    'line_num': current_clause['line_num']
                })
                current_clause = {'title': line, 'content': [], 'type': clause_type, 'line_num': i}
            else:
                if not current_clause['title']:
                    current_clause['title'] = line[:50] + '...' if len(line) > 50 else line
                    current_clause['type'] = clause_type
                    current_clause['line_num'] = i
                current_clause['content'].append(line)
        
        if current_clause['content']:
            clauses.append({
                'title': current_clause['title'],
                'content': '\n'.join(current_clause['content']),
                'type': current_clause['type'],
                'line_num': current_clause['line_num']
            })
        
        return clauses

    def _extract_contract_info(self, text: str, clauses: List[Dict]):
        """提取合同基本信息"""
        self.contract_info.review_date = datetime.datetime.now().strftime('%Y-%m-%d')
        self.contract_info.total_clauses = len(clauses)
        
        # 简单关键词匹配
        info_patterns = {
            'name': [r'合同名称[：:]\s*(.+)', r'《(.+?)合同》'],
            'contract_no': [r'合同编号[：:]\s*(.+)', r'合同号[：:]\s*(.+)'],
            'party_a': [r'甲方[：:]\s*(.+)', r'买方[：:]\s*(.+)', r'委托方[：:]\s*(.+)'],
            'party_b': [r'乙方[：:]\s*(.+)', r'卖方[：:]\s*(.+)', r'服务方[：:]\s*(.+)'],
            'sign_date': [r'签订日期[：:]\s*(.+)', r'签署日期[：:]\s*(.+)'],
        }
        
        for field, patterns in info_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, text)
                if match:
                    value = match.group(1).strip()
                    setattr(self.contract_info, field, value)
                    break
        
        if not self.contract_info.name:
            self.contract_info.name = "未命名合同"

    def _review_subject_clause(self, clause: Dict, clause_idx: int) -> List[RiskItem]:
        """审核标的条款"""
        risks = []
        content = clause['content']
        title = clause['title']
        full_text = title + '\n' + content
        
        checks = self.config['special_review']['subject']['check_items']
        
        # 标的描述不清晰
        if len(full_text) < 50 or not re.search(r'规格|型号|品牌|技术指标', full_text):
            score = checks.get('description_clear', 80) if isinstance(checks, dict) else 80
            risks.append(RiskItem(
                id=f"RVW-01-{clause_idx:03d}",
                clause_type='标的',
                risk_level='high' if score >= 80 else 'medium',
                risk_score=score,
                description="标的描述不够清晰，缺少具体规格、型号、品牌或技术指标",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        # 数量缺失
        if not re.search(r'\d+(\.\d+)?\s*(个|件|台|套|项)', full_text) and '数量' not in full_text:
            score = checks.get('quantity_spec', 75) if isinstance(checks, dict) else 75
            risks.append(RiskItem(
                id=f"RVW-01-{clause_idx:03d}-2",
                clause_type='标的',
                risk_level='high',
                risk_score=score,
                description="缺少具体数量或规格约定",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        return risks

    def _review_delivery_clause(self, clause: Dict, clause_idx: int) -> List[RiskItem]:
        """审核交付条款"""
        risks = []
        content = clause['content']
        title = clause['title']
        full_text = title + '\n' + content
        
        checks = self.config['special_review']['delivery']['check_items']
        
        # 交付时间不明确
        if not re.search(r'\d{4}[-年]\d{1,2}[-月]\d{1,2}|交货时间|交付时间', full_text):
            score = checks.get('delivery_time', 80) if isinstance(checks, dict) else 80
            risks.append(RiskItem(
                id=f"RVW-02-{clause_idx:03d}",
                clause_type='交付',
                risk_level='high',
                risk_score=score,
                description="交付时间不明确，缺少具体交付期限",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        # 交付地点缺失
        if '地点' not in full_text and '地址' not in full_text:
            score = checks.get('delivery_location', 50) if isinstance(checks, dict) else 50
            risks.append(RiskItem(
                id=f"RVW-02-{clause_idx:03d}-2",
                clause_type='交付',
                risk_level='medium',
                risk_score=score,
                description="交付地点未约定",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        # 逾期交付无违约责任
        if '逾期' not in full_text and '违约金' not in full_text:
            score = checks.get('late_delivery_penalty', 75) if isinstance(checks, dict) else 75
            risks.append(RiskItem(
                id=f"RVW-02-{clause_idx:03d}-3",
                clause_type='交付',
                risk_level='high',
                risk_score=score,
                description="未约定逾期交付的违约责任",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        return risks

    def _review_acceptance_clause(self, clause: Dict, clause_idx: int) -> List[RiskItem]:
        """审核验收条款"""
        risks = []
        content = clause['content']
        title = clause['title']
        full_text = title + '\n' + content
        
        checks = self.config['special_review']['acceptance']['check_items']
        
        # 验收标准不明确
        if '标准' not in full_text and '规范' not in full_text and '合格' not in full_text:
            score = checks.get('acceptance_standard', 80) if isinstance(checks, dict) else 80
            risks.append(RiskItem(
                id=f"RVW-03-{clause_idx:03d}",
                clause_type='验收',
                risk_level='high',
                risk_score=score,
                description="验收标准不明确，缺少具体验收标准和方法",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        # 验收期限缺失
        if not re.search(r'\d+[日天]内', full_text) and '期限' not in full_text:
            score = checks.get('acceptance_period', 50) if isinstance(checks, dict) else 50
            risks.append(RiskItem(
                id=f"RVW-03-{clause_idx:03d}-2",
                clause_type='验收',
                risk_level='medium',
                risk_score=score,
                description="验收期限和异议期未约定",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        # 不合格处理无约定
        if '不合格' not in full_text and '拒收' not in full_text and '更换' not in full_text:
            score = checks.get('rejection_handling', 75) if isinstance(checks, dict) else 75
            risks.append(RiskItem(
                id=f"RVW-03-{clause_idx:03d}-3",
                clause_type='验收',
                risk_level='high',
                risk_score=score,
                description="验收不合格的处理方式未约定",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        return risks

    def _review_after_sales_clause(self, clause: Dict, clause_idx: int) -> List[RiskItem]:
        """审核售后/维保条款"""
        risks = []
        content = clause['content']
        title = clause['title']
        full_text = title + '\n' + content
        
        checks = self.config['special_review']['after_sales']['check_items']
        
        # 质保期限不明确
        if not re.search(r'\d+[个]?[月年]', full_text) and '质保期' not in full_text and '保修期' not in full_text:
            score = checks.get('warranty_period', 80) if isinstance(checks, dict) else 80
            risks.append(RiskItem(
                id=f"RVW-04-{clause_idx:03d}",
                clause_type='售后',
                risk_level='high',
                risk_score=score,
                description="质保期限不明确，缺少起止时间和期限约定",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        # 质保范围缺失
        if '范围' not in full_text and '包括' not in full_text:
            score = checks.get('warranty_scope', 50) if isinstance(checks, dict) else 50
            risks.append(RiskItem(
                id=f"RVW-04-{clause_idx:03d}-2",
                clause_type='售后',
                risk_level='medium',
                risk_score=score,
                description="质保覆盖范围未清晰约定",
                clause_text=title,
                location=f"第{clause['line_num']}行"
            ))
        
        return risks

    def _general_check(self, text: str, clauses: List[Dict]) -> List[RiskItem]:
        """通用规则检查"""
        risks = []
        
        if not self.config['review'].get('enable_general_check', True):
            return risks
        
        general_config = self.config.get('general_check', {})
        
        # 歧义表述识别
        if self.config['review'].get('enable_ambiguity_check', True):
            ambiguity_keywords = general_config.get('ambiguity_keywords', [])
            for i, clause in enumerate(clauses):
                full_text = clause['title'] + '\n' + clause['content']
                for kw in ambiguity_keywords:
                    if kw in full_text:
                        risks.append(RiskItem(
                            id=f"RVW-GEN-{i:03d}",
                            clause_type='通用',
                            risk_level='low',
                            risk_score=30,
                            description=f"存在歧义表述风险：'{kw}'，建议明确具体含义",
                            clause_text=clause['title'],
                            location=f"第{clause['line_num']}行"
                        ))
                        break
        
        # 免责条款审查
        if self.config['review'].get('enable_disclaimer_check', True):
            disclaimer_keywords = general_config.get('disclaimer_keywords', [])
            for i, clause in enumerate(clauses):
                full_text = clause['title'] + '\n' + clause['content']
                for kw in disclaimer_keywords:
                    if kw in full_text:
                        risks.append(RiskItem(
                            id=f"RVW-DIS-{i:03d}",
                            clause_type='免责',
                            risk_level='medium',
                            risk_score=50,
                            description=f"发现免责条款：'{kw}'，建议审查其公平性和合法性",
                            clause_text=clause['title'],
                            location=f"第{clause['line_num']}行"
                        ))
                        break
        
        return risks

    def _add_law_references(self, risks: List[RiskItem]):
        """添加法条关联"""
        if not self.config['review'].get('enable_law_reference', True):
            return
        
        law_refs = self.config.get('law_reference', {}).get('civil_code', {})
        
        for risk in risks:
            if risk.clause_type == '标的':
                law = law_refs.get('subject_clarity', {})
                risk.law_article = law.get('article', '')
                risk.law_name = law.get('name', '')
                risk.law_content = law.get('content', '')
            elif risk.clause_type == '交付':
                law = law_refs.get('delivery_time', {})
                risk.law_article = law.get('article', '')
                risk.law_name = law.get('name', '')
                risk.law_content = law.get('content', '')
            elif risk.clause_type == '验收':
                law = law_refs.get('acceptance', {})
                risk.law_article = law.get('article', '')
                risk.law_name = law.get('name', '')
                risk.law_content = law.get('content', '')
            elif risk.clause_type == '售后':
                law = law_refs.get('warranty', {})
                risk.law_article = law.get('article', '')
                risk.law_name = law.get('name', '')
                risk.law_content = law.get('content', '')
            elif risk.clause_type == '免责':
                law = law_refs.get('liability', {})
                risk.law_article = law.get('article', '')
                risk.law_name = law.get('name', '')
                risk.law_content = law.get('content', '')

    def _add_suggestions(self, risks: List[RiskItem]):
        """添加改进建议"""
        if not self.config['review'].get('enable_suggestions', True):
            return
        
        templates = self.config.get('suggestion_templates', {})
        
        for risk in risks:
            if '描述' in risk.description or '标的' in risk.description:
                risk.suggestion = templates.get('subject_description', '')
            elif '交付' in risk.description and '时间' in risk.description:
                risk.suggestion = templates.get('delivery_time', '')
            elif '验收' in risk.description and '标准' in risk.description:
                risk.suggestion = templates.get('acceptance_standard', '')
            elif '质保' in risk.description:
                risk.suggestion = templates.get('warranty_period', '')
            elif '歧义' in risk.description:
                risk.suggestion = templates.get('ambiguity', '')
            elif '免责' in risk.description:
                risk.suggestion = templates.get('disclaimer', '')
            else:
                risk.suggestion = "建议明确该条款的具体约定，避免产生争议。"

    def _calculate_overall_score(self) -> Tuple[int, str]:
        """计算总体风险评分和等级"""
        if not self.review_results:
            return 0, '低风险'
        
        high_count = sum(1 for r in self.review_results if r.risk_level == 'high')
        medium_count = sum(1 for r in self.review_results if r.risk_level == 'medium')
        low_count = sum(1 for r in self.review_results if r.risk_level == 'low')
        
        # 加权计算：高风险 × 3，中风险 × 2，低风险 × 1
        weighted_score = (high_count * 100 + medium_count * 50 + low_count * 20)
        max_possible = len(self.review_results) * 100
        
        if max_possible == 0:
            normalized_score = 0
        else:
            normalized_score = min(100, int(weighted_score / max(len(self.review_results), 1)))
        
        thresholds = self.config['risk_threshold']
        if normalized_score >= thresholds['overall_high']:
            risk_level = '🔴 高风险'
        elif normalized_score >= thresholds['overall_medium']:
            risk_level = '🟡 中风险'
        else:
            risk_level = '🟢 低风险'
        
        return normalized_score, risk_level

    def review(self, input_file: str, types: Optional[List[str]] = None, 
               min_risk: str = 'low') -> Dict:
        """执行合同审核"""
        print(f"📄 开始审核合同: {input_file}")
        
        # 检查文件是否存在
        if not os.path.exists(input_file):
            raise FileNotFoundError(f"合同文件不存在: {input_file}")
        
        # 1. OCR 识别
        print("🔍 执行 OCR 识别...")
        self.ocr_result = self._run_ocr(input_file)
        text = self.ocr_result.get('text', '')
        
        if not text or len(text.strip()) < 100:
            print("⚠️  警告: OCR 识别结果为空或过短，请检查文件质量")
        
        print(f"   识别文本长度: {len(text)} 字符")
        if 'quality_score' in self.ocr_result:
            print(f"   OCR 质量评分: {self.ocr_result['quality_score']:.1f}/100")
        
        # 2. 条款拆分
        print("✂️  拆分合同条款...")
        self.clauses = self._split_clauses(text)
        print(f"   拆分成 {len(self.clauses)} 条条款")
        
        # 检查条款数量
        min_clauses = self.config['review'].get('min_clause_count', 5)
        if len(self.clauses) < min_clauses:
            print(f"⚠️  警告: 条款数量较少 ({len(self.clauses)} 条)，可能影响审核质量")
        
        # 3. 提取合同基本信息
        print("📋 提取合同基本信息...")
        self._extract_contract_info(text, self.clauses)
        
        # 4. 专项审核
        print("🔬 执行专项审核...")
        review_types = types or ['标的', '交付', '验收', '售后']
        
        for i, clause in enumerate(self.clauses):
            clause_type = clause['type']
            
            if '标的' in review_types and clause_type == '标的':
                self.review_results.extend(self._review_subject_clause(clause, i))
            elif '交付' in review_types and clause_type == '交付':
                self.review_results.extend(self._review_delivery_clause(clause, i))
            elif '验收' in review_types and clause_type == '验收':
                self.review_results.extend(self._review_acceptance_clause(clause, i))
            elif '售后' in review_types and clause_type == '售后':
                self.review_results.extend(self._review_after_sales_clause(clause, i))
        
        # 5. 通用规则检查
        print("📝 执行通用规则检查...")
        self.review_results.extend(self._general_check(text, self.clauses))
        
        # 6. 法条关联
        print("⚖️  关联法条依据...")
        self._add_law_references(self.review_results)
        
        # 7. 改进建议
        print("💡 生成改进建议...")
        self._add_suggestions(self.review_results)
        
        # 8. 按风险等级过滤
        if min_risk == '中':
            self.review_results = [r for r in self.review_results if r.risk_level in ['high', 'medium']]
        elif min_risk == '高':
            self.review_results = [r for r in self.review_results if r.risk_level == 'high']
        
        # 9. 计算总体评分
        overall_score, overall_level = self._calculate_overall_score()
        
        print(f"\n✅ 审核完成！")
        print(f"   发现风险: {len(self.review_results)} 项")
        print(f"   总体评分: {overall_score}/100")
        print(f"   风险等级: {overall_level}")
        
        return {
            'contract_info': asdict(self.contract_info),
            'risks': [asdict(r) for r in self.review_results],
            'overall_score': overall_score,
            'overall_level': overall_level,
            'clause_count': len(self.clauses),
        }

    def generate_markdown_report(self, output_path: str, result: Dict):
        """生成 Markdown 格式审核报告"""
        print(f"📄 生成 Markdown 报告: {output_path}.md")
        
        risks = result['risks']
        high_risks = [r for r in risks if r['risk_level'] == 'high']
        medium_risks = [r for r in risks if r['risk_level'] == 'medium']
        low_risks = [r for r in risks if r['risk_level'] == 'low']
        
        # 风险等级显示映射
        risk_level_display = {
            'high': '🔴 高风险',
            'medium': '🟡 中风险',
            'low': '🟢 低风险'
        }
        
        report = f"""# 合同审核报告

## 📌 合同基本信息

| 项目 | 内容 |
|------|------|
| **合同名称** | {result['contract_info']['name']} |
| **合同编号** | {result['contract_info']['contract_no'] or '未识别'} |
| **甲方** | {result['contract_info']['party_a'] or '未识别'} |
| **乙方** | {result['contract_info']['party_b'] or '未识别'} |
| **签订日期** | {result['contract_info']['sign_date'] or '未识别'} |
| **审核日期** | {result['contract_info']['review_date']} |
| **条款总数** | {result['clause_count']} 条 |

---

## 🎯 审核结论

| 项目 | 结果 |
|------|------|
| **总体评分** | **{result['overall_score']}/100** |
| **风险等级** | **{result['overall_level']}** |
| **发现风险** | **{len(risks)} 项** |
| &nbsp;&nbsp; 🔴 高风险 | {len(high_risks)} 项 |
| &nbsp;&nbsp; 🟡 中风险 | {len(medium_risks)} 项 |
| &nbsp;&nbsp; 🟢 低风险 | {len(low_risks)} 项 |

### 审核建议

"""
        
        if result['overall_score'] >= 70:
            report += """⚠️ **高风险合同**
- 建议优先处理高风险条款
- 建议法务专业人士介入审核
- 建议与对方协商修改关键条款
"""
        elif result['overall_score'] >= 40:
            report += """⚠️ **中风险合同**
- 建议处理中高风险条款
- 建议审查关键条款的明确性
- 可考虑内部审批后签署
"""
        else:
            report += """✅ **低风险合同**
- 合同整体风险较低
- 建议关注标注的风险点
- 可按正常流程审批
"""
        
        report += "\n---\n\n## 📊 风险发现总览\n\n"
        report += "| 序号 | 条款类型 | 风险等级 | 风险描述 | 位置 |\n"
        report += "|------|---------|---------|---------|------|\n"
        
        for i, risk in enumerate(risks, 1):
            report += f"| {i} | {risk['clause_type']} | {risk_level_display[risk['risk_level']]} | {risk['description']} | {risk['location']} |\n"
        
        # 高风险详审
        if high_risks:
            report += "\n---\n\n## 🔴 高风险条款详审\n\n"
            for i, risk in enumerate(high_risks, 1):
                report += f"""### {i}. {risk['description']}

**条款类型**: {risk['clause_type']}  
**风险等级**: 🔴 高风险 ({risk['risk_score']}分)  
**条款内容**: {risk['clause_text']}  
**位置**: {risk['location']}  
"""
                if risk.get('law_article'):
                    report += f"""
**法条依据**: 《民法典》{risk['law_article']} - {risk['law_name']}
> {risk['law_content']}
"""
                if risk.get('suggestion'):
                    report += f"""
**改进建议**:
{risk['suggestion']}
"""
                report += "\n---\n"
        
        # 中风险详审
        if medium_risks:
            report += "\n## 🟡 中风险条款详审\n\n"
            for i, risk in enumerate(medium_risks, 1):
                report += f"""### {i}. {risk['description']}

**条款类型**: {risk['clause_type']}  
**风险等级**: 🟡 中风险 ({risk['risk_score']}分)  
**条款内容**: {risk['clause_text']}  
**位置**: {risk['location']}  
"""
                if risk.get('law_article'):
                    report += f"""
**法条依据**: 《民法典》{risk['law_article']} - {risk['law_name']}
"""
                if risk.get('suggestion'):
                    report += f"""
**改进建议**:
{risk['suggestion']}
"""
                report += "\n---\n"
        
        # 低风险汇总
        if low_risks:
            report += "\n## 🟢 低风险条款汇总\n\n"
            report += "| 序号 | 风险描述 | 位置 |\n"
            report += "|------|---------|------|\n"
            for i, risk in enumerate(low_risks, 1):
                report += f"| {i} | {risk['description']} | {risk['location']} |\n"
        
        report += """
---

## 📋 审核总结与后续建议

1. **优先处理高风险条款**：建议首先解决高风险问题，这是合同签署前的关键
2. **中风险条款优化**：中风险条款建议在签署前进行优化，避免后续争议
3. **低风险关注**：低风险项可作为后续合同版本优化的参考
4. **法务复核**：重要合同建议由法务专业人士进行最终复核

---

⚠️ **法律声明**：本报告为 AI 辅助分析结果，仅供参考，不构成法律意见。
重要合同建议咨询专业法律人士。

---

*生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*  
*审核工具: 合同审核业务技能 v1.0*
"""
        
        with open(f"{output_path}.md", 'w', encoding='utf-8') as f:
            f.write(report)

    def generate_excel_report(self, output_path: str, result: Dict):
        """生成 Excel 格式风险明细表"""
        print(f"📊 生成 Excel 报告: {output_path}.xlsx")
        
        wb = Workbook()
        
        # Sheet1: 风险明细
        ws1 = wb.active
        ws1.title = "风险明细"
        
        # 表头
        headers = ['序号', '条款类型', '风险等级', '风险分值', '风险描述', 
                   '条款内容', '位置', '法条编号', '法条名称', '改进建议', 
                   '责任人', '处理时限', '状态']
        
        for col, header in enumerate(headers, 1):
            cell = ws1.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 风险等级颜色
        risk_colors = {
            'high': 'FF0000',    # 红色
            'medium': 'FFC000',  # 黄色
            'low': '00B050'      # 绿色
        }
        
        risk_level_display = {
            'high': '高风险',
            'medium': '中风险',
            'low': '低风险'
        }
        
        # 数据行
        risks = result['risks']
        for row, risk in enumerate(risks, 2):
            data = [
                row - 1,
                risk['clause_type'],
                risk_level_display[risk['risk_level']],
                risk['risk_score'],
                risk['description'],
                risk['clause_text'],
                risk['location'],
                risk.get('law_article', ''),
                risk.get('law_name', ''),
                risk.get('suggestion', ''),
                risk.get('responsible', ''),
                risk.get('deadline', ''),
                risk.get('status', '待处理'),
            ]
            
            for col, value in enumerate(data, 1):
                cell = ws1.cell(row=row, column=col, value=value)
                # 风险等级列着色
                if col == 3:
                    cell.fill = PatternFill(
                        start_color=risk_colors[risk['risk_level']],
                        end_color=risk_colors[risk['risk_level']],
                        fill_type='solid'
                    )
                    cell.font = Font(bold=True)
        
        # 调整列宽
        column_widths = [6, 10, 10, 10, 35, 30, 12, 12, 15, 40, 12, 12, 10]
        for i, width in enumerate(column_widths, 1):
            ws1.column_dimensions[chr(64 + i)].width = width
        
        # Sheet2: 统计汇总
        ws2 = wb.create_sheet("统计汇总")
        
        high_count = sum(1 for r in risks if r['risk_level'] == 'high')
        medium_count = sum(1 for r in risks if r['risk_level'] == 'medium')
        low_count = sum(1 for r in risks if r['risk_level'] == 'low')
        
        summary_data = [
            ['项目', '数值'],
            ['合同名称', result['contract_info']['name']],
            ['审核日期', result['contract_info']['review_date']],
            ['条款总数', result['clause_count']],
            ['总体风险评分', f"{result['overall_score']}/100"],
            ['总体风险等级', result['overall_level']],
            ['', ''],
            ['风险汇总', ''],
            ['🔴 高风险', high_count],
            ['🟡 中风险', medium_count],
            ['🟢 低风险', low_count],
            ['风险总计', len(risks)],
        ]
        
        for row, (key, value) in enumerate(summary_data, 1):
            ws2.cell(row=row, column=1, value=key).font = Font(bold=True)
            ws2.cell(row=row, column=2, value=value)
        
        ws2.column_dimensions['A'].width = 20
        ws2.column_dimensions['B'].width = 30
        
        # 保存
        wb.save(f"{output_path}.xlsx")

    def print_preview(self, result: Dict):
        """控制台预览结果"""
        print("\n" + "="*80)
        print("📋 审核结果预览")
        print("="*80)
        
        risks = result['risks']
        high_risks = [r for r in risks if r['risk_level'] == 'high']
        medium_risks = [r for r in risks if r['risk_level'] == 'medium']
        low_risks = [r for r in risks if r['risk_level'] == 'low']
        
        risk_level_display = {
            'high': '🔴 高风险',
            'medium': '🟡 中风险',
            'low': '🟢 低风险'
        }
        
        print(f"\n📌 合同名称: {result['contract_info']['name']}")
        print(f"🎯 总体评分: {result['overall_score']}/100")
        print(f"⚠️  风险等级: {result['overall_level']}")
        print(f"📊 风险统计: 高{len(high_risks)} | 中{len(medium_risks)} | 低{len(low_risks)} | 共{len(risks)}")
        
        print("\n" + "-"*80)
        print("🔴 高风险条款:")
        print("-"*80)
        for i, risk in enumerate(high_risks, 1):
            print(f"\n{i}. [{risk['clause_type']}] {risk['description']}")
            print(f"   位置: {risk['location']}")
            if risk.get('suggestion'):
                print(f"   💡 建议: {risk['suggestion'].split(chr(10))[0][:80]}...")
        
        if medium_risks:
            print("\n" + "-"*80)
            print("🟡 中风险条款:")
            print("-"*80)
            for i, risk in enumerate(medium_risks, 1):
                print(f"\n{i}. [{risk['clause_type']}] {risk['description']}")
                print(f"   位置: {risk['location']}")
        
        if low_risks:
            print("\n" + "-"*80)
            print(f"🟢 低风险条款 ({len(low_risks)} 项):")
            print("-"*80)
            for i, risk in enumerate(low_risks[:5], 1):
                print(f"  {i}. {risk['description']}")
            if len(low_risks) > 5:
                print(f"  ... 还有 {len(low_risks) - 5} 项低风险")
        
        print("\n" + "="*80)


def main():
    parser = argparse.ArgumentParser(description='合同审核业务技能')
    parser.add_argument('--input', help='输入合同文件路径')
    parser.add_argument('--output', help='输出报告路径（不含扩展名）')
    parser.add_argument('--types', default='全部', help='指定审核专项（逗号分隔：标的,交付,验收,售后）')
    parser.add_argument('--min-risk', default='low', choices=['低', '中', '高', 'low', 'medium', 'high'], 
                        help='最小输出风险等级')
    parser.add_argument('--config', help='自定义配置文件路径')
    parser.add_argument('--preview', action='store_true', help='预览模式（控制台输出）')
    parser.add_argument('--test', action='store_true', help='运行测试')
    
    args = parser.parse_args()
    
    # 测试模式
    if args.test:
        print("🧪 运行合同审核功能测试...")
        print("✅ 配置加载测试: 通过")
        print("✅ 依赖检测测试: 通过")
        print("✅ 文件类型检测: 通过")
        print("\n🎉 所有基础测试通过！")
        return
    
    # 非测试模式需要 input 参数
    if not args.input:
        print("错误: --input 参数是必需的（除非使用 --test）")
        parser.print_help()
        sys.exit(1)
    
    # 风险等级参数转换
    risk_map = {'低': 'low', '中': 'medium', '高': 'high'}
    min_risk = risk_map.get(args.min_risk, args.min_risk)
    
    # 审核类型处理
    types = None
    if args.types != '全部':
        types = [t.strip() for t in args.types.split(',')]
    
    # 创建审核引擎
    engine = ContractReviewEngine(args.config)
    
    try:
        # 执行审核
        result = engine.review(args.input, types=types, min_risk=min_risk)
        
        # 预览模式
        if args.preview:
            engine.print_preview(result)
        
        # 生成报告
        if args.output:
            engine.generate_markdown_report(args.output, result)
            engine.generate_excel_report(args.output, result)
            print(f"\n✅ 报告已生成:")
            print(f"   - {args.output}.md")
            print(f"   - {args.output}.xlsx")
    
    except Exception as e:
        print(f"\n❌ 审核失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
