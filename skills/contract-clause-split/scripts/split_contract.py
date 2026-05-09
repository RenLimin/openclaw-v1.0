#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合同条款拆分工具 v2.0
=====================
100% Python 脚本实现，无需调用大模型

功能:
  - PDF / Word 合同文本解析
  - 智能按条款拆分
  - 10大分类自动匹配
  - 风险关键词识别
  - 关键信息（金额、日期、甲乙双方）提取
  - Excel 多 Sheet 输出
  - 签章页图片提取

作者: Ella 🦊
版本: v2.0
日期: 2026-04-23
"""

import os
import re
import sys
import argparse
import yaml
from pathlib import Path
from typing import List, Dict, Tuple, Optional

# 颜色输出
class Colors:
    RED = '\033[91m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    PURPLE = '\033[95m'
    CYAN = '\033[96m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'

# =============================================================================
# 配置加载模块
# =============================================================================

def load_config(config_dir: Path) -> Tuple[Dict, Dict, Dict]:
    """加载所有配置文件"""
    print(f"{Colors.CYAN}📖 加载配置文件...{Colors.ENDC}")
    
    config_files = [
        'classification-rules.yaml',
        'risk-keywords.yaml',
        'extract-rules.yaml'
    ]
    
    configs = {}
    for cfg_file in config_files:
        cfg_path = config_dir / cfg_file
        if cfg_path.exists():
            with open(cfg_path, 'r', encoding='utf-8') as f:
                configs[cfg_file.split('.')[0]] = yaml.safe_load(f)
            print(f"  ✅ {cfg_file}")
        else:
            print(f"  {Colors.YELLOW}⚠️  警告: 配置文件不存在 {cfg_file}{Colors.ENDC}")
    
    return (
        configs.get('classification-rules', {}),
        configs.get('risk-keywords', {}),
        configs.get('extract-rules', {})
    )

# =============================================================================
# 文档解析模块 (PDF + Word)
# =============================================================================

def extract_from_word(file_path: Path) -> str:
    """从 Word 文档提取文本（表格作为整体处理）"""
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = []
        
        # 提取普通段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 3:
                paragraphs.append(text)
        
        # 提取表格内容（整个表格作为一个整体！）
        for i, table in enumerate(doc.tables, 1):
            table_lines = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and len(cell_text) > 1:
                        row_cells.append(cell_text)
                
                if row_cells:
                    table_lines.append(' | '.join(row_cells))
            
            if table_lines:
                table_text = '\n'.join(table_lines)
                paragraphs.append(f"[表格{i}] {table_text}")
        
        full_text = '\n'.join(paragraphs)
        print(f"  ✅ 提取了 {len(paragraphs)} 段文本，共 {len(full_text)} 字符")
        print(f"     (含 {len(doc.tables)} 个表格，已作为整体处理)")
        return full_text
    except ImportError:
        print(f"  {Colors.RED}❌ 错误: 请安装 python-docx: pip install python-docx{Colors.ENDC}")
        sys.exit(1)
    except Exception as e:
        print(f"  {Colors.RED}❌ Word 解析失败: {e}{Colors.ENDC}")
        return ""

def extract_from_pdf(file_path: Path) -> str:
    """从 PDF 文档提取文本"""
    try:
        import pdfplumber
        paragraphs = []
        
        with pdfplumber.open(str(file_path)) as pdf:
            print(f"  📄 PDF 共 {len(pdf.pages)} 页")
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    for para in text.split('\n'):
                        para = para.strip()
                        if para and len(para) > 3:
                            paragraphs.append(para)
        
        full_text = '\n'.join(paragraphs)
        print(f"  ✅ 提取了 {len(paragraphs)} 段文本，共 {len(full_text)} 字符")
        return full_text
    except ImportError:
        print(f"  {Colors.YELLOW}⚠️  警告: 请安装 pdfplumber: pip install pdfplumber{Colors.ENDC}")
        print(f"  尝试使用 PyPDF2 备用方案...")
        try:
            from PyPDF2 import PdfReader
            paragraphs = []
            reader = PdfReader(str(file_path))
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    for para in text.split('\n'):
                        para = para.strip()
                        if para and len(para) > 3:
                            paragraphs.append(para)
            full_text = '\n'.join(paragraphs)
            print(f"  ✅ 使用 PyPDF2 提取了 {len(paragraphs)} 段文本，共 {len(full_text)} 字符")
            return full_text
        except ImportError:
            print(f"  {Colors.RED}❌ 错误: 请安装 PyPDF2: pip install pypdf2{Colors.ENDC}")
            sys.exit(1)
    except Exception as e:
        print(f"  {Colors.RED}❌ PDF 解析失败: {e}{Colors.ENDC}")
        return ""

def extract_text(file_path: Path) -> str:
    """根据文件扩展名选择解析器"""
    print(f"\n{Colors.CYAN}📂 解析文档: {file_path.name}{Colors.ENDC}")
    suffix = file_path.suffix.lower()
    
    if suffix == '.docx':
        return extract_from_word(file_path)
    elif suffix == '.pdf':
        return extract_from_pdf(file_path)
    else:
        print(f"  {Colors.RED}❌ 不支持的格式: {suffix}{Colors.ENDC}")
        print("  支持的格式: .pdf, .docx")
        sys.exit(1)

# =============================================================================
# 条款拆分模块
# =============================================================================

def smart_split_clauses(full_text: str) -> List[Dict]:
    """智能拆分合同条款（按条款序号边界拆分，同条款多行合并）"""
    print(f"\n{Colors.CYAN}✂️  条款智能拆分中...{Colors.ENDC}")
    
    # 条款序号匹配正则（覆盖常见格式）
    clause_start_patterns = [
        r'^第[一二三四五六七八九十\d]+[条章节款]',  # 第一条、第一章、第1条
        r'^\d+\.\s',                                    # 1.  2.  3.
        r'^\d+\.\d+\s',                               # 1.1  2.1  3.1
        r'^\d+\.\d+\.\d+\s',                       # 1.1.1
        r'^[（(][一二三四五六七八九十\d]+[）)]\s*',  # （一） (1) （1）
        r'^[一二三四五六七八九十]{1,2}[、．.]\s*',     # 一、 二、 三、
        r'^\[表格\d+\]',                              # [表格1] [表格2]
    ]
    
    combined_pattern = re.compile('|'.join(clause_start_patterns))
    
    lines = full_text.split('\n')
    clauses = []
    current_clause_lines = []
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 3:
            continue
        
        # 检查是否为新条款开始
        is_new_clause = bool(combined_pattern.match(line))
        
        if is_new_clause:
            # 保存上一个条款
            if current_clause_lines:
                clause_content = '\n'.join(current_clause_lines)
                clauses.append({
                    'id': len(clauses) + 1,
                    'content': clause_content,
                    'length': len(clause_content)
                })
            # 开始新条款
            current_clause_lines = [line]
        else:
            # 继续当前条款
            current_clause_lines.append(line)
    
    # 保存最后一个条款
    if current_clause_lines:
        clause_content = '\n'.join(current_clause_lines)
        clauses.append({
            'id': len(clauses) + 1,
            'content': clause_content,
            'length': len(clause_content)
        })
    
    print(f"  ✅ 拆分完成，共 {len(clauses)} 个条款（按序号边界合并）")
    return clauses

# =============================================================================
# 条款分类模块
# =============================================================================

def classify_clauses(clauses: List[Dict], class_config: Dict) -> List[Dict]:
    """对条款进行分类（依据《民法典》合同编定义）"""
    print(f"\n{Colors.CYAN}🏷️  条款分类中...{Colors.ENDC}")
    
    rules = class_config.get('classification_rules', [])
    default_cat = class_config.get('default_category', {'code': '99', 'name': '待分类'})
    
    categorized = []
    category_counts = {}
    total_clauses = len(clauses)
    
    for idx, clause in enumerate(clauses):
        content = clause['content']
        matched_cat = default_cat.copy()
        best_score = -1
        best_priority = 0
        
        is_early_content = (idx / max(total_clauses, 1)) < 0.25  # 合同前25%内容
        
        # ==========================================
        # 前置：强制匹配检查（最高优先级，命中直接匹配，跳过排除规则）
        # 例如：表格内容即使有"乙方应"等词，也优先判定为标的条款
        # ==========================================
        force_matched = None
        # 表格 → 强制标的条款（设备清单）
        if "[表格" in content or "设备清单" in content or "配置清单" in content:
            for cr in rules:
                if cr.get('name') == "标的条款":
                    force_matched = cr
                    break
        # 明确金额描述 → 强制价格条款
        elif "人民币" in content and "元整" in content and "违约金" not in content:
            for cr in rules:
                if cr.get('name') == "价格条款":
                    force_matched = cr
                    break
        # 甲乙双方信息头 → 强制主体信息
        elif ("甲方：" in content or "甲方（" in content or "乙方：" in content or "乙方（" in content) and idx < 5:
            for cr in rules:
                if cr.get('name') == "主体信息":
                    force_matched = cr
                    break
        
        if force_matched is not None:
            matched_cat = force_matched
            best_score = 999
            best_priority = 999
        else:
            # ==========================================
            # 正常匹配流程
            # ==========================================
            for cat_rule in rules:
                cat_name = cat_rule.get('name', '')
                
                # ==========================================
                # 排除规则（命中任意排除关键词，直接跳过此分类）
                # ==========================================
                exclude_keywords = cat_rule.get('exclude_keywords', [])
                should_exclude = False
                for kw in exclude_keywords:
                    if kw in content:
                        should_exclude = True
                        break
                if should_exclude:
                    continue
                
                # ==========================================
                # 计算匹配分数
                # ==========================================
                score = 0
                
                # 核心特征匹配（×3 高权重，法律定义核心）
                core_features = cat_rule.get('core_features', [])
                for feature in core_features:
                    if feature in content:
                        score += 3
                
                # 普通关键词匹配（×1 权重）
                keywords = cat_rule.get('keywords', [])
                for kw in keywords:
                    if kw in content:
                        score += 1
                
                # ==========================================
                # 位置和场景加权
                # ==========================================
                # 合同前25%内容，基本信息额外加分
                if is_early_content and cat_name == "基本信息":
                    score += 3
                # 合同前25%内容，价格/违约责任减分，避免开头金额/违约词误判
                if is_early_content and cat_name in ["价格条款", "违约责任"]:
                    score -= 2
                
                # ==========================================
                # 优先级判定（高优先级分类优先匹配）
                # ==========================================
                priority = cat_rule.get('priority', 0)
                
                # 对比：优先级高的直接胜出；同优先级看分数
                if (priority > best_priority) or (priority == best_priority and score > best_score):
                    best_score = score
                    best_priority = priority
                    matched_cat = cat_rule
            
            # ==========================================
            # 第二步：计算匹配分数
            # ==========================================
            score = 0
            
            # 核心特征匹配（×3 高权重，法律定义核心）
            core_features = cat_rule.get('core_features', [])
            for feature in core_features:
                if feature in content:
                    score += 3
            
            # 普通关键词匹配（×1 权重）
            keywords = cat_rule.get('keywords', [])
            for kw in keywords:
                if kw in content:
                    score += 1
            
            # ==========================================
            # 第三步：位置和场景加权
            # ==========================================
            # 合同前25%内容，基本信息额外加分
            if is_early_content and cat_name == "基本信息":
                score += 3
            # 合同前25%内容，价格/违约责任减分，避免开头金额/违约词误判
            if is_early_content and cat_name in ["价格条款", "违约责任"]:
                score -= 2
            
            # ==========================================
            # 第四步：优先级判定（高优先级分类优先匹配）
            # ==========================================
            priority = cat_rule.get('priority', 0)
            
            # 对比：优先级高的直接胜出；同优先级看分数
            if (priority > best_priority) or (priority == best_priority and score > best_score):
                best_score = score
                best_priority = priority
                matched_cat = cat_rule
        
        cat_name = matched_cat.get('name', '待分类')
        cat_code = matched_cat.get('code', '99')
        
        clause.update({
            'category': cat_name,
            'category_code': cat_code,
            'keyword_matches': best_score,
            'position': idx
        })
        
        categorized.append(clause)
        
        # 统计
        category_counts[cat_name] = category_counts.get(cat_name, 0) + 1
    
    # 输出统计
    print(f"  {Colors.PURPLE}📊 分类统计:{Colors.ENDC}")
    for cat_name, count in sorted(category_counts.items()):
        print(f"    {cat_name}: {count} 条")
    
    # ==========================================
    # 第一步：按分类聚类排序（相同分类放一起）
    # ==========================================
    print(f"\n{Colors.CYAN}📑 按分类聚类排序...{Colors.ENDC}")
    
    # 定义标准分类顺序（合同常规结构）
    standard_category_order = [
        "基本信息",
        "主体信息", 
        "标的条款",
        "价格条款",
        "履行条款",
        "权利义务",
        "违约责任",
        "保密与知识产权",
        "争议解决",
        "其他条款",
        "待分类"
    ]
    
    # 按分类分组
    grouped = {}
    for clause in categorized:
        cat = clause['category']
        if cat not in grouped:
            grouped[cat] = []
        grouped[cat].append(clause)
    
    # 按标准顺序重新排列所有条款
    sorted_clauses = []
    for cat in standard_category_order:
        if cat in grouped:
            sorted_clauses.extend(grouped[cat])
            del grouped[cat]
    # 剩余分类追加到末尾
    for cat, clauses in grouped.items():
        sorted_clauses.extend(clauses)
    
    # ==========================================
    # 第二步：按分类统一编号（同分类使用相同主编号）
    # ==========================================
    print(f"{Colors.CYAN}🏷️  按分类统一编号...{Colors.ENDC}")
    
    # 建立分类编号映射
    category_id_map = {}
    for clause in sorted_clauses:
        cat = clause['category']
        if cat not in category_id_map:
            category_id_map[cat] = len(category_id_map) + 1
    
    # 计数器
    category_sub_id = {cat: 0 for cat in category_id_map.keys()}
    
    # 重新编号
    for clause in sorted_clauses:
        cat = clause['category']
        main_id = category_id_map[cat]
        category_sub_id[cat] += 1
        sub_id = category_sub_id[cat]
        clause['id'] = f"{main_id}.{sub_id}"
        clause['category_id'] = main_id
    
    print(f"  ✅ 完成，共 {len(category_id_map)} 个分类，{len(sorted_clauses)} 条")
    print(f"     分类顺序：{', '.join(category_id_map.keys())}")
    
    return sorted_clauses

# =============================================================================
# 风险识别模块
# =============================================================================

def identify_risks(clauses: List[Dict], risk_config: Dict) -> List[Dict]:
    """识别条款中的风险"""
    print(f"\n{Colors.CYAN}⚠️  风险识别中...{Colors.ENDC}")
    
    risk_rules = risk_config.get('risk_rules', {})
    high_risk = risk_rules.get('high_risk', {}).get('keywords', [])
    medium_risk = risk_rules.get('medium_risk', {}).get('keywords', [])
    low_risk = risk_rules.get('low_risk', {}).get('keywords', [])
    
    risks_found = {'high': 0, 'medium': 0, 'low': 0}
    
    for clause in clauses:
        content = clause['content']
        clause_risk = []
        risk_level = "✅ 无风险"
        
        # 按优先级匹配风险关键词
        for kw in high_risk:
            if kw in content:
                clause_risk.append(kw)
                risk_level = "🔴 高风险"
        
        if risk_level == "✅ 无风险":
            for kw in medium_risk:
                if kw in content:
                    clause_risk.append(kw)
                    risk_level = "🟡 中风险"
        
        if risk_level == "✅ 无风险":
            for kw in low_risk:
                if kw in content:
                    clause_risk.append(kw)
                    risk_level = "🟢 低风险"
        
        clause['risk_level'] = risk_level
        clause['risk_keywords'] = ', '.join(clause_risk)
        
        if '高风险' in risk_level:
            risks_found['high'] += 1
        elif '中风险' in risk_level:
            risks_found['medium'] += 1
        elif '低风险' in risk_level:
            risks_found['low'] += 1
    
    print(f"  🔴 高风险: {risks_found['high']} 条")
    print(f"  🟡 中风险: {risks_found['medium']} 条")
    print(f"  🟢 低风险: {risks_found['low']} 条")
    
    return clauses

# =============================================================================
# 关键信息提取模块
# =============================================================================

def extract_key_info(full_text: str, extract_config: Dict) -> Dict[str, str]:
    """提取合同关键信息"""
    print(f"\n{Colors.CYAN}🔑 提取关键信息...{Colors.ENDC}")
    
    rules = extract_config.get('extraction_rules', {})
    config = extract_config.get('extraction_config', {})
    
    # 只在文档前 N% 搜索关键信息
    search_chars = min(config.get('search_max_chars', 5000), len(full_text))
    search_text = full_text[:search_chars]
    
    key_info = {}
    
    for rule_name, rule in rules.items():
        display_name = rule.get('name', rule_name)
        patterns = rule.get('patterns', [])
        found_value = None
        
        for pattern in patterns:
            match = re.search(pattern, search_text)
            if match:
                # 如果有分组提取
                if match.groups():
                    found_value = ' '.join(match.groups()).strip()
                else:
                    found_value = match.group(0).strip()
                break
        
        if found_value and len(found_value) > 2:
            key_info[display_name] = found_value
            print(f"  {Colors.GREEN}✅ {display_name}: {found_value[:50]}{Colors.ENDC}")
        else:
            key_info[display_name] = "未提取到"
            print(f"  ⚪ {display_name}: 未提取到")
    
    return key_info

# =============================================================================
# 签署页图片提取模块
# =============================================================================

def _extract_signature_page(pdf_path: Path) -> List[Path]:
    """提取PDF中的签署页为高清图片
    
    检测逻辑：
      1. 检查最后3页（签署页通常在最后）
      2. 检测关键词：签章、签字、盖章、法定代表人、授权代表等
      3. 命中则渲染为高清PNG（300 DPI）
      4. 保存到 ~/.openclaw/output/contract-images/
    
    返回：
      提取的图片路径列表（支持多签署页）
    """
    try:
        import pdfplumber
        from PIL import Image
        import io
    except ImportError as e:
        print(f"  {Colors.YELLOW}⚠️  签署页提取需要 pdfplumber 和 pillow: pip install pdfplumber pillow{Colors.ENDC}")
        print(f"     跳过签署页提取，继续处理其他内容")
        return []
    
    print(f"\n{Colors.CYAN}🖼️  提取签署页图片...{Colors.ENDC}")
    
    # 创建输出目录
    output_dir = Path.home() / '.openclaw' / 'output' / 'contract-images'
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 签署页检测关键词
    signature_keywords = [
        '签章', '签字', '盖章', '盖 章', '签 字',
        '法定代表人', '授权代表', '委托代理人',
        '单位盖章', '公司盖章', '（盖章）', '（签字）',
        '年 月 日', '年月日', '甲方（盖章）', '乙方（盖章）',
        '甲方签字', '乙方签字', '甲方签章', '乙方签章'
    ]
    
    extracted_images = []
    
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            total_pages = len(pdf.pages)
            # 只检查最后3页（签署页通常在最后）
            start_page = max(0, total_pages - 3)
            
            for page_idx in range(start_page, total_pages):
                page = pdf.pages[page_idx]
                page_num = page_idx + 1
                
                # 提取页面文本
                page_text = page.extract_text() or ''
                
                # 检测是否为签署页
                is_signature_page = False
                matched_keywords = []
                for kw in signature_keywords:
                    if kw in page_text:
                        is_signature_page = True
                        matched_keywords.append(kw)
                        if len(matched_keywords) >= 2:  # 命中2个以上关键词则确定为签署页
                            break
                
                if is_signature_page:
                    print(f"  ✅ 检测到第 {page_num} 页为签署页")
                    print(f"     命中关键词: {', '.join(matched_keywords[:3])}")
                    
                    try:
                        # 渲染为高清图片（300 DPI）
                        img = page.to_image(resolution=300)
                        
                        # 保存图片
                        img_filename = f"{pdf_path.stem}_签署页_{page_num}.png"
                        img_path = output_dir / img_filename
                        
                        # 使用 PIL 保存
                        if hasattr(img, 'original'):  # pdfplumber 的 PageImage 对象
                            img.original.save(str(img_path), 'PNG', optimize=True)
                        else:
                            img.save(str(img_path), 'PNG', optimize=True)
                        
                        extracted_images.append(img_path)
                        print(f"     已保存: {img_path}")
                        
                    except Exception as e:
                        print(f"  {Colors.YELLOW}⚠️  第 {page_num} 页渲染失败: {e}{Colors.ENDC}")
                        continue
        
        if extracted_images:
            print(f"  ✅ 共提取 {len(extracted_images)} 个签署页图片")
        else:
            print(f"  ⚪ 未检测到符合条件的签署页（最后3页未命中签署关键词）")
        
        return extracted_images
        
    except Exception as e:
        print(f"  {Colors.YELLOW}⚠️  签署页提取异常: {e}{Colors.ENDC}")
        return []

def _insert_image_to_excel(ws, image_path: Path, start_row: int = 2) -> int:
    """将图片插入到 Excel 工作表中
    
    参数:
      ws: openpyxl Worksheet 对象
      image_path: 图片文件路径
      start_row: 起始行号（默认从第2行开始）
    
    返回:
      图片占用的行数（用于下一张图片的位置）
    """
    try:
        from openpyxl.drawing.image import Image
        from PIL import Image as PILImage
    except ImportError:
        print(f"  {Colors.YELLOW}⚠️  Excel 插入图片需要 openpyxl 和 pillow{Colors.ENDC}")
        return start_row
    
    try:
        # 打开图片获取尺寸
        with PILImage.open(str(image_path)) as pil_img:
            original_width, original_height = pil_img.size
            
            # 计算缩放比例（Excel 单元格宽度约 8px，高度约 20px）
            # 目标宽度：80 个单元格宽度 ≈ 640px（150 DPI 下适应 A4 纸）
            target_width_px = 640
            scale = target_width_px / original_width
            target_height_px = int(original_height * scale)
            
            # 转换为 Excel 单位（行高：1 = 20px，列宽：1 = 8px）
            target_width_cells = target_width_px / 8  # 约 80 列
            target_height_rows = int(target_height_px / 20) + 2  # +2 边距
        
        # 创建 openpyxl 图片对象
        img = Image(str(image_path))
        
        # 设置图片尺寸（按比例缩放）
        img.width = target_width_px
        img.height = target_height_px
        
        # 设置图片位置（从 B 列开始，留出 A 列边距）
        img.anchor = f'B{start_row}'
        
        # 插入图片
        ws.add_image(img)
        
        # 返回下一张图片的起始行
        next_row = start_row + target_height_rows + 3  # +3 图片间距
        return next_row
        
    except Exception as e:
        print(f"  {Colors.YELLOW}⚠️  插入图片失败 {image_path.name}: {e}{Colors.ENDC}")
        return start_row + 10  # 出错时跳过一些行

# =============================================================================
# Excel 导出模块
# =============================================================================

def export_to_excel(clauses: List[Dict], key_info: Dict, 
                    output_file: Path, input_file: Path, full_text: str,
                    signature_images: List[Path] = None) -> None:
    """导出结果到 Excel 多 Sheet"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        print(f"\n{Colors.RED}❌ 错误: 请安装 openpyxl: pip install openpyxl{Colors.ENDC}")
        sys.exit(1)
    
    print(f"\n{Colors.CYAN}📊 导出到 Excel...{Colors.ENDC}")
    
    wb = openpyxl.Workbook()
    
    # 样式定义
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    risk_high_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
    risk_medium_fill = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid")
    risk_low_fill = PatternFill(start_color="008000", end_color="008000", fill_type="solid")
    
    # =========================================================================
    # Sheet 1: 全部条款
    ws1 = wb.active
    ws1.title = "全部条款"
    
    headers = ['条款编号', '条款分类', '风险等级', '风险关键词', '条款内容', '字数']
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    for row, clause in enumerate(clauses, 2):
        ws1.cell(row=row, column=1, value=clause['id'])
        ws1.cell(row=row, column=2, value=clause['category'])
        ws1.cell(row=row, column=3, value=clause['risk_level'])
        ws1.cell(row=row, column=4, value=clause['risk_keywords'])
        ws1.cell(row=row, column=5, value=clause['content'])
        ws1.cell(row=row, column=6, value=clause['length'])
        
        # 风险等级着色
        risk_level = clause['risk_level']
        if '高风险' in risk_level:
            ws1.cell(row=row, column=3).fill = risk_high_fill
        elif '中风险' in risk_level:
            ws1.cell(row=row, column=3).fill = risk_medium_fill
        elif '低风险' in risk_level:
            ws1.cell(row=row, column=3).fill = risk_low_fill
    
    # 调整列宽
    ws1.column_dimensions['A'].width = 10
    ws1.column_dimensions['B'].width = 15
    ws1.column_dimensions['C'].width = 12
    ws1.column_dimensions['D'].width = 20
    ws1.column_dimensions['E'].width = 80
    ws1.column_dimensions['F'].width = 8
    
    # 冻结首行
    ws1.freeze_panes = 'A2'
    
    # =========================================================================
    # Sheet 2: 分类视图
    ws2 = wb.create_sheet("分类视图")
    categories = sorted(list(set(c['category'] for c in clauses)))
    
    row = 1
    for cat in categories:
        # 分类标题
        ws2.cell(row=row, column=1, value=cat)
        ws2.cell(row=row, column=1).font = Font(bold=True, size=14)
        ws2.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        row += 1
        
        # 分类下的条款
        cat_clauses = [c for c in clauses if c['category'] == cat]
        for clause in cat_clauses:
            ws2.cell(row=row, column=1, value=f"条款 {clause['id']}")
            ws2.cell(row=row, column=2, value=clause['risk_level'])
            ws2.cell(row=row, column=3, value=clause['content'][:200] + ('...' if len(clause['content']) > 200 else ''))
            row += 1
        row += 1  # 分类间空行
    
    ws2.column_dimensions['A'].width = 12
    ws2.column_dimensions['B'].width = 12
    ws2.column_dimensions['C'].width = 100
    
    # =========================================================================
    # Sheet 3: 签署页图片
    ws_sign = wb.create_sheet("签署页", 2)  # 位置 2（0-based）即第3个Sheet
    
    # 签署页标题
    ws_sign.cell(row=1, column=2, value="📝 合同签署页图片")
    ws_sign.cell(row=1, column=2).font = Font(bold=True, size=14, color="4472C4")
    
    # 提取说明
    ws_sign.cell(row=2, column=2, value="提取说明：自动检测 PDF 最后3页，命中'签章/签字/盖章/法定代表人'等关键词的页面渲染为高清图片")
    ws_sign.cell(row=2, column=2).font = Font(italic=True, color="666666")
    
    # 插入签署页图片（支持多图片）
    if signature_images and len(signature_images) > 0:
        current_row = 5  # 从第5行开始插入图片
        for idx, img_path in enumerate(signature_images, 1):
            ws_sign.cell(row=current_row - 1, column=2, value=f"签署页 {idx}: {img_path.name}")
            ws_sign.cell(row=current_row - 1, column=2).font = Font(bold=True)
            current_row = _insert_image_to_excel(ws_sign, img_path, current_row)
        print(f"  ✅ 已插入 {len(signature_images)} 张签署页图片到 Excel")
    else:
        ws_sign.cell(row=5, column=2, value="⚠️ 未提取到签署页图片")
        ws_sign.cell(row=5, column=2).font = Font(color="FFA500")
        ws_sign.cell(row=6, column=2, value="可能原因：")
        ws_sign.cell(row=7, column=2, value="  1. 非 PDF 格式文件（仅支持 PDF）")
        ws_sign.cell(row=8, column=2, value="  2. 最后3页未命中签署关键词")
        ws_sign.cell(row=9, column=2, value="  3. pdfplumber 或 pillow 未安装")
    
    # 调整列宽（给图片留足够空间）
    ws_sign.column_dimensions['A'].width = 5
    ws_sign.column_dimensions['B'].width = 120
    
    # =========================================================================
    # Sheet 4: 风险条款汇总（原 Sheet 3 后移）
    ws3 = wb.create_sheet("风险条款汇总")
    
    risk_headers = ['条款编号', '风险等级', '风险关键词', '条款内容']
    for col, header in enumerate(risk_headers, 1):
        cell = ws3.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
    
    risk_clauses = [c for c in clauses if '高风险' in c['risk_level'] or '中风险' in c['risk_level']]
    for row, clause in enumerate(risk_clauses, 2):
        ws3.cell(row=row, column=1, value=clause['id'])
        ws3.cell(row=row, column=2, value=clause['risk_level'])
        ws3.cell(row=row, column=3, value=clause['risk_keywords'])
        ws3.cell(row=row, column=4, value=clause['content'])
        
        if '高风险' in clause['risk_level']:
            ws3.cell(row=row, column=2).fill = risk_high_fill
        elif '中风险' in clause['risk_level']:
            ws3.cell(row=row, column=2).fill = risk_medium_fill
    
    ws3.column_dimensions['A'].width = 10
    ws3.column_dimensions['B'].width = 12
    ws3.column_dimensions['C'].width = 30
    ws3.column_dimensions['D'].width = 100
    ws3.freeze_panes = 'A2'
    
    # =========================================================================
    # Sheet 4: 关键信息
    ws4 = wb.create_sheet("关键信息")
    
    ws4.cell(row=1, column=1, value="关键信息项")
    ws4.cell(row=1, column=2, value="内容")
    ws4.cell(row=1, column=1).font = header_font
    ws4.cell(row=1, column=2).font = header_font
    ws4.cell(row=1, column=1).fill = header_fill
    ws4.cell(row=1, column=2).fill = header_fill
    
    for row, (key, value) in enumerate(key_info.items(), 2):
        ws4.cell(row=row, column=1, value=key)
        ws4.cell(row=row, column=2, value=value)
    
    # 增加源文件信息
    ws4.cell(row=row+1, column=1, value="源文件")
    ws4.cell(row=row+1, column=2, value=str(input_file))
    ws4.cell(row=row+2, column=1, value="解析日期")
    ws4.cell(row=row+2, column=2, value=str(__import__('datetime').datetime.now()))
    
    ws4.column_dimensions['A'].width = 20
    ws4.column_dimensions['B'].width = 80
    
    # =========================================================================
    # Sheet 5: 原始文本
    ws5 = wb.create_sheet("原始文本")
    ws5.cell(row=1, column=1, value="合同原文（用于核对）")
    ws5.cell(row=1, column=1).font = header_font
    ws5.cell(row=1, column=1).fill = header_fill
    
    ws5.cell(row=2, column=1, value=full_text[:32767])  # Excel 单元格限制
    ws5.column_dimensions['A'].width = 120
    
    # 保存
    wb.save(str(output_file))
    print(f"  ✅ 已导出: {output_file}")
    print(f"  📋 包含工作表: {', '.join(wb.sheetnames)}")

# =============================================================================
# 内置自我测试模块
# =============================================================================

def self_test(clauses: List[Dict]) -> Dict:
    """内置自我测试，检查输出质量（4项基本校验）
    
    返回测试报告字典，包含：
      - passed: 布尔值，是否全部通过
      - results: 各项测试结果详情
    """
    print(f"\n{Colors.YELLOW}{'='*60}{Colors.ENDC}")
    print(f"  🧪 内置自我测试检验")
    print(f"{Colors.YELLOW}{'='*60}{Colors.ENDC}")
    
    all_pass = True
    results = {}
    
    # 检查1: 条款数量（正常合同应在 10-25 条之间）
    clause_count = len(clauses)
    if 10 <= clause_count <= 25:
        print(f"  ✅ 条款数量检查通过: {clause_count} 条 (正常范围 10-25)")
        results['clause_count'] = {'status': 'pass', 'value': clause_count, 'expected': '10-25'}
    else:
        print(f"  {Colors.RED}❌ 条款数量异常: {clause_count} 条{Colors.ENDC}")
        results['clause_count'] = {'status': 'fail', 'value': clause_count, 'expected': '10-25'}
        all_pass = False
    
    # 检查2: 分类多样性（应有 ≥3 个分类）
    categories = set(c['category'] for c in clauses)
    cat_count = len(categories)
    if cat_count >= 3:
        print(f"  ✅ 分类多样性检查通过: {cat_count} 个分类")
        print(f"     分类列表: {', '.join(categories)}")
        results['category_diversity'] = {'status': 'pass', 'value': cat_count, 'expected': '≥3', 'categories': list(categories)}
    else:
        print(f"  {Colors.YELLOW}⚠️  分类数量偏少: {cat_count} 个{Colors.ENDC}")
        results['category_diversity'] = {'status': 'warning', 'value': cat_count, 'expected': '≥3', 'categories': list(categories)}
    
    # 检查3: 平均条款长度（应 > 50 字符）
    avg_len = sum(c['length'] for c in clauses) / len(clauses) if clauses else 0
    if avg_len > 50:
        print(f"  ✅ 条款长度检查通过: 平均 {int(avg_len)} 字符")
        results['avg_length'] = {'status': 'pass', 'value': int(avg_len), 'expected': '>50'}
    else:
        print(f"  {Colors.YELLOW}⚠️  条款偏短: 平均 {int(avg_len)} 字符{Colors.ENDC}")
        results['avg_length'] = {'status': 'warning', 'value': int(avg_len), 'expected': '>50'}
    
    # 检查4: 风险识别覆盖率（应有 ≥1 条风险条款）
    risk_count = sum(1 for c in clauses if '风险' in c['risk_level'] and '无风险' not in c['risk_level'])
    if risk_count > 0:
        print(f"  ✅ 风险识别检查通过: {risk_count} 条条款命中风险")
        results['risk_coverage'] = {'status': 'pass', 'value': risk_count, 'expected': '≥1'}
    else:
        print(f"  {Colors.YELLOW}⚠️  未命中任何风险条款{Colors.ENDC}")
        results['risk_coverage'] = {'status': 'warning', 'value': risk_count, 'expected': '≥1'}
    
    # 最终结果
    if all_pass:
        print(f"\n{Colors.GREEN}自我测试完成！✅ 全部通过{Colors.ENDC}")
    else:
        print(f"\n{Colors.YELLOW}自我测试完成！⚠️ 部分需人工核对{Colors.ENDC}")
    
    return {
        'passed': all_pass,
        'results': results,
        'summary': '全部通过' if all_pass else '部分需人工核对'
    }

# =============================================================================
# 预览输出模块
# =============================================================================

def print_preview(clauses: List[Dict], key_info: Dict):
    """在终端输出预览结果"""
    print(f"\n{Colors.PURPLE}{Colors.BOLD}" + "="*80)
    print(f"📋 合同条款拆分 - 结果预览")
    print("="*80 + Colors.ENDC)
    
    # 关键信息
    print(f"\n{Colors.CYAN}🔑 关键信息:{Colors.ENDC}")
    for k, v in key_info.items():
        print(f"  {k}: {v}")
    
    # 高风险条款
    high_risks = [c for c in clauses if '高风险' in c['risk_level']]
    if high_risks:
        print(f"\n{Colors.RED}🔴 高风险条款预览（前 5 条）:{Colors.ENDC}")
        for clause in high_risks[:5]:
            print(f"\n  [条款 {clause['id']}] 风险关键词: {clause['risk_keywords']}")
            print(f"  {Colors.BOLD}{clause['content'][:150]}{Colors.ENDC}...")
    
    # 中风险条款
    medium_risks = [c for c in clauses if '中风险' in c['risk_level']]
    if medium_risks:
        print(f"\n{Colors.YELLOW}🟡 中风险条款预览（前 5 条）:{Colors.ENDC}")
        for clause in medium_risks[:5]:
            print(f"\n  [条款 {clause['id']}] 风险关键词: {clause['risk_keywords']}")
            print(f"  {clause['content'][:150]}...")
    
    print(f"\n{Colors.GREEN}{Colors.BOLD}✅ 处理完成！{Colors.ENDC}")
    print(f"\n{Colors.CYAN}💡 提示:{Colors.ENDC}")
    print(f"  1. 请打开 Excel 文件查看完整结果")
    print(f"  2. 核对「关键信息」Sheet 是否正确识别了合同要素")
    print(f"  3. 重点关注「风险条款汇总」Sheet 中的高风险条款")
    print(f"  4. 条款分类规则可在 config/classification-rules.yaml 中调整")

# =============================================================================
# 主程序
# =============================================================================

def main():
    banner = f"""
{Colors.CYAN}╔══════════════════════════════════════════════════════════════════╗
║                                                                  ║
║          合同条款拆分工具 v2.0                                    ║
║                                                                  ║
║     100% Python 脚本实现 - 无需大模型调用                          ║
║                                                                  ║
╚══════════════════════════════════════════════════════════════════╝{Colors.ENDC}
"""
    print(banner)
    
    parser = argparse.ArgumentParser(
        description='合同条款智能拆分工具',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument('--input', '-i', required=True,
                      help='输入合同文件路径 (PDF/DOCX)')
    parser.add_argument('--output', '-o', default=None,
                      help='输出 Excel 文件路径（默认: 原文件名_解析结果.xlsx）')
    parser.add_argument('--config-dir', '-c', default=None,
                      help='配置文件目录（默认: 脚本所在目录/../config）')
    parser.add_argument('--preview', '-p', action='store_true',
                      help='在终端输出预览结果')
    parser.add_argument('--risk-check', action='store_true',
                      help='启用风险条款识别（默认启用）')
    parser.add_argument('--extract-signature', action='store_true', default=True,
                      help='提取签章页为图片（默认启用，仅支持 PDF）')
    parser.add_argument('--no-signature', action='store_true',
                      help='跳过签署页图片提取')
    parser.add_argument('--verbose', '-v', action='store_true',
                      help='显示详细处理日志')
    
    args = parser.parse_args()
    
    # 路径处理
    input_file = Path(args.input).expanduser().resolve()
    if not input_file.exists():
        print(f"{Colors.RED}❌ 错误: 输入文件不存在 {input_file}{Colors.ENDC}")
        sys.exit(1)
    
    # 输出文件
    if args.output:
        output_file = Path(args.output).expanduser().resolve()
    else:
        output_file = input_file.parent / f"{input_file.stem}_解析结果.xlsx"
    
    # 配置目录
    if args.config_dir:
        config_dir = Path(args.config_dir).expanduser().resolve()
    else:
        config_dir = Path(__file__).parent.parent / 'config'
    
    # 加载配置
    class_config, risk_config, extract_config = load_config(config_dir)
    
    # 文本提取
    full_text = extract_text(input_file)
    if not full_text:
        print(f"{Colors.RED}❌ 未提取到文本内容{Colors.ENDC}")
        sys.exit(1)
    
    # 条款拆分
    clauses = smart_split_clauses(full_text)
    
    # 条款分类
    clauses = classify_clauses(clauses, class_config)
    
    # 风险识别
    clauses = identify_risks(clauses, risk_config)
    
    # 内置自我测试
    test_report = self_test(clauses)
    
    # 关键信息提取
    key_info = extract_key_info(full_text, extract_config)
    
    # 签署页图片提取（仅 PDF，默认启用）
    signature_images = []
    if args.extract_signature and not args.no_signature:
        if input_file.suffix.lower() == '.pdf':
            signature_images = _extract_signature_page(input_file)
        else:
            print(f"\n{Colors.CYAN}ℹ️  签署页提取仅支持 PDF 格式，跳过{Colors.ENDC}")
    
    # 导出 Excel
    export_to_excel(clauses, key_info, output_file, input_file, full_text, signature_images)
    
    # 终端预览
    if args.preview:
        print_preview(clauses, key_info)
    
    print(f"\n{Colors.GREEN}🎉 全部处理完成！{Colors.ENDC}")
    print(f"   📂 输出文件: {output_file}")

if __name__ == '__main__':
    main()
