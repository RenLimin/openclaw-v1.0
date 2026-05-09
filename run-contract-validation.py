#!/usr/bin/env python3
"""
合同条款拆分 - 人工验证测试
生成可读的验证报告
"""
import json
from docx import Document
from pathlib import Path

print("=" * 80)
print("📄 合同条款拆分 - 人工验证测试")
print("=" * 80)

# 1. 加载配置
config_path = Path(__file__).parent / 'skills/contract-clause-split/config/contract-split-config.json'
with open(config_path, 'r', encoding='utf-8') as f:
    config = json.load(f)

# 2. 加载测试合同
test_file = '/Users/bangcle/Downloads/contract/XSZS2603090130北京聚信得仁 - 采购合同（设备）-移动应用合格平台检测设备采购项目-聚信得仁-梆梆.docx'
print(f"\n📂 测试文件: {Path(test_file).name}")

doc = Document(test_file)
full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

print(f"📝 文档总字数: {len(full_text)} 字")
print(f"📋 段落总数: {len(doc.paragraphs)}")

# 3. 显示合同前 1000 字预览
print("\n" + "=" * 80)
print("📖 合同内容预览（前 1000 字）:")
print("=" * 80)
print(full_text[:1000])
print("\n... (内容截断，完整合同已加载)")

# 4. 条款拆分演示（按段落拆分）
print("\n" + "=" * 80)
print("✂️  条款拆分演示（前 15 个非空段落）:")
print("=" * 80)

valid_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip() and len(p.text.strip()) > 10]

for i, para in enumerate(valid_paragraphs[:15], 1):
    category = "待分类"
    
    # 简单规则匹配演示
    for cat_key, cat_val in config['clause_categories'].items():
        for keyword in cat_val.get('keywords', []):
            if keyword in para:
                category = cat_val['name']
                break
    
    print(f"\n【条款 {i}】【{category}】")
    print(f"   {para[:120]}{'...' if len(para) > 120 else ''}")

# 5. 自动识别关键信息
print("\n" + "=" * 80)
print("🔍 关键信息自动提取（基于规则匹配）:")
print("=" * 80)

# 提取合同编号
contract_no = ""
for para in valid_paragraphs:
    if "合同编号" in para or "编号：" in para:
        contract_no = para
        break

# 提取甲方乙方
party_a = ""
party_b = ""
for para in valid_paragraphs[:20]:
    if "甲方" in para and ("需方" in para or "买方" in para):
        party_a = para
    if "乙方" in para and ("供方" in para or "卖方" in para):
        party_b = para

# 提取金额
amount = ""
for para in valid_paragraphs[:30]:
    if ("元" in para and ("万" in para or "仟" in para or "佰" in para)) or "￥" in para or "价格" in para:
        if len(para) < 100:
            amount = para
            break

print(f"📌 合同编号/名称: {contract_no[:80] if contract_no else '未识别到'}")
print(f"👤 甲方信息: {party_a[:80] if party_a else '未识别到'}")
print(f"👤 乙方信息: {party_b[:80] if party_b else '未识别到'}")
print(f"💰 合同金额相关: {amount[:100] if amount else '未识别到'}")

# 6. 风险条款识别演示
print("\n" + "=" * 80)
print("⚠️  潜在风险条款识别（基于关键词匹配）:")
print("=" * 80)

risk_keywords = config['risk_keywords']
risk_found = []

for para in valid_paragraphs:
    for level, keywords in risk_keywords.items():
        for kw in keywords:
            if kw in para:
                risk_found.append((level, kw, para[:80] + "..."))
                break

if risk_found:
    for level, kw, content in risk_found[:10]:
        print(f"   [{level.upper()}] 关键词 '{kw}':")
        print(f"        {content}\n")
else:
    print("   ✅ 未识别到明显风险关键词")

# 7. 统计信息
print("\n" + "=" * 80)
print("📊 拆分统计:")
print("=" * 80)
print(f"   总段落数: {len(doc.paragraphs)}")
print(f"   有效条款数: {len(valid_paragraphs)}")
print(f"   识别到风险条款: {len(risk_found)} 条")
print(f"   条款分类规则数: {len(config['clause_categories'])}")

print("\n" + "=" * 80)
print("✅ 合同条款拆分演示完成！")
print("=" * 80)
print("\n💡 人工验证要点:")
print("   1. 合同内容提取是否完整正确")
print("   2. 条款拆分边界是否合理")
print("   3. 关键信息（甲乙双方、金额）是否正确识别")
print("   4. 风险条款识别是否准确")
