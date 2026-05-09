#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OCR 引擎统一入口 v2.0
功能：多格式支持（PDF/Word/图片）、多引擎调度、自动检测类型、统一输出格式、质量评估集成

新增功能：
  - PDFProcessor: 自动检测文本版/扫描版 PDF，自动选择最优提取方式
  - WordProcessor: 直接读取 Word 文档，支持段落+表格+图片提取
  - ImageProcessor: 图片预处理（去噪/增强/二值化/方向校正），批量处理
  - 统一 JSON 输出格式：text、pages、tables、images、metadata、quality_scores、structure_tags

作者: OCR Engine Team
版本: v2.0
日期: 2026-04-24
"""

import os
import re
import sys
import time
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
import yaml

# 导入子模块
from noise_reducer import NoiseReducer
from quality_assess import QualityAssessor
from table_extractor import TableExtractor


class OCREngineError(Exception):
    """OCR 引擎异常基类"""
    pass


class EngineNotFoundError(OCREngineError):
    """引擎未找到异常"""
    pass


# =============================================================================
# 数据结构定义
# =============================================================================

@dataclass
class PageContent:
    """单页内容"""
    page_num: int
    text: str = ""
    image_path: str = ""
    quality_score: float = 0.0
    structure_tags: List[str] = None


@dataclass
class TableContent:
    """表格内容"""
    table_index: int
    page_num: int
    headers: List[str] = None
    rows: List[List[str]] = None
    raw_data: Dict = None


@dataclass
class ImageContent:
    """图片内容"""
    image_index: int
    page_num: int
    path: str
    ocr_text: str = ""
    quality_score: float = 0.0


@dataclass
class QualityScores:
    """质量评分"""
    overall_score: float = 0.0
    quality_level: str = ""
    page_scores: Dict[int, float] = None
    character_recognition_rate: float = 0.0
    noise_level: str = ""
    table_completeness: float = 0.0


# =============================================================================
# 引擎可用性检测
# =============================================================================

def check_tesseract_available() -> Tuple[bool, str]:
    """检测 Tesseract OCR 是否可用"""
    try:
        import pytesseract
        version = pytesseract.get_tesseract_version()
        return True, f"Tesseract v{version}"
    except ImportError:
        return False, "pytesseract 未安装 (pip install pytesseract)"
    except Exception as e:
        return False, str(e)


def check_poppler_available() -> Tuple[bool, str]:
    """检测 Poppler（pdftoppm）是否可用"""
    if shutil.which('pdftoppm'):
        return True, "Poppler 可用"
    return False, "Poppler 未安装 (brew install poppler 或 apt install poppler-utils)"


def check_pymupdf_available() -> Tuple[bool, str]:
    """检测 PyMuPDF (fitz) 是否可用"""
    try:
        import fitz
        return True, f"PyMuPDF v{fitz.__version__}"
    except ImportError:
        return False, "PyMuPDF 未安装 (pip install pymupdf)"
    except Exception as e:
        return False, str(e)


def check_pdfplumber_available() -> Tuple[bool, str]:
    """检测 pdfplumber 是否可用"""
    try:
        import pdfplumber
        return True, f"pdfplumber v{pdfplumber.__version__}"
    except ImportError:
        return False, "pdfplumber 未安装 (pip install pdfplumber)"
    except Exception as e:
        return False, str(e)


def check_python_docx_available() -> Tuple[bool, str]:
    """检测 python-docx 是否可用"""
    try:
        from docx import Document
        return True, "python-docx 可用"
    except ImportError:
        return False, "python-docx 未安装 (pip install python-docx)"
    except Exception as e:
        return False, str(e)


def check_pillow_available() -> Tuple[bool, str]:
    """检测 Pillow 图像处理库是否可用"""
    try:
        from PIL import Image, ImageEnhance, ImageFilter
        return True, "Pillow 可用"
    except ImportError:
        return False, "Pillow 未安装 (pip install pillow)"
    except Exception as e:
        return False, str(e)


def check_easyocr_available() -> Tuple[bool, str]:
    """检测 EasyOCR 是否可用"""
    try:
        import easyocr
        return True, "EasyOCR 可用"
    except ImportError:
        return False, "EasyOCR 未安装 (pip install easyocr)"
    except Exception as e:
        return False, str(e)


def check_all_dependencies() -> Dict:
    """检查所有依赖项"""
    results = {
        'tesseract': check_tesseract_available(),
        'poppler': check_poppler_available(),
        'pymupdf': check_pymupdf_available(),
        'pdfplumber': check_pdfplumber_available(),
        'python_docx': check_python_docx_available(),
        'pillow': check_pillow_available(),
        'easyocr': check_easyocr_available(),
    }

    # 计算总体状态
    core_deps = ['tesseract', 'poppler', 'pymupdf', 'python_docx', 'pillow']
    all_good = all(v[0] for k, v in results.items() if k in core_deps)

    return {
        'all_good': all_good,
        'dependencies': results,
        'engines_available': {
            'tesseract': results['tesseract'][0],
            'easyocr': results['easyocr'][0],
            'paddleocr': False  # 暂未实现
        }
    }


# =============================================================================
# PDF 处理器类
# =============================================================================

class PDFProcessor:
    """PDF 文档处理器 - 自动处理文本版和扫描版 PDF"""

    def __init__(self, engine: str = 'tesseract', lang: str = 'chi_sim+eng'):
        self.engine = engine
        self.lang = lang
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖"""
        pymupdf_ok, _ = check_pymupdf_available()
        poppler_ok, _ = check_poppler_available()
        tesseract_ok, _ = check_tesseract_available()

        if not pymupdf_ok:
            raise EngineNotFoundError("PyMuPDF 未安装，无法处理 PDF")
        if not poppler_ok or not tesseract_ok:
            print("⚠️  扫描版 PDF 处理依赖缺失，仅支持文本版 PDF")

    def detect_pdf_type(self, pdf_path: Path) -> Tuple[str, Dict]:
        """
        检测 PDF 类型

        Returns:
            ('text' | 'scanned' | 'mixed', 详细信息)
        """
        import fitz
        doc = fitz.open(str(pdf_path))
        total_chars = 0
        page_chars = []
        text_pages = 0

        for page in doc:
            text = page.get_text()
            char_count = len(text.strip())
            total_chars += char_count
            page_chars.append(char_count)
            if char_count > 100:
                text_pages += 1

        avg_chars = total_chars / max(len(doc), 1)
        text_ratio = text_pages / max(len(doc), 1)

        details = {
            'total_pages': len(doc),
            'total_chars': total_chars,
            'avg_chars_per_page': avg_chars,
            'text_pages': text_pages,
            'text_ratio': text_ratio,
            'page_chars': page_chars
        }

        # 判断类型
        if text_ratio >= 0.9:
            return 'text', details
        elif text_ratio <= 0.1:
            return 'scanned', details
        else:
            return 'mixed', details

    def extract_text_pdf(self, pdf_path: Path) -> Dict:
        """从文本版 PDF 提取文字（使用 pdfplumber 获得更好效果）"""
        try:
            import pdfplumber
        except ImportError:
            import fitz
            # fallback to PyMuPDF
            doc = fitz.open(str(pdf_path))
            pages = []
            full_text = ""

            for i, page in enumerate(doc):
                text = page.get_text()
                pages.append(PageContent(
                    page_num=i + 1,
                    text=text,
                    quality_score=95.0,
                    structure_tags=['text']
                ))
                full_text += text + "\n"

            return {
                'text': full_text,
                'pages': [asdict(p) for p in pages],
                'method': 'pymupdf_text_extract'
            }

        # 使用 pdfplumber
        pages = []
        full_text = ""
        tables = []

        with pdfplumber.open(str(pdf_path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages.append(PageContent(
                    page_num=i + 1,
                    text=text,
                    quality_score=95.0,
                    structure_tags=['text']
                ))
                full_text += text + "\n"

                # 尝试提取 PDF 中的表格
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if table and len(table) > 0:
                        tables.append(TableContent(
                            table_index=len(tables),
                            page_num=i + 1,
                            headers=table[0] if table else [],
                            rows=table[1:] if len(table) > 1 else [],
                            raw_data={'source': 'pdfplumber'}
                        ))

        return {
            'text': full_text,
            'pages': [asdict(p) for p in pages],
            'tables': [asdict(t) for t in tables],
            'method': 'pdfplumber_text_extract'
        }

    def ocr_scanned_pdf(self, pdf_path: Path, dpi: int = 300) -> Dict:
        """使用 OCR 处理扫描版 PDF"""
        from pdf2image import convert_from_path
        import pytesseract

        start_time = time.time()

        # PDF 转图片
        print(f"  🖨️  PDF 转图片中 (DPI={dpi}) ...")
        pages = convert_from_path(str(pdf_path), dpi=dpi)
        print(f"  ✅ 共 {len(pages)} 页")

        # OCR 逐页识别
        print(f"  🤖  OCR 识别中 ({self.lang}) ...")
        page_contents = []
        full_text = ""

        for i, page_img in enumerate(pages, 1):
            print(f"  🔍  识别第 {i}/{len(pages)} 页 ...", end='\r')
            text = pytesseract.image_to_string(page_img, lang=self.lang)

            # 简单质量评估：字符数越多质量越高
            quality = min(100.0, len(text.strip()) / 2)

            page_contents.append(PageContent(
                page_num=i,
                text=text,
                quality_score=quality,
                structure_tags=['ocr']
            ))
            full_text += text + "\n"

        print(f"\n  ✅ OCR 完成")

        return {
            'text': full_text,
            'pages': [asdict(p) for p in page_contents],
            'method': f'{self.engine}_ocr',
            'processing_time': time.time() - start_time
        }

    def process(self, pdf_path: Path) -> Dict:
        """处理 PDF 文档 - 自动选择最优方式"""
        pdf_type, pdf_info = self.detect_pdf_type(pdf_path)

        if pdf_type == 'text':
            print(f"  ✅ 检测为【文本版 PDF】，直接提取文字")
            result = self.extract_text_pdf(pdf_path)
        elif pdf_type == 'scanned':
            print(f"  ⚠️  检测为【扫描版 PDF】，启动 OCR 引擎")
            result = self.ocr_scanned_pdf(pdf_path)
        else:
            print(f"  ⚠️  检测为【混合格式 PDF】，优先提取文本，扫描页使用 OCR")
            # 混合模式：先尝试文本提取，效果不好再用 OCR
            text_result = self.extract_text_pdf(pdf_path)
            avg_chars = pdf_info['avg_chars_per_page']
            if avg_chars > 50:
                result = text_result
            else:
                result = self.ocr_scanned_pdf(pdf_path)

        result['pdf_type'] = pdf_type
        result['pdf_info'] = pdf_info
        return result


# =============================================================================
# Word 处理器类
# =============================================================================

class WordProcessor:
    """Word 文档处理器 - 支持段落、表格、图片提取"""

    def __init__(self):
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖"""
        docx_ok, _ = check_python_docx_available()
        if not docx_ok:
            raise EngineNotFoundError("python-docx 未安装，无法处理 Word 文档")

    def detect_heading_level(self, para_text: str, style_name: str = "") -> Optional[int]:
        """检测标题级别"""
        # 先检查 Word 内置标题样式
        if style_name.startswith('Heading') or style_name.startswith('标题'):
            try:
                level_str = re.sub(r'[^0-9]', '', style_name)
                if level_str:
                    return int(level_str)
            except ValueError:
                pass

        # 正则匹配标题模式
        heading_patterns = {
            1: [r'^第[一二三四五六七八九十百千\d]+章\s', r'^第[一二三四五六七八九十百千\d]+编\s'],
            2: [r'^第[一二三四五六七八九十百千\d]+条\s', r'^第[一二三四五六七八九十百千\d]+节\s', r'^\d+\.\s'],
            3: [r'^\d+\.\d+\s', r'^[（(][一二三四五六七八九十\d]+[）)]\s*'],
            4: [r'^\d+\.\d+\.\d+\s', r'^[一二三四五六七八九十]{1,2}[、．.]\s*'],
            5: [r'^\d+\.\d+\.\d+\.\d+\s'],
        }

        text = para_text.strip()
        for level, patterns in heading_patterns.items():
            for pattern in patterns:
                if re.match(pattern, text):
                    return level

        return None

    def extract_images_from_doc(self, doc, temp_dir: Path) -> List[Dict]:
        """从 Word 文档中提取图片并进行 OCR"""
        images = []
        try:
            import pytesseract
            from PIL import Image as PILImage

            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    image_path = temp_dir / f"word_img_{len(images)}.png"

                    with open(image_path, 'wb') as f:
                        f.write(image_data)

                    # 对图片进行 OCR
                    try:
                        pil_img = PILImage.open(image_path)
                        ocr_text = pytesseract.image_to_string(pil_img, lang='chi_sim+eng')
                        quality = min(100.0, len(ocr_text.strip()) / 2)
                    except Exception:
                        ocr_text = ""
                        quality = 0.0

                    images.append(asdict(ImageContent(
                        image_index=len(images),
                        page_num=0,  # Word 图片无页码信息
                        path=str(image_path),
                        ocr_text=ocr_text,
                        quality_score=quality
                    )))
        except Exception as e:
            print(f"  ⚠️  图片提取部分失败: {e}")

        return images

    def process(self, word_path: Path) -> Dict:
        """处理 Word 文档"""
        from docx import Document

        print(f"  📄 解析 Word 文档: {word_path.name}")

        doc = Document(str(word_path))
        temp_dir = Path(tempfile.mkdtemp(prefix='wordocr_'))

        # 提取段落
        pages = []
        full_text = ""
        structure_tags = []

        current_page_text = ""
        para_count = 0

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            heading_level = self.detect_heading_level(text, para.style.name if para.style else "")
            if heading_level:
                structure_tags.append(f"heading_{heading_level}")
                tag = f"[H{heading_level}]"
            else:
                tag = "[P]"

            current_page_text += f"{tag} {text}\n"
            full_text += f"{text}\n"
            para_count += 1

            # Word 无明确分页，每 50 段虚拟一页
            if para_count % 50 == 0:
                page_num = para_count // 50
                pages.append(asdict(PageContent(
                    page_num=page_num,
                    text=current_page_text,
                    quality_score=100.0,
                    structure_tags=structure_tags.copy()
                )))
                current_page_text = ""
                structure_tags = []

        # 最后一页
        if current_page_text:
            pages.append(asdict(PageContent(
                page_num=max(1, len(pages) + 1),
                text=current_page_text,
                quality_score=100.0,
                structure_tags=structure_tags
            )))

        # 提取表格（复用 word-analyzer 逻辑）
        tables = []
        for table_idx, table in enumerate(doc.tables):
            headers = []
            rows_data = []

            for row_idx, row in enumerate(table.rows):
                row_text = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = row_text
                else:
                    rows_data.append(row_text)

            tables.append(asdict(TableContent(
                table_index=table_idx,
                page_num=0,  # Word 表格无页码信息
                headers=headers,
                rows=rows_data,
                raw_data={'source': 'python_docx'}
            )))

        # 提取图片
        images = self.extract_images_from_doc(doc, temp_dir)

        # 元数据
        core_props = doc.core_properties
        metadata = {
            'title': core_props.title or "",
            'author': core_props.author or "",
            'created': str(core_props.created) if core_props.created else "",
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'image_count': len(images)
        }

        print(f"  ✅ 提取完成: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格, {len(images)} 图片")

        return {
            'text': full_text,
            'pages': pages,
            'tables': tables,
            'images': images,
            'metadata': metadata,
            'method': 'python_docx_parse',
            'temp_dir': str(temp_dir)
        }


# =============================================================================
# 图片处理器类
# =============================================================================

class ImageProcessor:
    """图片文档处理器 - 支持预处理、方向校正、批量处理"""

    def __init__(self, engine: str = 'tesseract', lang: str = 'chi_sim+eng'):
        self.engine = engine
        self.lang = lang
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖"""
        pillow_ok, _ = check_pillow_available()
        tesseract_ok, _ = check_tesseract_available()

        if not pillow_ok:
            raise EngineNotFoundError("Pillow 未安装，无法处理图片")
        if not tesseract_ok:
            raise EngineNotFoundError("Tesseract 未安装，无法进行 OCR")

    def preprocess_image(self, img) -> Any:
        """图片预处理：去噪、增强、二值化"""
        from PIL import ImageEnhance, ImageFilter

        # 1. 对比度增强
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)

        # 2. 亮度增强
        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(1.2)

        # 3. 锐化
        img = img.filter(ImageFilter.SHARPEN)

        # 4. 中值滤波去噪
        img = img.filter(ImageFilter.MedianFilter(size=3))

        # 5. 转灰度图
        img = img.convert('L')

        # 6. 二值化（自适应阈值）
        threshold = 128
        img = img.point(lambda x: 0 if x < threshold else 255, '1')

        return img

    def detect_orientation(self, img) -> int:
        """检测图片方向（使用 Tesseract）"""
        try:
            import pytesseract
            osd = pytesseract.image_to_osd(img)
            angle_match = re.search(r'Orientation in degrees: (\d+)', osd)
            if angle_match:
                return int(angle_match.group(1))
        except Exception:
            pass
        return 0

    def correct_orientation(self, img):
        """校正图片方向"""
        angle = self.detect_orientation(img)
        if angle != 0:
            print(f"  🔄 检测到图片旋转 {angle}°，自动校正")
            img = img.rotate(-angle, expand=True)
        return img

    def process_single_image(self, image_path: Path, page_num: int = 1) -> Dict:
        """处理单张图片"""
        from PIL import Image as PILImage
        import pytesseract

        print(f"  🖼️  处理图片: {image_path.name}")

        img = PILImage.open(image_path)

        # 方向校正
        img = self.correct_orientation(img)

        # 图片预处理
        processed_img = self.preprocess_image(img)

        # OCR 识别
        print(f"  🔍  OCR 识别中...")
        text = pytesseract.image_to_string(processed_img, lang=self.lang)

        # 质量评分
        quality = min(100.0, len(text.strip()) / 2)

        return {
            'page_num': page_num,
            'text': text,
            'original_path': str(image_path),
            'quality_score': quality,
            'structure_tags': ['image_ocr']
        }

    def process(self, input_path: Path) -> Dict:
        """处理图片（支持单张或目录批量处理）"""
        pages = []
        full_text = ""
        images = []

        if input_path.is_dir():
            # 目录批量处理
            image_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']
            image_files = sorted([
                f for f in input_path.iterdir()
                if f.is_file() and f.suffix.lower() in image_extensions
            ])

            print(f"  📁 批量处理目录: {input_path.name}, 共 {len(image_files)} 张图片")

            for idx, img_file in enumerate(image_files, 1):
                result = self.process_single_image(img_file, idx)
                pages.append(result)
                full_text += result['text'] + "\n"
                images.append(asdict(ImageContent(
                    image_index=idx - 1,
                    page_num=idx,
                    path=str(img_file),
                    ocr_text=result['text'],
                    quality_score=result['quality_score']
                )))
        else:
            # 单张图片
            result = self.process_single_image(input_path, 1)
            pages.append(result)
            full_text = result['text']
            images.append(asdict(ImageContent(
                image_index=0,
                page_num=1,
                path=str(input_path),
                ocr_text=result['text'],
                quality_score=result['quality_score']
            )))

        return {
            'text': full_text,
            'pages': pages,
            'tables': [],
            'images': images,
            'method': f'{self.engine}_image_ocr',
            'image_count': len(images)
        }


# =============================================================================
# OCR Runner 主类 v2.0
# =============================================================================

class OCRRunner:
    """OCR 运行器主类 v2.0 - 多格式统一入口"""

    def __init__(self, engine: str = 'auto', config_path: str = None, lang: str = 'chi_sim+eng'):
        """
        初始化 OCR 运行器

        Args:
            engine: 引擎选择 ('auto', 'tesseract', 'easyocr', 'paddleocr')
            config_path: 配置文件路径
            lang: 语言设置
        """
        if config_path is None:
            script_dir = Path(__file__).parent
            config_path = script_dir.parent / 'config' / 'engine-config.yaml'

        self.config_path = Path(config_path)
        self.config = self._load_config()
        self.lang = lang

        # 检测依赖
        self.dep_status = check_all_dependencies()

        # 选择引擎
        self.engine = self._select_engine(engine)

        # 初始化子模块
        self.noise_reducer = NoiseReducer()
        self.quality_assessor = QualityAssessor()
        self.table_extractor = TableExtractor()

        # 初始化各格式处理器
        self.pdf_processor = PDFProcessor(engine=self.engine, lang=self.lang)
        self.word_processor = WordProcessor()
        self.image_processor = ImageProcessor(engine=self.engine, lang=self.lang)

    def _load_config(self) -> Dict:
        """加载配置文件"""
        if not self.config_path.exists():
            return {}

        with open(self.config_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)

    def _select_engine(self, requested_engine: str) -> str:
        """自动选择可用引擎"""
        available = self.dep_status['engines_available']

        if requested_engine == 'auto':
            # 按优先级选择
            if available['tesseract']:
                return 'tesseract'
            if available['easyocr']:
                return 'easyocr'
            raise EngineNotFoundError("没有可用的 OCR 引擎，请安装 Tesseract 或 EasyOCR")

        if requested_engine not in available:
            raise EngineNotFoundError(f"未知引擎: {requested_engine}")

        if not available[requested_engine]:
            raise EngineNotFoundError(f"引擎 {requested_engine} 不可用，请检查安装")

        return requested_engine

    def _calculate_quality_scores(self, result: Dict) -> QualityScores:
        """计算质量评分"""
        text = result.get('text', '')
        pages = result.get('pages', [])

        # 使用 QualityAssessor 评估
        assess_result = self.quality_assessor.assess_full(text)

        # 分页质量
        page_scores = {}
        for page in pages:
            page_scores[page['page_num']] = page.get('quality_score', 80.0)

        return QualityScores(
            overall_score=assess_result.get('overall_score', 0.0),
            quality_level=assess_result.get('quality_level_name', 'unknown'),
            page_scores=page_scores,
            character_recognition_rate=assess_result.get('char_valid_rate', 0.0),
            noise_level=assess_result.get('noise_level_name', 'unknown'),
            table_completeness=1.0 if result.get('tables') else 0.0
        )

    def run(self, input_path: str, clean_noise: bool = True,
            auto_correct: bool = True, assess_quality: bool = True) -> Dict:
        """
        执行文档识别 - 统一入口

        Args:
            input_path: 输入文件/目录路径 (PDF/Word/图片)
            clean_noise: 是否清理噪音
            auto_correct: 是否自动纠错
            assess_quality: 是否进行质量评估

        Returns:
            完整识别结果（统一 JSON 格式）
        """
        input_file = Path(input_path)
        if not input_file.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")

        print(f"\n{'='*60}")
        print(f"  📂 开始处理: {input_file.name}")
        print(f"{'='*60}")

        start_time = time.time()

        # 检测文件类型并分发到对应处理器
        result = {}

        if input_file.suffix.lower() == '.pdf':
            result = self.pdf_processor.process(input_file)
            result['input_type'] = 'pdf'

        elif input_file.suffix.lower() in ['.docx', '.doc']:
            result = self.word_processor.process(input_file)
            result['input_type'] = 'word'

        elif input_file.suffix.lower() in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif'] or input_file.is_dir():
            result = self.image_processor.process(input_file)
            result['input_type'] = 'image'

        else:
            raise ValueError(f"不支持的文件格式: {input_file.suffix}")

        # 后处理：噪音清理
        if clean_noise and 'text' in result:
            print(f"  🧹 清理噪音中 ...")
            cleaned = self.noise_reducer.full_clean(result['text'])
            result['text'] = cleaned
            result['clean_stats'] = self.noise_reducer.get_stats()

        # 质量评估
        if assess_quality and 'text' in result:
            print(f"  📊 质量评估中 ...")
            quality_scores = self._calculate_quality_scores(result)
            result['quality_scores'] = asdict(quality_scores)
            print(f"  ✅ 整体质量评分: {quality_scores.overall_score:.1f}/100 ({quality_scores.quality_level})")

        # 表格提取（如果处理器未提取）
        if 'tables' not in result or not result['tables']:
            tables = self.table_extractor.extract_tables(result.get('text', ''))
            if tables:
                result['tables'] = tables
                print(f"  📋 从文本中检测到 {len(tables)} 个表格")

        # 确保统一输出结构存在
        result.setdefault('pages', [])
        result.setdefault('tables', [])
        result.setdefault('images', [])
        result.setdefault('metadata', {})

        # 添加元信息
        result['input_file'] = str(input_file)
        result['engine_used'] = self.engine
        result['total_processing_time'] = time.time() - start_time

        print(f"\n{'='*60}")
        print(f"  ✅ 处理完成!")
        print(f"  📄 总页数: {len(result['pages'])}")
        print(f"  📋 表格数: {len(result['tables'])}")
        print(f"  🖼️  图片数: {len(result['images'])}")
        print(f"  ⏱️  总耗时: {result['total_processing_time']:.2f} 秒")
        print(f"{'='*60}\n")

        return result


# =============================================================================
# 命令行接口
# =============================================================================

def print_dependency_check():
    """打印依赖检查结果"""
    status = check_all_dependencies()

    print(f"\n{'='*60}")
    print(f"  🧪 OCR 引擎 v2.0 依赖检查")
    print(f"{'='*60}")

    dep_names = {
        'tesseract': 'Tesseract OCR 引擎',
        'poppler': 'Poppler (PDF 渲染)',
        'pymupdf': 'PyMuPDF (PDF 处理)',
        'pdfplumber': 'pdfplumber (PDF 文本提取)',
        'python_docx': 'python-docx (Word 文档)',
        'pillow': 'Pillow (图像处理)',
        'easyocr': 'EasyOCR (备用引擎)',
    }

    for name, (available, msg) in status['dependencies'].items():
        icon = '✅' if available else '❌'
        print(f"  {icon} {dep_names.get(name, name):25s} : {msg}")

    print()
    print(f"  可用 OCR 引擎:")
    for engine, available in status['engines_available'].items():
        if available:
            print(f"    ✅ {engine}")

    if status['all_good']:
        print(f"\n  ✅ 核心依赖全部满足!")
    else:
        print(f"\n  ⚠️  部分核心依赖缺失，请安装后重试")

    print(f"{'='*60}\n")


def main():
    import argparse
    import json

    parser = argparse.ArgumentParser(description='OCR 文档识别引擎 v2.0 - 支持 PDF/Word/图片')
    parser.add_argument('--input', '-i', help='输入文件/目录路径 (PDF/Word/图片)')
    parser.add_argument('--output', '-o', help='输出 JSON 文件路径')
    parser.add_argument('--engine', '-e', choices=['auto', 'tesseract', 'easyocr'],
                        default='auto', help='选择 OCR 引擎')
    parser.add_argument('--lang', '-l', default='chi_sim+eng', help='语言设置 (默认: chi_sim+eng)')
    parser.add_argument('--no-clean', action='store_true', help='跳过噪音清理')
    parser.add_argument('--no-assess', action='store_true', help='跳过质量评估')
    parser.add_argument('--check-deps', action='store_true', help='检查依赖项')
    parser.add_argument('--text-only', action='store_true', help='仅输出识别文本')
    parser.add_argument('--test', action='store_true', help='运行测试')

    args = parser.parse_args()

    if args.check_deps:
        print_dependency_check()
        return

    if args.test:
        print("🧪 运行 OCR 引擎 v2.0 测试...")
        print_dependency_check()
        print("✅ 测试完成!")
        return

    if not args.input:
        print("❌ 请指定 --input 参数，或使用 --help 查看帮助")
        return

    try:
        runner = OCRRunner(engine=args.engine, lang=args.lang)
        result = runner.run(
            args.input,
            clean_noise=not args.no_clean,
            assess_quality=not args.no_assess
        )

        if args.text_only:
            print(result.get('text', ''))
        elif args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            print(f"✅ 结果已保存到: {args.output}")
        else:
            # 摘要输出
            print(f"\n📊 处理摘要:")
            print(f"  文件类型: {result.get('input_type')}")
            print(f"  使用引擎: {result.get('engine_used')}")
            print(f"  总页数: {len(result.get('pages', []))}")
            print(f"  表格数: {len(result.get('tables', []))}")
            print(f"  图片数: {len(result.get('images', []))}")
            print(f"  总耗时: {result.get('total_processing_time', 0):.2f} 秒")

            if 'quality_scores' in result:
                q = result['quality_scores']
                print(f"  整体质量: {q['overall_score']:.1f}/100 ({q['quality_level']})")
                print(f"  字符识别率: {q['character_recognition_rate']:.1%}")

            if 'clean_stats' in result:
                stats = result['clean_stats']
                print(f"  噪音清理: 公章移除={stats.get('seal_removed', 0)}, "
                      f"表格垃圾={stats.get('table_garbage_removed', 0)}")

            print(f"\n📝 文本预览 (前 500 字符):")
            print(f"{'-'*60}")
            text = result.get('text', '')
            print(text[:500] + ('...' if len(text) > 500 else ''))
            print(f"{'-'*60}")

    except Exception as e:
        print(f"❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
