"""
Document Parser module for Iris RAG engine.
Provides multi-format document parsing and intelligent text chunking.
"""
import os
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

# Try to import optional dependencies
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import nltk
    from nltk.tokenize import sent_tokenize
    NLTK_AVAILABLE = True
    try:
        sent_tokenize("Test sentence.")
    except LookupError:
        try:
            import ssl
            try:
                _create_unverified_https_context = ssl._create_unverified_context
            except AttributeError:
                pass
            else:
                ssl._create_default_https_context = _create_unverified_https_context
            nltk.download('punkt_tab', quiet=True)
        except:
            try:
                nltk.download('punkt', quiet=True)
            except:
                NLTK_AVAILABLE = False
except ImportError:
    NLTK_AVAILABLE = False


@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document."""
    text: str
    index: int
    start_pos: int
    end_pos: int
    metadata: Dict[str, Any]


class DocumentParser:
    """Parser for multiple document formats with intelligent chunking."""
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50,
                 use_semantic_boundaries: bool = True):
        """
        Initialize document parser.
        
        Args:
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between consecutive chunks
            use_semantic_boundaries: Try to split at sentence boundaries
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_semantic_boundaries = use_semantic_boundaries
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a file and extract its content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dict with content, metadata, and file info
        """
        path = Path(file_path)
        file_name = path.name
        file_type = path.suffix.lower()
        file_size = path.stat().st_size
        
        # Calculate file hash for change detection
        file_hash = self._calculate_file_hash(file_path)
        
        # Parse based on file type
        if file_type == '.txt':
            content = self._parse_txt(file_path)
        elif file_type == '.md':
            content = self._parse_markdown(file_path)
        elif file_type == '.pdf':
            content = self._parse_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            content = self._parse_docx(file_path)
        elif file_type in ['.xlsx', '.xls', '.csv']:
            content = self._parse_excel(file_path)
        else:
            # Try to parse as text
            try:
                content = self._parse_txt(file_path)
            except:
                raise ValueError(f"Unsupported file type: {file_type}")
        
        return {
            "content": content,
            "file_name": file_name,
            "file_type": file_type,
            "file_size": file_size,
            "file_hash": file_hash,
            "file_path": str(path.absolute())
        }
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate MD5 hash of a file."""
        md5_hash = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                md5_hash.update(chunk)
        return md5_hash.hexdigest()
    
    def _parse_txt(self, file_path: str) -> str:
        """Parse plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _parse_markdown(self, file_path: str) -> str:
        """Parse Markdown file."""
        content = self._parse_txt(file_path)
        # Convert markdown to plain text for better chunking
        if MARKDOWN_AVAILABLE:
            html = markdown.markdown(content)
            # Simple HTML to text conversion
            text = re.sub(r'<[^>]+>', '', html)
            return text
        return content
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF parsing")
        
        text_parts = []
        reader = PdfReader(file_path)
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        
        return "\n\n".join(text_parts)
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing")
        
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _parse_excel(self, file_path: str) -> str:
        """Parse Excel spreadsheet."""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl is required for Excel parsing")
        
        text_parts = []
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(v) if v is not None else "" for v in row]
                if any(row_values):
                    text_parts.append(" | ".join(row_values))
        
        wb.close()
        return "\n".join(text_parts)
    
    def chunk_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks.
        
        Args:
            text: Input text to chunk
            metadata: Additional metadata to attach to each chunk
            
        Returns:
            List of chunk dictionaries
        """
        if not text.strip():
            return []
        
        base_metadata = metadata or {}
        chunks = []
        
        if self.use_semantic_boundaries and NLTK_AVAILABLE:
            chunks = self._chunk_semantic(text, base_metadata)
        else:
            chunks = self._chunk_fixed(text, base_metadata)
        
        return chunks
    
    def _chunk_fixed(self, text: str, base_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Simple fixed-size chunking with overlap."""
        chunks = []
        chunk_index = 0
        pos = 0
        text_len = len(text)
        
        while pos < text_len:
            end_pos = min(pos + self.chunk_size, text_len)
            
            # Try to end at a newline or period if possible
            if end_pos < text_len:
                for break_char in ['\n', '. ', '! ', '? ', '; ']:
                    break_pos = text.rfind(break_char, pos, end_pos)
                    if break_pos > pos + self.chunk_size // 2:
                        end_pos = break_pos + len(break_char)
                        break
            
            chunk_text = text[pos:end_pos].strip()
            if chunk_text:
                chunks.append({
                    "text": chunk_text,
                    "index": chunk_index,
                    "start_pos": pos,
                    "end_pos": end_pos,
                    "metadata": {
                        **base_metadata,
                        "chunk_method": "fixed"
                    }
                })
                chunk_index += 1
            
            # Move forward with overlap
            pos = end_pos - self.chunk_overlap
            if pos <= 0 or pos >= text_len:
                break
        
        return chunks
    
    def _chunk_semantic(self, text: str, base_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Semantic chunking based on sentence boundaries."""
        sentences = sent_tokenize(text)
        chunks = []
        current_chunk = []
        current_length = 0
        chunk_index = 0
        pos = 0
        
        for sentence in sentences:
            sent_len = len(sentence)
            
            if current_length + sent_len > self.chunk_size and current_chunk:
                # Save current chunk
                chunk_text = " ".join(current_chunk)
                chunks.append({
                    "text": chunk_text,
                    "index": chunk_index,
                    "start_pos": pos,
                    "end_pos": pos + len(chunk_text),
                    "metadata": {
                        **base_metadata,
                        "chunk_method": "semantic"
                    }
                })
                chunk_index += 1
                
                # Keep overlap sentences
                overlap_chars = 0
                overlap_sentences = []
                for s in reversed(current_chunk):
                    if overlap_chars + len(s) <= self.chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_chars += len(s)
                    else:
                        break
                
                current_chunk = overlap_sentences
                current_length = sum(len(s) for s in current_chunk)
                pos += len(chunk_text) - overlap_chars
            
            current_chunk.append(sentence)
            current_length += sent_len
        
        # Add the last chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "text": chunk_text,
                "index": chunk_index,
                "start_pos": pos,
                "end_pos": pos + len(chunk_text),
                "metadata": {
                    **base_metadata,
                    "chunk_method": "semantic"
                }
            })
        
        return chunks
    
    def parse_directory(self, dir_path: str, recursive: bool = True,
                        extensions: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        """
        Parse all documents in a directory.
        
        Args:
            dir_path: Path to directory
            recursive: Whether to scan subdirectories
            extensions: List of allowed extensions (e.g., ['.txt', '.pdf'])
            
        Returns:
            List of parsed documents
        """
        default_extensions = ['.txt', '.md', '.pdf', '.docx', '.xlsx', '.xls', '.csv']
        allowed_extensions = set(extensions or default_extensions)
        
        documents = []
        path = Path(dir_path)
        
        glob_pattern = '**/*' if recursive else '*'
        
        for file_path in path.glob(glob_pattern):
            if file_path.is_file() and file_path.suffix.lower() in allowed_extensions:
                try:
                    doc = self.parse_file(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    print(f"Failed to parse {file_path}: {e}")
        
        return documents
