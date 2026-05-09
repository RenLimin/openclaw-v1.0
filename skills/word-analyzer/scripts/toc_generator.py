#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 目录（TOC）管理模块
========================
自动检测标题层级、生成目录、校验目录与页码一致性

功能:
  - 自动检测标题层级
  - 自动生成目录（带页码）
  - 目录与实际页码一致性校验
  - 目录更新

作者: Word-Analyzer Team
版本: v1.0
日期: 2026-04-24
"""

import re
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("❌ 错误: 请安装 python-docx: pip install python-docx")
    exit(1)

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False

from word_parser import WordParser


class TOCGenerator:
    """Word 目录生成与校验器"""

    def __init__(self, file_path: str, config_path: str = None):
        """
        初始化目录管理器

        Args:
            file_path: Word 文档路径 (.docx)
            config_path: 配置文件路径（可选）
        """
        self.file_path = Path(file_path)
        self.config = self._load_config(config_path)
        self.parser = WordParser(file_path, config_path)
        self.doc = None

    def _load_config(self, config_path: str = None) -> Dict:
        """加载配置文件"""
        default_config = {
            'toc_rules': {
                'toc_location_keywords': ["目录", "目次", "Contents"],
                'max_search_lines': 50,
                'page_separators': [r'\.{2,}', r'\s{2,}', r'\t+'],
                'max_level': 4
            },
            'toc_style': {
                'title': '目录',
                'title_font': '黑体',
                'title_size': 16,
                'title_bold': True,
                'entry_font': '宋体',
                'entry_size': 12,
                'leader': '.',
                'indent_per_level': 1.0  # cm
            }
        }

        if config_path and HAS_YAML:
            cfg_path = Path(config_path)
            if cfg_path.exists():
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    user_config = yaml.safe_load(f)
                    if 'toc_rules' in user_config:
                        default_config['toc_rules'].update(user_config['toc_rules'])
                    if 'toc_style' in user_config:
                        default_config['toc_style'].update(user_config['toc_style'])

        return default_config

    def _estimate_page_number(self, paragraph_index: int, total_paragraphs: int,
                              total_pages: int = 0) -> int:
        """
        估算段落所在页码（近似算法）
        注意: python-docx 无法直接获取精确页码，此为估算值
        """
        if total_pages <= 0:
            # 如果无法获取总页数，按每页约50段估算
            estimated_pages = max(1, total_paragraphs // 50 + 1)
        else:
            estimated_pages = total_pages

        # 线性估算页码
        page = int((paragraph_index / max(1, total_paragraphs)) * estimated_pages) + 1
        return min(page, estimated_pages)

    def detect_existing_toc(self) -> Optional[Dict[str, Any]]:
        """
        检测文档中是否已存在目录

        Returns:
            目录信息字典，如果未找到则返回 None
        """
        paragraphs = self.parser.extract_all_paragraphs()
        toc_rules = self.config['toc_rules']
        keywords = toc_rules['toc_location_keywords']
        max_search = min(toc_rules['max_search_lines'], len(paragraphs))

        # 查找目录起始位置
        toc_start = None
        for i in range(max_search):
            text = paragraphs[i]['text']
            for kw in keywords:
                if kw in text and len(text) < 20:  # 短文本包含关键词通常是目录标题
                    toc_start = i
                    break
            if toc_start:
                break

        if toc_start is None:
            return None

        # 识别目录条目（包含页码特征的行）
        toc_entries = []
        page_separators = self.config['toc_rules']['page_separators']
        separator_pattern = re.compile('|'.join(page_separators))

        for i in range(toc_start + 1, len(paragraphs)):
            para = paragraphs[i]
            text = para['text'].strip()

            if not text:
                continue

            # 检测目录条目特征：包含分隔符且结尾是数字
            match = separator_pattern.search(text)
            if match:
                before = text[:match.start()].strip()
                after = text[match.end():].strip()

                # 检查后面是否是页码（数字）
                if after.isdigit():
                    toc_entries.append({
                        'index': i,
                        'text': text,
                        'heading_text': before,
                        'page_number': int(after),
                        'separator': text[match.start():match.end()]
                    })
            # 如果连续3行都不是目录格式，认为目录结束
            elif len(toc_entries) > 3 and i - toc_start > 10:
                break

        if not toc_entries:
            return None

        return {
            'title_index': toc_start,
            'title_text': paragraphs[toc_start]['text'],
            'entries': toc_entries,
            'entry_count': len(toc_entries)
        }

    def validate_toc(self) -> Dict[str, Any]:
        """
        校验目录的正确性：
        1. 目录条目是否与实际标题一致
        2. 目录页码是否与估算页码一致

        Returns:
            校验报告
        """
        toc_info = self.detect_existing_toc()
        headings = self.parser.extract_headings()

        report = {
            'has_toc': toc_info is not None,
            'success': True,
            'missing_in_toc': [],
            'extra_in_toc': [],
            'page_mismatch': [],
            'text_mismatch': [],
            'stats': {
                'total_headings': len(headings),
                'toc_entries': toc_info['entry_count'] if toc_info else 0
            }
        }

        if not toc_info:
            report['success'] = False
            report['message'] = '文档中未检测到目录'
            return report

        # 提取目录中的标题文本
        toc_headings = {entry['heading_text']: entry for entry in toc_info['entries']}

        # 检查实际标题是否都在目录中
        for heading in headings:
            text = heading['text'].strip()
            # 模糊匹配
            found = False
            for toc_text in toc_headings.keys():
                if text in toc_text or toc_text in text:
                    found = True
                    # 检查页码
                    estimated_page = self._estimate_page_number(
                        heading['index'],
                        len(self.parser.doc.paragraphs) if self.parser.doc else 100
                    )
                    toc_page = toc_headings[toc_text]['page_number']
                    if abs(estimated_page - toc_page) > 2:  # 允许2页误差
                        report['page_mismatch'].append({
                            'heading': text,
                            'toc_page': toc_page,
                            'estimated_page': estimated_page
                        })
                    break
            if not found:
                report['missing_in_toc'].append(text)

        # 检查目录中的条目是否都存在
        actual_heading_texts = {h['text'].strip() for h in headings}
        for toc_heading in toc_headings.keys():
            found = any(toc_heading in ah or ah in toc_heading for ah in actual_heading_texts)
            if not found:
                report['extra_in_toc'].append(toc_heading)

        # 计算匹配度
        matched = len(headings) - len(report['missing_in_toc'])
        report['match_rate'] = matched / max(1, len(headings))
        report['success'] = (len(report['missing_in_toc']) == 0 and
                           len(report['extra_in_toc']) == 0 and
                           len(report['page_mismatch']) == 0)

        return report

    def generate_toc_entries(self) -> List[Dict[str, Any]]:
        """根据文档中的标题生成目录条目"""
        headings = self.parser.extract_headings()
        max_level = self.config['toc_rules']['max_level']

        # 过滤掉层级过深的标题
        filtered_headings = [h for h in headings if h['level'] <= max_level]

        # 生成目录条目
        entries = []
        total_paragraphs = len(self.parser.doc.paragraphs) if self.parser.doc else 100

        for heading in filtered_headings:
            page_num = self._estimate_page_number(heading['index'], total_paragraphs)
            entries.append({
                'level': heading['level'],
                'text': heading['text'],
                'page_number': page_num,
                'paragraph_index': heading['index']
            })

        return entries

    def _set_font(self, run, font_name: str, size: float, bold: bool = False):
        """设置字体（兼容中文）"""
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        # 设置中文字体
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def generate_toc(self, output_path: str = None, insert_at_beginning: bool = True) -> str:
        """
        生成目录并保存到文档

        Args:
            output_path: 输出文件路径，默认为原文件名加 _with_toc
            insert_at_beginning: 是否在文档开头插入目录

        Returns:
            输出文件路径
        """
        if output_path is None:
            output_path = self.file_path.with_stem(self.file_path.stem + '_with_toc')

        # 加载文档
        self.doc = Document(str(self.file_path))

        # 获取样式配置
        style = self.config['toc_style']
        entries = self.generate_toc_entries()

        if insert_at_beginning:
            # 在文档开头插入
            insert_pos = 0

            # 添加目录标题
            title_para = self.doc.paragraphs[insert_pos].insert_paragraph_before()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(style['title'])
            self._set_font(title_run, style['title_font'], style['title_size'], style['title_bold'])

            insert_pos += 1

            # 添加空行
            empty_para = self.doc.paragraphs[insert_pos].insert_paragraph_before()
            insert_pos += 1

            # 添加目录条目
            for entry in entries:
                para = self.doc.paragraphs[insert_pos].insert_paragraph_before()

                # 缩进
                indent = Pt(style['indent_per_level'] * 28.35 * (entry['level'] - 1))
                para.paragraph_format.left_indent = indent

                # 标题文本
                text_run = para.add_run(entry['text'])
                self._set_font(text_run, style['entry_font'], style['entry_size'])

                # 前导符和页码
                leader = style['leader'] * 2
                page_text = f"{leader}{entry['page_number']}"
                page_run = para.add_run(page_text)
                self._set_font(page_run, style['entry_font'], style['entry_size'])

                insert_pos += 1

            # 添加分页符
            page_break_para = self.doc.paragraphs[insert_pos].insert_paragraph_before()
            page_break_para.add_run().add_break(break_type=7)  # 7 = WD_BREAK.PAGE

        else:
            # 在文档末尾添加目录
            self.doc.add_paragraph()
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(style['title'])
            self._set_font(title_run, style['title_font'], style['title_size'], style['title_bold'])

            self.doc.add_paragraph()

            for entry in entries:
                para = self.doc.add_paragraph()
                indent = Pt(style['indent_per_level'] * 28.35 * (entry['level'] - 1))
                para.paragraph_format.left_indent = indent

                text_run = para.add_run(entry['text'])
                self._set_font(text_run, style['entry_font'], style['entry_size'])

                leader = style['leader'] * 2
                page_text = f"{leader}{entry['page_number']}"
                page_run = para.add_run(page_text)
                self._set_font(page_run, style['entry_font'], style['entry_size'])

        # 保存文档
        self.doc.save(str(output_path))
        print(f"✅ 目录已生成，包含 {len(entries)} 个条目")
        print(f"💾 文档已保存: {output_path}")

        return str(output_path)

    def update_toc(self, output_path: str = None) -> str:
        """
        更新现有目录的页码

        Args:
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        # 检测现有目录
        toc_info = self.detect_existing_toc()

        if not toc_info:
            print("⚠️  文档中未检测到现有目录，将重新生成目录")
            return self.generate_toc(output_path)

        # 生成新的目录条目
        new_entries = self.generate_toc_entries()

        # 加载文档进行更新
        self.doc = Document(str(self.file_path))

        # 更新目录条目中的页码
        # 注意：这里简单实现，实际 Word 的 TOC 域需要特殊处理
        # 我们采用替换策略：删除旧目录，插入新目录
        # 先删除旧目录条目（倒序删除）
        for entry in reversed(toc_info['entries']):
            if entry['index'] < len(self.doc.paragraphs):
                p = self.doc.paragraphs[entry['index']]
                p.clear()

        # 在目录标题后插入新目录
        insert_pos = toc_info['title_index'] + 1

        style = self.config['toc_style']
        for entry in new_entries:
            if insert_pos < len(self.doc.paragraphs):
                para = self.doc.paragraphs[insert_pos].insert_paragraph_before()
            else:
                para = self.doc.add_paragraph()

            indent = Pt(style['indent_per_level'] * 28.35 * (entry['level'] - 1))
            para.paragraph_format.left_indent = indent

            text_run = para.add_run(entry['text'])
            self._set_font(text_run, style['entry_font'], style['entry_size'])

            leader = style['leader'] * 2
            page_text = f"{leader}{entry['page_number']}"
            page_run = para.add_run(page_text)
            self._set_font(page_run, style['entry_font'], style['entry_size'])

            insert_pos += 1

        if output_path is None:
            output_path = self.file_path.with_stem(self.file_path.stem + '_updated_toc')

        self.doc.save(str(output_path))
        print(f"✅ 目录已更新，共 {len(new_entries)} 个条目")
        print(f"💾 文档已保存: {output_path}")

        return str(output_path)


def main():
    parser = argparse.ArgumentParser(description='Word 目录管理工具')
    parser.add_argument('--input', '-i', required=True, help='输入 Word 文档路径')
    parser.add_argument('--config', '-c', help='配置文件路径')

    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('--validate', action='store_true', help='校验现有目录')
    group.add_argument('--generate', action='store_true', help='生成目录')
    group.add_argument('--update', action='store_true', help='更新目录页码')

    parser.add_argument('--output', '-o', help='输出文件路径')

    args = parser.parse_args()

    toc_gen = TOCGenerator(args.input, args.config)

    if args.validate:
        result = toc_gen.validate_toc()
        print(json.dumps(result, ensure_ascii=False, indent=2))

        if result['has_toc']:
            print(f"\n✅ 检测到目录，共 {result['stats']['toc_entries']} 个条目")
            print(f"🎯 标题匹配度: {result['match_rate'] * 100:.1f}%")

            if result['missing_in_toc']:
                print(f"\n⚠️  目录缺失的标题 ({len(result['missing_in_toc'])}):")
                for h in result['missing_in_toc'][:5]:
                    print(f"  - {h}")

            if result['page_mismatch']:
                print(f"\n⚠️  页码可能不一致 ({len(result['page_mismatch'])}):")
                for m in result['page_mismatch'][:5]:
                    print(f"  - {m['heading']}: 目录写{m['toc_page']}, 估算{m['estimated_page']}")
        else:
            print("\n❌ 未检测到目录")

    elif args.generate:
        toc_gen.generate_toc(args.output)

    elif args.update:
        toc_gen.update_toc(args.output)


if __name__ == "__main__":
    main()
