#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
履约义务拆分引擎 v1.0
=====================
完整流程:
  输入文件 → OCR识别 → 条款拆分 → 义务提取 → 标准比对 → 风险评估 → Excel导出

作者: Ella 🦊
版本: v1.0
日期: 2026-04-24
"""

import os
import re
import sys
import argparse
import yaml
import json
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
from difflib import SequenceMatcher

# =============================================================================
# 颜色输出
# =============================================================================

class Colors:
    RED = '\033[91m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    PURPLE = '\033[95m'
    CYAN = '\033[96m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'

# =============================================================================
# 异常类定义
# =============================================================================

class ObligationSplitError(Exception):
    """履约义务拆分基异常"""
    pass

class FileNotFoundError(ObligationSplitError):
    """文件不存在"""
    pass

class UnsupportedFormatError(ObligationSplitError):
    """不支持的文件格式"""
    pass

class OCRError(ObligationSplitError):
    """OCR识别失败"""
    pass

class ClauseSplitError(ObligationSplitError):
    """条款拆分失败"""
    pass

class ProductMatchError(ObligationSplitError):
    """产品匹配失败"""
    pass

# =============================================================================
# 主引擎类
# =============================================================================

class ObligationSplitter:
    """履约义务拆分主引擎"""

    def __init__(self, config_path: str = None):
        """
        初始化拆分引擎

        Args:
            config_path: 配置文件目录路径
        """
        # 设置默认配置路径
        if config_path is None:
            script_dir = Path(__file__).parent
            config_path = script_dir.parent / 'config'

        self.config_dir = Path(config_path)
        self.config = self._load_config()

        # 添加依赖技能路径
        skill_base = Path(__file__).parent.parent.parent
        for skill_name, rel_path in self.config.get('skill_paths', {}).items():
            full_path = skill_base / rel_path
            if str(full_path) not in sys.path:
                sys.path.insert(0, str(full_path))

        self.ocr_engine = None
        self.clause_splitter = None
        self.excel_io = None

        # 处理结果缓存
        self._raw_text = None
        self._clauses = None
        self._obligations = None
        self._matched_product = None
        self._differences = None
        self._risk_summary = None

    def _load_config(self) -> Dict:
        """加载配置文件"""
        config_file = self.config_dir / 'config.yaml'
        if not config_file.exists():
            print(f"{Colors.YELLOW}⚠️  警告: 配置文件不存在，使用默认配置{Colors.ENDC}")
            return {}

        with open(config_file, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)

    # =========================================================================
    # 依赖检查
    # =========================================================================

    def check_dependencies(self) -> Dict[str, Any]:
        """
        检查所有依赖是否满足

        Returns:
            依赖检查结果字典
        """
        print(f"\n{Colors.CYAN}📦 依赖检查中...{Colors.ENDC}")

        results = {
            'python_packages': {},
            'skills': {},
            'system': {},
            'all_passed': True
        }

        # 检查 Python 包
        packages = [
            ('docx', 'python-docx'),
            ('pypdf', 'pypdf'),
            ('pdfplumber', 'pdfplumber'),
            ('yaml', 'pyyaml'),
            ('openpyxl', 'openpyxl'),
            ('PIL', 'pillow'),
            ('pandas', 'pandas'),
        ]

        for import_name, pkg_name in packages:
            try:
                __import__(import_name)
                results['python_packages'][pkg_name] = True
                print(f"  ✅ {pkg_name}")
            except ImportError:
                results['python_packages'][pkg_name] = False
                results['all_passed'] = False
                print(f"  {Colors.RED}❌ {pkg_name} (未安装){Colors.ENDC}")

        # 检查技能依赖
        skill_base = Path(__file__).parent.parent.parent
        skills = [
            ('ocr-engine', 'ocr-engine/scripts/ocr_runner.py'),
            ('contract-clause-split', 'contract-clause-split/scripts/split_contract.py'),
            ('excel-engine', 'excel-engine/scripts/excel_io.py'),
        ]

        for skill_name, script_path in skills:
            full_path = skill_base / script_path
            exists = full_path.exists()
            results['skills'][skill_name] = exists
            if exists:
                print(f"  ✅ {skill_name}")
            else:
                results['all_passed'] = False
                print(f"  {Colors.RED}❌ {skill_name} (技能不存在){Colors.ENDC}")

        # 检查系统依赖（Tesseract）
        try:
            import subprocess
            result = subprocess.run(
                ['tesseract', '--version'],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode == 0:
                version = result.stdout.split('\n')[0]
                results['system']['tesseract'] = True
                results['system']['tesseract_version'] = version
                print(f"  ✅ Tesseract OCR ({version})")
            else:
                results['system']['tesseract'] = False
                results['all_passed'] = False
                print(f"  {Colors.RED}❌ Tesseract OCR (执行失败){Colors.ENDC}")
        except Exception:
            results['system']['tesseract'] = False
            results['all_passed'] = False
            print(f"  {Colors.RED}❌ Tesseract OCR (未安装){Colors.ENDC}")

        if results['all_passed']:
            print(f"\n{Colors.GREEN}✅ 所有依赖检查通过！{Colors.ENDC}")
        else:
            print(f"\n{Colors.RED}❌ 部分依赖未通过，请安装后重试{Colors.ENDC}")

        return results

    # =========================================================================
    # 阶段1: 文件解析与文本提取
    # =========================================================================

    def _detect_file_type(self, file_path: Path) -> str:
        """检测文件类型"""
        suffix = file_path.suffix.lower()

        if suffix in ['.pdf']:
            return 'pdf'
        elif suffix in ['.docx', '.doc']:
            return 'word'
        elif suffix in ['.jpg', '.jpeg', '.png', '.bmp', '.gif']:
            return 'image'
        else:
            raise UnsupportedFormatError(f"不支持的文件格式: {suffix}")

    def _extract_from_word(self, file_path: Path) -> str:
        """从 Word 文档提取文本"""
        try:
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) > 3:
                    paragraphs.append(text)

            # 提取表格
            for i, table in enumerate(doc.tables, 1):
                table_lines = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_lines.append(' | '.join(row_cells))
                if table_lines:
                    paragraphs.append(f"[表格{i}] {chr(10).join(table_lines)}")

            return '\n'.join(paragraphs)
        except ImportError:
            raise ObligationSplitError("请安装 python-docx: pip install python-docx")
        except Exception as e:
            raise ObligationSplitError(f"Word 解析失败: {e}")

    def _extract_from_pdf_text(self, file_path: Path) -> str:
        """从文本版 PDF 提取内容"""
        try:
            import pdfplumber
            paragraphs = []

            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        for para in text.split('\n'):
                            para = para.strip()
                            if para and len(para) > 3:
                                paragraphs.append(para)

            return '\n'.join(paragraphs)
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                paragraphs = []
                reader = PdfReader(str(file_path))
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        for para in text.split('\n'):
                            para = para.strip()
                            if para and len(para) > 3:
                                paragraphs.append(para)
                return '\n'.join(paragraphs)
            except ImportError:
                raise ObligationSplitError("请安装 pdfplumber 或 PyPDF2")
        except Exception as e:
            raise ObligationSplitError(f"PDF 解析失败: {e}")

    def _is_scanned_pdf(self, file_path: Path) -> bool:
        """检测是否为扫描版 PDF"""
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                # 检查前3页是否有文本
                for i, page in enumerate(pdf.pages[:3]):
                    text = page.extract_text()
                    if text and len(text.strip()) > 100:
                        return False
                return True
        except:
            return True

    def _extract_with_ocr(self, file_path: Path) -> str:
        """使用 OCR 引擎提取文本"""
        print(f"  {Colors.CYAN}🔍 调用 OCR 引擎识别...{Colors.ENDC}")

        try:
            from ocr_runner import OCRRunner
            ocr_config = self.config.get('ocr', {})
            ocr = OCRRunner(
                engine=ocr_config.get('default_engine', 'tesseract'),
                config_path=str(Path(__file__).parent.parent.parent / 'ocr-engine' / 'config')
            )

            result = ocr.run(
                str(file_path),
                clean_noise=ocr_config.get('clean_noise', True),
                auto_correct=ocr_config.get('auto_correct', True)
            )

            return result.get('text', '')
        except ImportError as e:
            raise OCRError(f"OCR 引擎导入失败: {e}")
        except Exception as e:
            raise OCRError(f"OCR 识别失败: {e}")

    def extract_text(self, file_path: str) -> str:
        """
        从文件中提取文本（自动识别格式和类型）

        Args:
            file_path: 输入文件路径

        Returns:
            提取的完整文本
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        print(f"\n{Colors.CYAN}📂 解析文件: {path.name}{Colors.ENDC}")

        file_type = self._detect_file_type(path)

        if file_type == 'word':
            text = self._extract_from_word(path)
            print(f"  ✅ Word 文档解析完成，共 {len(text)} 字符")
        elif file_type == 'image':
            text = self._extract_with_ocr(path)
            print(f"  ✅ 图片 OCR 识别完成，共 {len(text)} 字符")
        else:  # pdf
            if self._is_scanned_pdf(path):
                print(f"  📄 检测到扫描版 PDF，启用 OCR...")
                text = self._extract_with_ocr(path)
                print(f"  ✅ PDF OCR 识别完成，共 {len(text)} 字符")
            else:
                text = self._extract_from_pdf_text(path)
                print(f"  ✅ PDF 文本解析完成，共 {len(text)} 字符")

        self._raw_text = text
        return text

    # =========================================================================
    # 阶段2: 条款拆分
    # =========================================================================

    def split_clauses(self, text: str) -> List[Dict]:
        """
        拆分合同条款（调用 contract-clause-split 逻辑）

        Args:
            text: 合同文本

        Returns:
            条款列表
        """
        print(f"\n{Colors.CYAN}✂️  条款拆分中...{Colors.ENDC}")

        # 条款序号匹配正则
        clause_start_patterns = [
            r'^第[一二三四五六七八九十\d]+[条章节款]',
            r'^\d+\.\s',
            r'^\d+\.\d+\s',
            r'^\d+\.\d+\.\d+\s',
            r'^[（(][一二三四五六七八九十\d]+[）)]\s*',
            r'^[一二三四五六七八九十]{1,2}[、．.]\s*',
            r'^\[表格\d+\]',
        ]

        combined_pattern = re.compile('|'.join(clause_start_patterns))

        lines = text.split('\n')
        clauses = []
        current_clause_lines = []

        for line in lines:
            line = line.strip()
            if not line or len(line) < 3:
                continue

            is_new_clause = bool(combined_pattern.match(line))

            if is_new_clause:
                if current_clause_lines:
                    clause_content = '\n'.join(current_clause_lines)
                    clauses.append({
                        'id': len(clauses) + 1,
                        'content': clause_content,
                        'length': len(clause_content)
                    })
                current_clause_lines = [line]
            else:
                current_clause_lines.append(line)

        # 保存最后一个条款
        if current_clause_lines:
            clause_content = '\n'.join(current_clause_lines)
            clauses.append({
                'id': len(clauses) + 1,
                'content': clause_content,
                'length': len(clause_content)
            })

        # 条款分类
        clauses = self._classify_clauses(clauses)

        print(f"  ✅ 拆分完成，共 {len(clauses)} 个条款")
        self._clauses = clauses
        return clauses

    def _classify_clauses(self, clauses: List[Dict]) -> List[Dict]:
        """简单的条款分类"""
        categories = {
            '基本信息': ['合同名称', '合同编号', '签订日期', '签订地点', '甲方：', '乙方：'],
            '标的条款': ['产品', '设备', '服务', '清单', '规格', '型号'],
            '价格条款': ['金额', '价款', '费用', '支付', '付款', '人民币', '元整'],
            '履行条款': ['交付', '工期', '期限', '时间', '地点', '方式'],
            '权利义务': ['义务', '责任', '权利', '应', '应当', '必须'],
            '违约责任': ['违约', '赔偿', '违约金', '赔偿金', '责任'],
            '保密条款': ['保密', '机密', '秘密', '泄露'],
            '争议解决': ['争议', '诉讼', '仲裁', '管辖', '法院'],
        }

        for clause in clauses:
            content = clause['content']
            clause['category'] = '其他条款'

            for cat_name, keywords in categories.items():
                for kw in keywords:
                    if kw in content:
                        clause['category'] = cat_name
                        break
                if clause['category'] != '其他条款':
                    break

        return clauses

    # =========================================================================
    # 阶段3: 义务提取
    # =========================================================================

    def extract_obligations(self, clauses: List[Dict]) -> List[Dict]:
        """
        从条款中提取履约义务

        Args:
            clauses: 条款列表

        Returns:
            义务列表
        """
        print(f"\n{Colors.CYAN}🎯 义务提取中...{Colors.ENDC}")

        obligation_types = self.config.get('obligation_types', [])
        extraction_rules = self.config.get('extraction_rules', {})
        obligations = []

        for clause in clauses:
            content = clause['content']

            # 识别义务类型
            for obl_type in obligation_types:
                keywords = obl_type.get('keywords', [])
                core_features = obl_type.get('core_features', [])

                score = 0
                for feature in core_features:
                    if feature in content:
                        score += 3
                for kw in keywords:
                    if kw in content:
                        score += 1

                if score >= 2:
                    # 提取字段
                    obligation = {
                        'id': len(obligations) + 1,
                        'type_code': obl_type['code'],
                        'type_name': obl_type['name'],
                        'content': self._clean_obligation_content(content),
                        'source_clause_id': clause['id'],
                        'source_clause_category': clause['category'],
                    }

                    # 提取履行时间
                    obligation['performance_time'] = self._extract_performance_time(content)

                    # 提取责任方
                    obligation['responsible_party'] = self._extract_responsible_party(content)

                    # 提取金额/比例
                    obligation['amount'] = self._extract_amount(content)

                    obligations.append(obligation)
                    break

        # 去重合并
        obligations = self._deduplicate_obligations(obligations)

        print(f"  ✅ 提取完成，共 {len(obligations)} 项义务")
        self._obligations = obligations
        return obligations

    def _clean_obligation_content(self, content: str) -> str:
        """清理义务内容，去除序号等"""
        # 移除开头的条款序号
        content = re.sub(r'^第[一二三四五六七八九十\d]+[条章节款][：:\s]*', '', content)
        content = re.sub(r'^\d+\.\s*', '', content)
        content = re.sub(r'^\d+\.\d+\s*', '', content)
        content = re.sub(r'^[（(][一二三四五六七八九十\d]+[）)]\s*', '', content)
        content = re.sub(r'^[一二三四五六七八九十]{1,2}[、．.]\s*', '', content)
        return content.strip()

    def _extract_performance_time(self, content: str) -> str:
        """提取履行时间"""
        time_rules = self.config.get('extraction_rules', {}).get('performance_time', [])
        for rule in time_rules:
            pattern = rule.get('pattern', '')
            match = re.search(pattern, content)
            if match:
                return match.group(0)
        return ''

    def _extract_responsible_party(self, content: str) -> str:
        """提取责任方"""
        party_rules = self.config.get('extraction_rules', {}).get('responsible_party', [])
        for rule in party_rules:
            pattern = rule.get('pattern', '')
            match = re.search(pattern, content)
            if match:
                return match.group(1)
        return ''

    def _extract_amount(self, content: str) -> str:
        """提取金额/比例"""
        amount_rules = self.config.get('extraction_rules', {}).get('amount', [])
        for rule in amount_rules:
            pattern = rule.get('pattern', '')
            match = re.search(pattern, content)
            if match:
                return match.group(0)
        return ''

    def _deduplicate_obligations(self, obligations: List[Dict]) -> List[Dict]:
        """义务去重合并"""
        dedup_config = self.config.get('deduplication', {})
        threshold = dedup_config.get('similarity_threshold', 0.8)

        unique_obligations = []
        seen_contents = []

        for obl in obligations:
            content = obl['content']

            # 检查是否与已有义务重复
            is_duplicate = False
            for i, seen in enumerate(seen_contents):
                similarity = SequenceMatcher(None, content, seen).ratio()
                if similarity >= threshold:
                    # 合并义务内容
                    if len(content) > len(seen):
                        unique_obligations[i]['content'] = content
                        unique_obligations[i]['performance_time'] = (
                            unique_obligations[i].get('performance_time', '') or obl.get('performance_time', '')
                        )
                    is_duplicate = True
                    break

            if not is_duplicate:
                unique_obligations.append(obl)
                seen_contents.append(content)

        if len(obligations) != len(unique_obligations):
            print(f"  ℹ️  合并了 {len(obligations) - len(unique_obligations)} 项重复义务")

        # 重新编号
        for i, obl in enumerate(unique_obligations):
            obl['id'] = i + 1

        return unique_obligations

    # =========================================================================
    # 阶段4: 产品匹配
    # =========================================================================

    def match_product(self, text: str, product_code: str = None) -> Optional[Dict]:
        """
        从合同文本中匹配标准产品

        Args:
            text: 合同文本
            product_code: 可选，指定产品编码（跳过自动匹配）

        Returns:
            匹配到的产品信息，未找到返回 None
        """
        print(f"\n{Colors.CYAN}📦 产品匹配中...{Colors.ENDC}")

        product_library = self.config.get('product_library', [])

        # 如果指定了产品编码，直接查找
        if product_code:
            for product in product_library:
                if product.get('product_code') == product_code:
                    print(f"  ✅ 使用指定产品: {product['product_name']} ({product_code})")
                    self._matched_product = product
                    return product
            print(f"  {Colors.YELLOW}⚠️  未找到指定产品编码: {product_code}{Colors.ENDC}")
            return None

        # 自动匹配：从文本中查找产品编码或名称
        for product in product_library:
            code = product.get('product_code', '')
            name = product.get('product_name', '')

            if code and code in text:
                print(f"  ✅ 匹配到产品: {name} (编码: {code})")
                self._matched_product = product
                return product

            if name and name in text:
                print(f"  ✅ 匹配到产品: {name}")
                self._matched_product = product
                return product

        print(f"  {Colors.YELLOW}⚠️  未自动匹配到产品，请手动指定 --product-code{Colors.ENDC}")
        return None

    # =========================================================================
    # 阶段5: 义务比对
    # =========================================================================

    def compare_obligations(self, contract_obligations: List[Dict],
                           standard_product: Dict = None) -> List[Dict]:
        """
        比对合同义务与标准产品义务

        Args:
            contract_obligations: 合同提取的义务
            standard_product: 标准产品信息

        Returns:
            差异列表
        """
        print(f"\n{Colors.CYAN}⚖️  义务比对中...{Colors.ENDC}")

        if standard_product is None:
            print(f"  ℹ️  无标准产品信息，跳过大比对")
            self._differences = []
            return []

        standard_obligations = standard_product.get('obligations', [])
        differences = []
        diff_types = self.config.get('difference_types', [])
        diff_type_map = {dt['code']: dt for dt in diff_types}

        # 检查合同中缺失的标准义务
        for std_obl in standard_obligations:
            std_type = std_obl.get('type', '')
            std_content = std_obl.get('content', '')

            found = False
            matched_contract_obl = None

            for contract_obl in contract_obligations:
                if contract_obl['type_code'] == std_type:
                    matched_contract_obl = contract_obl
                    similarity = SequenceMatcher(None, std_content, contract_obl['content']).ratio()
                    if similarity >= 0.6:
                        found = True
                        # 检查是否有内容差异
                        if similarity < 0.95:
                            diff_info = diff_type_map.get('content_changed', {})
                            differences.append({
                                'id': len(differences) + 1,
                                'type_code': 'content_changed',
                                'type_name': diff_info.get('name', '内容变更'),
                                'standard_obligation': std_content,
                                'contract_obligation': contract_obl['content'],
                                'description': diff_info.get('description', ''),
                                'suggestion': diff_info.get('suggestion', ''),
                            })
                        break

            if not found:
                diff_info = diff_type_map.get('removed', {})
                differences.append({
                    'id': len(differences) + 1,
                    'type_code': 'removed',
                    'type_name': diff_info.get('name', '删除义务'),
                    'standard_obligation': std_content,
                    'contract_obligation': matched_contract_obl['content'] if matched_contract_obl else '（无）',
                    'description': diff_info.get('description', ''),
                    'suggestion': diff_info.get('suggestion', ''),
                })

        # 检查合同中新增的义务
        for contract_obl in contract_obligations:
            found = False
            for std_obl in standard_obligations:
                if contract_obl['type_code'] == std_obl.get('type', ''):
                    similarity = SequenceMatcher(
                        None,
                        std_obl.get('content', ''),
                        contract_obl['content']
                    ).ratio()
                    if similarity >= 0.6:
                        found = True
                        break

            if not found:
                diff_info = diff_type_map.get('added', {})
                differences.append({
                    'id': len(differences) + 1,
                    'type_code': 'added',
                    'type_name': diff_info.get('name', '新增义务'),
                    'standard_obligation': '（无）',
                    'contract_obligation': contract_obl['content'],
                    'description': diff_info.get('description', ''),
                    'suggestion': diff_info.get('suggestion', ''),
                })

        print(f"  ✅ 比对完成，共发现 {len(differences)} 处差异")
        self._differences = differences
        return differences

    # =========================================================================
    # 阶段6: 风险评估
    # =========================================================================

    def assess_risk(self, obligations: List[Dict], differences: List[Dict]) -> Dict:
        """
        评估风险等级

        Args:
            obligations: 义务列表
            differences: 差异列表

        Returns:
            风险摘要
        """
        print(f"\n{Colors.CYAN}⚠️  风险评估中...{Colors.ENDC}")

        risk_config = self.config.get('risk_assessment', {})
        high_risk = risk_config.get('high_risk', {})
        medium_risk = risk_config.get('medium_risk', {})

        risk_summary = {
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'details': [],
            'overall_level': '低',
        }

        # 根据差异类型评估风险
        for diff in differences:
            diff_type = diff['type_code']

            if diff_type == 'removed':
                # 删除标准义务 → 高风险
                risk_summary['high_count'] += 1
                risk_summary['details'].append({
                    'obligation': diff['standard_obligation'],
                    'risk_level': '高',
                    'reason': '删除了标准产品定义的义务',
                    'legal_basis': high_risk.get('legal_basis', []),
                })
            elif diff_type == 'content_changed':
                # 内容变更 → 评估差异程度
                risk_summary['medium_count'] += 1
                risk_summary['details'].append({
                    'obligation': diff['standard_obligation'],
                    'risk_level': '中',
                    'reason': '义务内容与标准定义存在差异',
                    'legal_basis': medium_risk.get('legal_basis', []),
                })
            elif diff_type == 'added':
                # 新增义务 → 中低风险
                risk_summary['medium_count'] += 1
                risk_summary['details'].append({
                    'obligation': diff['contract_obligation'],
                    'risk_level': '中',
                    'reason': '合同中新增了标准产品未约定的义务',
                    'legal_basis': medium_risk.get('legal_basis', []),
                })

        # 计算整体风险等级
        if risk_summary['high_count'] > 0:
            risk_summary['overall_level'] = '高'
        elif risk_summary['medium_count'] > 2:
            risk_summary['overall_level'] = '高'
        elif risk_summary['medium_count'] > 0:
            risk_summary['overall_level'] = '中'

        # 为每项义务添加风险等级
        for obl in obligations:
            obl['risk_level'] = '低'
            obl['legal_basis'] = ''

            # 检查是否在风险详情中
            for detail in risk_summary['details']:
                if detail['obligation'] in obl['content']:
                    obl['risk_level'] = detail['risk_level']
                    obl['legal_basis'] = '; '.join(detail['legal_basis'])
                    break

        level_color = Colors.RED if risk_summary['overall_level'] == '高' else \
                      Colors.YELLOW if risk_summary['overall_level'] == '中' else Colors.GREEN

        print(f"  {level_color}✅ 评估完成: {risk_summary['overall_level']}风险 "
              f"(高风险: {risk_summary['high_count']}, 中风险: {risk_summary['medium_count']}){Colors.ENDC}")

        self._risk_summary = risk_summary
        return risk_summary

    # =========================================================================
    # 阶段7: 导出 Excel
    # =========================================================================

    def export_excel(self, output_path: str, obligations: List[Dict],
                     differences: List[Dict], clauses: List[Dict],
                     product: Dict = None) -> str:
        """
        导出 Excel 文件

        Args:
            output_path: 输出文件路径
            obligations: 义务列表
            differences: 差异列表
            clauses: 条款列表
            product: 产品信息

        Returns:
            输出文件路径
        """
        print(f"\n{Colors.CYAN}📊 导出 Excel 中...{Colors.ENDC}")

        excel_config = self.config.get('excel_output', {})

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            wb = Workbook()

            # 样式定义
            header_font = Font(bold=True, size=12, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            yellow_fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
            green_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # ============================================
            # Sheet1: 履约义务汇总
            # ============================================
            ws1 = wb.active
            ws1.title = excel_config.get('sheet_names', {}).get('obligation_summary', '履约义务汇总')

            headers1 = ['序号', '义务类型', '义务内容', '履行时间', '责任方', '金额/比例', '风险等级', '法条依据', '备注']
            for col, header in enumerate(headers1, 1):
                cell = ws1.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_align
                cell.border = thin_border

            for row, obl in enumerate(obligations, 2):
                risk_level = obl.get('risk_level', '低')
                risk_fill = red_fill if risk_level == '高' else yellow_fill if risk_level == '中' else green_fill
                risk_emoji = '🔴' if risk_level == '高' else '🟡' if risk_level == '中' else '🟢'

                row_data = [
                    obl['id'],
                    obl['type_name'],
                    obl['content'],
                    obl.get('performance_time', ''),
                    obl.get('responsible_party', ''),
                    obl.get('amount', ''),
                    f"{risk_emoji} {risk_level}风险",
                    obl.get('legal_basis', ''),
                    ''
                ]

                for col, value in enumerate(row_data, 1):
                    cell = ws1.cell(row=row, column=col, value=value)
                    cell.alignment = left_align if col == 3 else center_align
                    cell.border = thin_border
                    if col == 7:
                        cell.fill = risk_fill

            # 调整列宽
            col_widths = [8, 12, 50, 20, 10, 15, 12, 25, 15]
            for i, width in enumerate(col_widths, 1):
                ws1.column_dimensions[get_column_letter(i)].width = width

            # ============================================
            # Sheet2: 差异比对报告
            # ============================================
            ws2 = wb.create_sheet(excel_config.get('sheet_names', {}).get('difference_report', '差异比对报告'))

            headers2 = ['序号', '标准义务', '合同义务', '差异类型', '差异说明', '建议处理']
            for col, header in enumerate(headers2, 1):
                cell = ws2.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_align
                cell.border = thin_border

            for row, diff in enumerate(differences, 2):
                diff_type = diff['type_code']
                type_fill = red_fill if diff_type == 'removed' else yellow_fill

                row_data = [
                    diff['id'],
                    diff['standard_obligation'],
                    diff['contract_obligation'],
                    diff['type_name'],
                    diff.get('description', ''),
                    diff.get('suggestion', '')
                ]

                for col, value in enumerate(row_data, 1):
                    cell = ws2.cell(row=row, column=col, value=value)
                    cell.alignment = left_align if col in [2, 3, 5, 6] else center_align
                    cell.border = thin_border
                    if col == 4:
                        cell.fill = type_fill

            col_widths2 = [8, 40, 40, 12, 30, 30]
            for i, width in enumerate(col_widths2, 1):
                ws2.column_dimensions[get_column_letter(i)].width = width

            # ============================================
            # Sheet3: 合同条款原文
            # ============================================
            ws3 = wb.create_sheet(excel_config.get('sheet_names', {}).get('original_clauses', '合同条款原文'))

            headers3 = ['条款编号', '分类标签', '条款内容']
            for col, header in enumerate(headers3, 1):
                cell = ws3.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_align
                cell.border = thin_border

            for row, clause in enumerate(clauses, 2):
                row_data = [
                    clause['id'],
                    clause.get('category', '其他条款'),
                    clause['content']
                ]

                for col, value in enumerate(row_data, 1):
                    cell = ws3.cell(row=row, column=col, value=value)
                    cell.alignment = left_align if col == 3 else center_align
                    cell.border = thin_border

            col_widths3 = [10, 15, 80]
            for i, width in enumerate(col_widths3, 1):
                ws3.column_dimensions[get_column_letter(i)].width = width

            # 保存文件
            output = Path(output_path)
            if not output.parent.exists():
                output.parent.mkdir(parents=True)

            wb.save(str(output))

            print(f"  ✅ Excel 导出完成: {output_path}")
            return output_path

        except ImportError:
            raise ObligationSplitError("请安装 openpyxl: pip install openpyxl")
        except Exception as e:
            raise ObligationSplitError(f"Excel 导出失败: {e}")

    # =========================================================================
    # 阶段8: 生成审核报告
    # =========================================================================

    def generate_audit_report(self, obligations: List[Dict], differences: List[Dict],
                              risk_summary: Dict, product: Dict = None) -> str:
        """
        生成审核报告（Markdown 格式）

        Args:
            obligations: 义务列表
            differences: 差异列表
            risk_summary: 风险摘要
            product: 产品信息

        Returns:
            Markdown 格式报告
        """
        print(f"\n{Colors.CYAN}📋 生成审核报告中...{Colors.ENDC}")

        report_lines = []
        report_lines.append("# 履约义务拆分审核报告\n")

        # 基本信息
        from datetime import datetime
        report_lines.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        if product:
            report_lines.append(f"**匹配产品**: {product.get('product_name', '')} ({product.get('product_code', '')})\n")

        # 风险概览
        report_lines.append("## 📊 风险概览\n")
        level_emoji = '🔴' if risk_summary['overall_level'] == '高' else \
                      '🟡' if risk_summary['overall_level'] == '中' else '🟢'
        report_lines.append(f"**整体风险等级**: {level_emoji} {risk_summary['overall_level']}风险\n")
        report_lines.append(f"- 🔴 高风险: {risk_summary['high_count']} 项")
        report_lines.append(f"- 🟡 中风险: {risk_summary['medium_count']} 项")
        report_lines.append(f"- 🟢 低风险: {risk_summary['low_count']} 项\n")

        # 义务汇总
        report_lines.append("## 📋 履约义务汇总\n")
        report_lines.append("| 序号 | 义务类型 | 义务内容 | 履行时间 | 责任方 | 风险等级 |")
        report_lines.append("|------|---------|---------|---------|--------|---------|")

        for obl in obligations:
            risk_emoji = '🔴' if obl.get('risk_level') == '高' else \
                         '🟡' if obl.get('risk_level') == '中' else '🟢'
            report_lines.append(
                f"| {obl['id']} | {obl['type_name']} | {obl['content'][:30]}... | "
                f"{obl.get('performance_time', '')} | {obl.get('responsible_party', '')} | "
                f"{risk_emoji} {obl.get('risk_level', '低')} |"
            )

        report_lines.append("")

        # 差异详情
        if differences:
            report_lines.append("## ⚠️ 差异比对详情\n")
            for diff in differences:
                diff_emoji = '🔴' if diff['type_code'] == 'removed' else '🟡'
                report_lines.append(f"### {diff_emoji} {diff['type_name']}\n")
                report_lines.append(f"**标准义务**: {diff['standard_obligation']}\n")
                report_lines.append(f"**合同义务**: {diff['contract_obligation']}\n")
                report_lines.append(f"**说明**: {diff.get('description', '')}\n")
                report_lines.append(f"**建议**: {diff.get('suggestion', '')}\n")

        # 处理建议
        report_lines.append("## 💡 处理建议\n")

        if risk_summary['overall_level'] == '高':
            report_lines.append("> ⚠️ **高风险预警**: 建议审慎评估，与对方协商调整后再签署。\n")
            for detail in risk_summary.get('details', []):
                if detail['risk_level'] == '高':
                    report_lines.append(f"- 🔴 **{detail['reason']}**: {detail['obligation'][:50]}...")
        elif risk_summary['overall_level'] == '中':
            report_lines.append("> ℹ️ **中等风险**: 建议与对方协商调整，或评估风险接受度后决定。\n")
        else:
            report_lines.append("> ✅ **低风险**: 义务条款基本符合标准，可正常推进。\n")

        report_lines.append("")

        # 法条依据
        report_lines.append("## 📜 相关法条依据\n")
        report_lines.append("- 《中华人民共和国民法典》第509条（全面履行义务）")
        report_lines.append("- 《中华人民共和国民法典》第577条（违约责任）")
        report_lines.append("- 《中华人民共和国民法典》第510条（合同约定不明）")

        report = '\n'.join(report_lines)
        print(f"  ✅ 审核报告生成完成")
        return report

    # =========================================================================
    # 完整流程入口
    # =========================================================================

    def run_full_pipeline(self, input_file: str, output_file: str = None,
                          product_code: str = None, preview: bool = False) -> Dict:
        """
        运行完整的履约义务拆分流程

        Args:
            input_file: 输入合同文件路径
            output_file: 输出 Excel 文件路径（可选）
            product_code: 指定产品编码（可选，跳过自动匹配）
            preview: 是否仅预览，不导出文件

        Returns:
            完整处理结果字典
        """
        print(f"\n{Colors.BOLD}{Colors.CYAN}═════════════════════════════════════════{Colors.ENDC}")
        print(f"{Colors.BOLD}{Colors.CYAN}  履约义务拆分 - 完整流程{Colors.ENDC}")
        print(f"{Colors.BOLD}{Colors.CYAN}═════════════════════════════════════════{Colors.ENDC}")

        try:
            # 阶段1: 文本提取
            text = self.extract_text(input_file)

            # 阶段2: 条款拆分
            clauses = self.split_clauses(text)

            # 阶段3: 义务提取
            obligations = self.extract_obligations(clauses)

            # 阶段4: 产品匹配
            product = self.match_product(text, product_code)

            # 阶段5: 义务比对
            differences = self.compare_obligations(obligations, product)

            # 阶段6: 风险评估
            risk_summary = self.assess_risk(obligations, differences)

            # 生成审核报告
            audit_report = self.generate_audit_report(obligations, differences, risk_summary, product)

            # 导出 Excel
            excel_path = None
            if not preview and output_file:
                excel_path = self.export_excel(output_file, obligations, differences, clauses, product)

                # 保存审核报告
                report_path = Path(output_file).with_suffix('.md')
                with open(report_path, 'w', encoding='utf-8') as f:
                    f.write(audit_report)
                print(f"  ✅ 审核报告已保存: {report_path}")

            # 预览模式：打印摘要
            if preview or not output_file:
                print(f"\n{Colors.CYAN}📋 处理结果预览{Colors.ENDC}")
                print(f"  义务数量: {len(obligations)}")
                print(f"  差异数量: {len(differences)}")
                print(f"  整体风险: {risk_summary['overall_level']}")

                print(f"\n{Colors.CYAN}义务列表:{Colors.ENDC}")
                for obl in obligations:
                    risk_emoji = '🔴' if obl.get('risk_level') == '高' else \
                                 '🟡' if obl.get('risk_level') == '中' else '🟢'
                    print(f"  {obl['id']}. [{obl['type_name']}] {risk_emoji} {obl['content'][:50]}...")

            # 汇总结果
            result = {
                'success': True,
                'input_file': input_file,
                'output_file': excel_path,
                'obligations': obligations,
                'clauses': clauses,
                'differences': differences,
                'matched_product': product,
                'risk_summary': risk_summary,
                'audit_report': audit_report,
            }

            print(f"\n{Colors.GREEN}✅ 处理完成！{Colors.ENDC}")
            return result

        except ObligationSplitError as e:
            print(f"\n{Colors.RED}❌ 处理失败: {e}{Colors.ENDC}")
            return {
                'success': False,
                'error': str(e),
            }
        except Exception as e:
            print(f"\n{Colors.RED}❌ 未知错误: {e}{Colors.ENDC}")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'error': str(e),
            }

# =============================================================================
# 命令行入口
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description='履约义务拆分工具')
    parser.add_argument('--input', '-i', help='输入合同文件路径')
    parser.add_argument('--output', '-o', help='输出 Excel 文件路径')
    parser.add_argument('--product-code', '-p', help='指定产品编码（跳过自动匹配）')
    parser.add_argument('--preview', action='store_true', help='仅预览结果，不导出文件')
    parser.add_argument('--check-deps', action='store_true', help='检查依赖')
    parser.add_argument('--test', action='store_true', help='运行测试')

    args = parser.parse_args()

    splitter = ObligationSplitter()

    if args.check_deps:
        splitter.check_dependencies()
        return

    if args.test:
        print("测试模式：检查依赖后运行简单测试...")
        deps = splitter.check_dependencies()
        if deps['all_passed']:
            print("\n✅ 依赖检查通过，可以正常使用！")
        else:
            print("\n❌ 请先安装缺失的依赖")
        return

    if not args.input:
        parser.print_help()
        print("\n示例:")
        print("  python3 obligation_splitter.py --input 合同.pdf --output 义务明细表.xlsx")
        print("  python3 obligation_splitter.py --input 合同.pdf --preview")
        print("  python3 obligation_splitter.py --input 合同.pdf --product-code AS-SERVER-001")
        return

    # 运行完整流程
    splitter.run_full_pipeline(
        input_file=args.input,
        output_file=args.output,
        product_code=args.product_code,
        preview=args.preview
    )

if __name__ == '__main__':
    main()
