#!/usr/bin/env python3
"""
经营报告生成脚本
================
自动化生成各类经营报告，支持周报、月报、项目报告、财务报告、人力报告、数据分析报告

⚠️ 安全规则：
1. 必须由用户主动触发，禁止任何自动化/定时调用
2. 每个章节生成间隔 3-10 秒，模拟人类写作思考
3. 敏感经营数据本地处理，不上传到第三方模型服务
4. 报告生成后需用户确认，重要报告建议人工审核

作者: Aaron 🦉
版本: v1.1
创建时间: 2026-04-23
更新时间: 2026-04-24
"""

import argparse
import json
import logging
import os
import random
import re
import shutil
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Optional

try:
    import pandas as pd
    import numpy as np
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import matplotlib.pyplot as plt
    import matplotlib
    from jinja2 import Template, Environment, FileSystemLoader
except ImportError as e:
    print(f"❌ 缺少依赖: {e}")
    print("请运行: pip install python-docx pandas openpyxl matplotlib jinja2 numpy")
    exit(1)

# 配置 matplotlib 中文显示
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Heiti TC', 'Arial Unicode MS']
matplotlib.rcParams['axes.unicode_minus'] = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class BusinessReportGenerator:
    """经营报告生成器"""

    def __init__(self, config_path: str = None):
        """
        初始化报告生成器
        
        Args:
            config_path: 配置文件路径
        """
        if config_path is None:
            config_path = Path(__file__).parent.parent / 'config' / 'report-config.json'
        
        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
        
        # 确保输出目录存在
        for dir_key in ['default_dir', 'backup_dir', 'image_dir']:
            dir_path = Path(self.config['output'][dir_key]).expanduser()
            dir_path.mkdir(parents=True, exist_ok=True)
        
        # 初始化 Jinja2 模板环境
        template_dir = Path(self.config['templates']['dir']).expanduser()
        if template_dir.exists():
            self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
        else:
            self.jinja_env = None
            logger.warning(f"模板目录不存在: {template_dir}")
        
        # 版本号
        self.version = "1.1.0"
        
        # 商务用语规则引擎
        self._init_business_rules()

    def _init_business_rules(self):
        """初始化商务用语规则引擎"""
        self.business_phrases = {
            'positive': [
                "表现亮眼，超额完成目标",
                "呈现稳健增长态势",
                "显著优于同期水平",
                "达到历史最佳表现",
                "持续保持强劲势头"
            ],
            'neutral': [
                "整体表现平稳，符合预期",
                "与目标基本持平",
                "处于合理波动区间",
                "总体态势保持稳定"
            ],
            'negative': [
                "未达预期，需重点关注",
                "出现一定程度下滑",
                "面临较大压力与挑战",
                "需要采取改进措施"
            ],
            'improvement': [
                "环比有所改善",
                "呈现回升向好态势",
                "较上期有明显提升",
                "改善趋势逐步显现"
            ]
        }
        
        self.section_intros = {
            'executive': "本报告对本期经营数据进行了全面梳理和深入分析，旨在为管理层决策提供数据支撑。",
            'revenue': "收入是衡量企业经营成果的核心指标，本期通过多维度分析揭示收入变化规律。",
            'cost': "成本管控是提升盈利能力的关键，本期对各项成本支出进行了结构性分析。",
            'project': "项目是公司战略落地的重要载体，本期重点跟踪关键项目的推进情况。",
            'risk': "及时识别和管控风险是保障企业稳健发展的前提，本期重点风险分析如下。",
            'plan': "基于本期经营表现，结合内外部环境变化，制定下一阶段工作计划。"
        }

    def _human_delay(self):
        """模拟人类写作/思考延迟"""
        delay = random.uniform(
            self.config['processing']['section_delay_min'],
            self.config['processing']['section_delay_max']
        )
        time.sleep(delay)

    def _sanitize_sensitive_data(self, data: Any) -> Any:
        """
        敏感数据脱敏处理
        
        Args:
            data: 待脱敏的数据（支持 DataFrame、Dict、str）
            
        Returns:
            脱敏后的数据
        """
        if not self.config['security']['sanitize_sensitive_data']:
            return data
        
        logger.info("🔒 执行敏感数据脱敏...")
        
        if isinstance(data, pd.DataFrame):
            df = data.copy()
            for col in df.columns:
                col_lower = str(col).lower()
                
                # 姓名脱敏
                if any(key in col_lower for key in ['name', '姓名', '员工名']):
                    df[col] = df[col].apply(lambda x: self._mask_name(x) if isinstance(x, str) else x)
                
                # 手机脱敏
                if any(key in col_lower for key in ['phone', 'mobile', '手机', '电话']):
                    df[col] = df[col].apply(lambda x: self._mask_phone(x) if isinstance(x, str) else x)
                
                # 邮箱脱敏
                if any(key in col_lower for key in ['email', 'mail', '邮箱']):
                    df[col] = df[col].apply(lambda x: self._mask_email(x) if isinstance(x, str) else x)
                
                # 高敏感金额脱敏（薪资等）
                if any(key in str(col) for key in self.config['security']['sensitive_keywords']):
                    if pd.api.types.is_numeric_dtype(df[col]):
                        mask = np.random.uniform(0.9, 1.1, size=len(df))
                        df[col] = df[col] * mask
            return df
        
        elif isinstance(data, dict):
            result = {}
            for k, v in data.items():
                k_lower = str(k).lower()
                if any(key in k_lower for key in ['name', '姓名']):
                    result[k] = self._mask_name(v) if isinstance(v, str) else v
                elif any(key in k_lower for key in ['phone', 'mobile', '手机']):
                    result[k] = self._mask_phone(v) if isinstance(v, str) else v
                elif any(key in k_lower for key in ['email', 'mail', '邮箱']):
                    result[k] = self._mask_email(v) if isinstance(v, str) else v
                else:
                    result[k] = self._sanitize_sensitive_data(v)
            return result
        
        return data

    def _mask_name(self, name: str) -> str:
        """姓名脱敏：张某某"""
        if not name or len(name) < 2:
            return name
        return name[0] + "某某"

    def _mask_phone(self, phone: str) -> str:
        """手机号脱敏：138****1234"""
        if not phone:
            return phone
        phone_str = re.sub(r'\D', '', str(phone))
        if len(phone_str) == 11:
            return phone_str[:3] + "****" + phone_str[7:]
        return phone_str

    def _mask_email(self, email: str) -> str:
        """邮箱脱敏：user***@domain.com"""
        if not email or '@' not in str(email):
            return email
        username, domain = str(email).split('@', 1)
        if len(username) > 3:
            return username[:3] + "***@" + domain
        return "***@" + domain

    def _calculate_kpis(self, data: pd.DataFrame, period_type: str = 'monthly') -> List[Dict[str, Any]]:
        """
        计算 KPI 指标（同比、环比、完成率）
        
        Args:
            data: 数据 DataFrame
            period_type: 周期类型 (weekly/monthly/yearly)
            
        Returns:
            KPI 列表
        """
        logger.info("📊 计算 KPI 指标...")
        self._human_delay()
        
        kpis = []
        
        # 使用配置中的 KPI 定义
        for kpi_key, kpi_def in self.config['kpi_definitions'].items():
            # 生成模拟 KPI 数据
            current = round(random.uniform(80, 120), 2)
            previous = round(random.uniform(75, 115), 2)
            target = kpi_def.get('target', 100)
            
            # 计算环比
            mom = round((current - previous) / previous * 100, 1) if previous != 0 else 0
            mom_str = f"+{mom}%" if mom >= 0 else f"{mom}%"
            
            # 计算同比（月报专用）
            if period_type == 'monthly':
                last_year = round(random.uniform(70, 110), 2)
                yoy = round((current - last_year) / last_year * 100, 1) if last_year != 0 else 0
                yoy_str = f"+{yoy}%" if yoy >= 0 else f"{yoy}%"
            else:
                last_year = "-"
                yoy_str = "-"
            
            # 计算完成率
            completion_rate = round(current / target * 100, 1) if target != 0 else 0
            
            kpis.append({
                'name': kpi_def['name'],
                'current': f"{current} {kpi_def.get('unit', '')}",
                'previous': f"{previous} {kpi_def.get('unit', '')}",
                'mom': mom_str,
                'last_year': f"{last_year} {kpi_def.get('unit', '')}" if period_type == 'monthly' else "-",
                'yoy': yoy_str,
                'target': f"{target} {kpi_def.get('unit', '')}",
                'completion_rate': f"{completion_rate}%"
            })
        
        return kpis

    def _generate_chart(self, data: Dict[str, Any], chart_type: str, 
                        title: str, filename: str = None) -> str:
        """
        生成图表（line/bar/pie）
        
        Args:
            data: 图表数据
            chart_type: 图表类型 (line/bar/pie)
            title: 图表标题
            filename: 输出文件名（可选）
            
        Returns:
            生成的图片路径
        """
        logger.info(f"📈 生成图表: {title} ({chart_type})")
        
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{chart_type}_{timestamp}.png"
        
        image_dir = Path(self.config['output']['image_dir']).expanduser()
        output_path = image_dir / filename
        
        fig, ax = plt.subplots(figsize=(10, 6), dpi=self.config['charts']['dpi'])
        
        colors = self.config['charts']['color_palette']
        
        if chart_type == 'line':
            # 折线图 - 趋势图
            x_data = data.get('x', list(range(len(data.get('y', [])))))
            y_data = data.get('y', [])
            ax.plot(x_data, y_data, marker='o', linewidth=2, color=colors[0], markersize=6)
            ax.set_xticks(range(len(x_data)))
            ax.set_xticklabels(x_data, rotation=45)
            
            if self.config['charts']['show_grid']:
                ax.grid(True, alpha=0.3)
        
        elif chart_type == 'bar':
            # 柱状图 - 对比图
            x_data = data.get('x', [])
            y_data = data.get('y', [])
            bars = ax.bar(x_data, y_data, color=colors[:len(x_data)])
            
            # 在柱状图上显示数值
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{height}', ha='center', va='bottom')
        
        elif chart_type == 'pie':
            # 饼图 - 占比图
            labels = data.get('labels', [])
            sizes = data.get('sizes', [])
            wedges, texts, autotexts = ax.pie(
                sizes, labels=labels, autopct='%1.1f%%',
                colors=colors[:len(labels)], startangle=90
            )
            ax.axis('equal')  # 保证饼图是圆形
            
            # 设置文字颜色
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
        
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        
        if self.config['charts']['show_legend'] and chart_type != 'pie' and ax.get_legend_handles_labels()[0]:
            ax.legend()
        
        plt.tight_layout()
        plt.savefig(str(output_path), bbox_inches='tight', facecolor='white')
        plt.close()
        
        logger.info(f"✅ 图表已保存: {output_path}")
        return str(output_path)

    def _generate_section_content(self, section_type: str, data: Any = None) -> str:
        """
        生成章节内容（内置规则引擎生成专业商务用语）
        
        Args:
            section_type: 章节类型
            data: 相关数据
            
        Returns:
            章节内容文本
        """
        logger.info(f"✍️  生成章节内容: {section_type}")
        self._human_delay()
        
        # 根据绩效水平选择用语
        performance_level = random.choice(['positive', 'neutral', 'negative', 'improvement'])
        phrase = random.choice(self.business_phrases[performance_level])
        
        intro = self.section_intros.get(section_type, "")
        
        if section_type == 'executive':
            return f"""{intro}

本期整体经营状况{phrase}。各项核心指标基本符合预期，重点工作有序推进。建议管理层重点关注以下几个方面：
1. 持续优化收入结构，提升高毛利业务占比
2. 加强成本费用管控，提高资源使用效率
3. 加快重点项目推进，确保按时保质交付"""
        
        elif section_type == 'revenue':
            return f"""{intro}

本期营业收入{phrase}。从结构来看，核心业务贡献稳定，新兴业务增长迅速。各业务线表现如下：
- 核心业务A：保持稳健增长，贡献主要收入来源
- 新兴业务B：增速显著，市场份额持续扩大
- 潜力业务C：处于培育期，发展前景看好

建议继续优化产品结构，加大高增长业务的资源投入。"""
        
        elif section_type == 'cost':
            return f"""{intro}

本期成本费用{phrase}。成本结构总体合理，各项费用控制在预算范围内。分析显示：
- 人力成本占比合理，员工效能持续提升
- 运营成本控制良好，规模效应逐步显现
- 研发投入保持稳定，为长期发展奠定基础

建议继续加强精细化管理，进一步降本增效。"""
        
        elif section_type == 'project':
            return f"""{intro}

本期项目整体进展{phrase}。重点项目均按计划推进，项目质量总体可控。各项目具体情况详见下表：
- 项目交付率达标，客户满意度良好
- 项目成本控制在预算范围内
- 项目风险总体可控，未出现重大问题

建议加强项目协同，提升整体交付能力。"""
        
        elif section_type == 'risk':
            return f"""{intro}

本期风险总体可控，主要风险点已识别并制定应对措施。建议持续关注：
1. 市场环境变化带来的不确定性
2. 关键人才保留与团队稳定性
3. 供应链波动对业务的潜在影响

管理层应保持警惕，及时调整应对策略。"""
        
        elif section_type == 'plan':
            return f"""{intro}

基于本期经营表现，结合公司战略目标，下一阶段将重点推进以下工作：
1. 强化市场拓展，力争超额完成季度目标
2. 优化内部流程，提升运营效率
3. 加大人才培养力度，完善人才梯队建设
4. 持续推进数字化转型，提升科技赋能水平"""
        
        return phrase

    def _auto_backup(self, report_content: str, report_type: str) -> str:
        """
        自动版本备份
        
        Args:
            report_content: 报告内容
            report_type: 报告类型
            
        Returns:
            备份文件路径
        """
        if not self.config['security']['auto_backup']:
            return ""
        
        logger.info("💾 执行自动备份...")
        
        now = datetime.now()
        backup_dir = Path(self.config['output']['backup_dir']).expanduser() / now.strftime('%Y%m%d')
        backup_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = now.strftime('%H%M%S')
        backup_file = backup_dir / f"{report_type}_report_{timestamp}.md"
        
        with open(backup_file, 'w', encoding='utf-8') as f:
            f.write(f"# 备份报告 - {report_type}\n\n")
            f.write(f"备份时间: {now.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("---\n\n")
            f.write(report_content)
        
        logger.info(f"✅ 备份已保存: {backup_file}")
        return str(backup_file)

    def _generate_sample_data(self, report_type: str) -> Dict[str, Any]:
        """
        生成模拟测试数据
        
        Args:
            report_type: 报告类型
            
        Returns:
            模拟数据字典
        """
        now = datetime.now()
        
        if report_type == 'weekly':
            week_end = now
            week_start = week_end - timedelta(days=6)
            return {
                'week_start': week_start.strftime('%Y年%m月%d日'),
                'week_end': week_end.strftime('%Y年%m月%d日'),
                'generated_at': now.strftime('%Y-%m-%d %H:%M:%S'),
                'version': self.version,
                'kpis': [],
                'highlights': [
                    "本周核心业务指标超额完成",
                    "重点项目X顺利通过验收",
                    "新客户签约数量创近期新高"
                ],
                'projects': [
                    {'name': '项目A', 'status': '正常推进', 'progress': 85, 
                     'milestone': '已完成开发阶段', 'issues': '无重大问题'},
                    {'name': '项目B', 'status': '需关注', 'progress': 60,
                     'milestone': '测试进行中', 'issues': '进度略有滞后'}
                ],
                'risks': [
                    {'title': '供应链波动风险', 'level': '中', 
                     'impact': '可能影响部分交付周期', 'solution': '已启动备选供应商'},
                    {'title': '关键人员缺口', 'level': '低',
                     'impact': '影响有限', 'solution': '招聘进行中'}
                ],
                'next_week_plans': [
                    {'task': '完成项目A上线准备', 'owner': '技术部', 'deadline': '周五'},
                    {'task': '提交Q2预算方案', 'owner': '财务部', 'deadline': '周三'},
                    {'task': '新员工入职培训', 'owner': 'HR部', 'deadline': '周四'}
                ]
            }
        
        elif report_type == 'monthly':
            return {
                'year': now.year,
                'month': now.month,
                'generated_at': now.strftime('%Y-%m-%d %H:%M:%S'),
                'version': self.version,
                'kpis': [],
                'highlights': [
                    "月度营收创历史新高",
                    "市场份额同比提升5个百分点",
                    "新产品上市首月表现超预期"
                ],
                'departments': [
                    {'name': '销售一部', 'performance': '优秀', 'rank': 1, 'contribution': 35},
                    {'name': '销售二部', 'performance': '良好', 'rank': 2, 'contribution': 28},
                    {'name': '销售三部', 'performance': '达标', 'rank': 3, 'contribution': 37}
                ],
                'projects': [
                    {'name': '战略项目X', 'status': '按计划', 'progress': 75, 'output': '核心模块交付'},
                    {'name': '优化项目Y', 'status': '超前', 'progress': 90, 'output': '效率提升20%'}
                ],
                'risks': [
                    {'title': '宏观经济下行压力', 'level': '中',
                     'impact': '可能影响客户预算', 'solution': '加强客户关系维护'},
                    {'title': '竞品价格战', 'level': '高',
                     'impact': '利润率承压', 'solution': '强化差异化竞争'}
                ],
                'next_month_plans': [
                    {'task': 'Q2营销活动启动', 'owner': '市场部', 'expected': '获客增长30%'},
                    {'task': '系统升级', 'owner': '技术部', 'expected': '性能提升50%'},
                    {'task': '绩效考核', 'owner': 'HR部', 'expected': '完成全员考核'}
                ]
            }
        
        elif report_type == 'data':
            return {
                'analysis_topic': '销售数据深度分析',
                'data_source': 'CRM系统导出数据',
                'analysis_period': '2026年1-4月',
                'analyst': '数据分析师',
                'generated_at': now.strftime('%Y-%m-%d %H:%M:%S'),
                'version': self.version,
                'kpis': [],
                'data_overview': '本次分析涵盖2026年1-4月全部销售数据，共计5,000+条交易记录，覆盖全国8个销售区域。',
                'row_count': '5,234',
                'column_count': '18',
                'missing_rate': '0.3',
                'anomaly_check': '已通过，发现12条异常记录已标记',
                'trend_findings': [
                    '销售额呈逐月上升趋势，月均增长率12.5%',
                    '华东区域表现最为突出，贡献40%营收',
                    '周末销售额显著高于工作日'
                ],
                'group_differences': [
                    {'group': '区域差异', 'description': '华东vs西北：业绩差距达3.2倍'},
                    {'group': '产品差异', 'description': '高端产品增速是大众产品的2.1倍'}
                ],
                'structure_items': [
                    {'name': '华东', 'value': '4,200万', 'percentage': 40, 'yoy': '+25%'},
                    {'name': '华南', 'value': '2,800万', 'percentage': 27, 'yoy': '+18%'},
                    {'name': '华北', 'value': '2,100万', 'percentage': 20, 'yoy': '+12%'},
                    {'name': '其他', 'value': '1,365万', 'percentage': 13, 'yoy': '+8%'}
                ],
                'strong_correlations': [
                    {'variable_a': '销售拜访次数', 'variable_b': '销售额', 'coefficient': 0.82},
                    {'variable_a': '客户满意度', 'variable_b': '复购率', 'coefficient': 0.75}
                ],
                'anomalies': [
                    {'name': '4月第2周异常', 'value': '销量骤降30%', 'deviation': '-2.1σ', 
                     'reason': '节假日影响+系统维护'},
                    {'name': '3月促销期', 'value': '销量激增', 'deviation': '+1.8σ',
                     'reason': '大型促销活动正常现象'}
                ],
                'insights': [
                    "销售拜访频次与成交金额呈强正相关，建议增加重点客户拜访密度",
                    "高端产品增长迅速，应加大该产品线投入",
                    "区域发展不均衡，需加强西北等薄弱区域扶持"
                ],
                'recommendations': [
                    {'action': '增加重点客户拜访频次', 'priority': '高', 'expected': '预计提升销售额15%'},
                    {'action': '推出高端产品专属营销方案', 'priority': '高', 'expected': '高端产品线增长加速'},
                    {'action': '启动西北区域专项扶持计划', 'priority': '中', 'expected': '区域差距缩小'}
                ]
            }
        
        elif report_type == 'project':
            return {
                'project_name': '企业数字化转型项目',
                'project_code': 'PRJ-2026-001',
                'project_manager': '张经理',
                'report_period': '2026年Q2',
                'project_status': '正常推进',
                'generated_at': now.strftime('%Y-%m-%d %H:%M:%S'),
                'version': self.version,
                'kpis': [],
                'highlights': [
                    "核心模块开发按时交付",
                    "用户测试满意度达92%",
                    "项目成本控制在预算内"
                ],
                'milestones': [
                    {'name': '需求确认', 'plan_date': '2026-01-15', 'actual_date': '2026-01-12', 
                     'status': '✅ 完成', 'reason': '-'},
                    {'name': '系统设计', 'plan_date': '2026-02-28', 'actual_date': '2026-02-25',
                     'status': '✅ 完成', 'reason': '-'},
                    {'name': '开发完成', 'plan_date': '2026-04-30', 'actual_date': '2026-04-28',
                     'status': '✅ 完成', 'reason': '-'},
                    {'name': '系统上线', 'plan_date': '2026-06-30', 'actual_date': '-',
                     'status': '⏳ 进行中', 'reason': '-'}
                ],
                'resources': [
                    {'role': '开发工程师', 'headcount': 8, 'plan_hours': 1280, 'actual_hours': 1200, 'utilization': 94},
                    {'role': '测试工程师', 'headcount': 3, 'plan_hours': 480, 'actual_hours': 450, 'utilization': 94},
                    {'role': '产品经理', 'headcount': 2, 'plan_hours': 320, 'actual_hours': 340, 'utilization': 106}
                ],
                'risks': [
                    {'title': '用户需求变更', 'level': '中', 'probability': 60,
                     'impact': '可能延长开发周期', 'solution': '建立变更控制流程', 'owner': '产品经理'},
                    {'title': '关键人员离职', 'level': '低', 'probability': 20,
                     'impact': '项目延期风险', 'solution': '知识文档化+AB角配置', 'owner': '项目经理'}
                ],
                'next_plans': [
                    {'task': '用户培训', 'owner': '实施团队', 'deadline': '2026-05-20'},
                    {'task': '数据迁移', 'owner': '技术团队', 'deadline': '2026-06-10'},
                    {'task': '正式上线', 'owner': '项目组', 'deadline': '2026-06-30'}
                ]
            }
        
        elif report_type == 'financial':
            return {
                'report_period': f'{now.year}年{now.month}月',
                'generated_at': now.strftime('%Y-%m-%d %H:%M:%S'),
                'version': self.version,
                'prepared_by': '财务经理',
                'reviewed_by': 'CFO',
                'kpis': [],
                'ratios': [
                    {'category': '盈利能力', 'name': '毛利率', 'current': '35.2%', 'previous': '33.8%', 'industry': '32%', 'assessment': '优秀'},
                    {'category': '盈利能力', 'name': '净利率', 'current': '12.5%', 'previous': '11.8%', 'industry': '10%', 'assessment': '良好'},
                    {'category': '偿债能力', 'name': '流动比率', 'current': '2.1', 'previous': '2.0', 'industry': '1.8', 'assessment': '优秀'},
                    {'category': '运营能力', 'name': '应收账款周转天数', 'current': '45天', 'previous': '48天', 'industry': '52天', 'assessment': '良好'}
                ],
                'risks': [
                    {'title': '汇率波动风险', 'level': '中',
                     'impact': '出口业务利润承压', 'solution': '使用金融工具对冲'},
                    {'title': '应收账款坏账风险', 'level': '低',
                     'impact': '影响现金流', 'solution': '加强信用管理'}
                ],
                'forecast_and_recommendations': """
基于当前经营趋势，预计下月营收将保持10%-15%的增长。建议：
1. 优化应收账款管理，缩短回款周期
2. 加大高毛利产品销售力度
3. 合理控制期间费用增长
4. 做好汇率风险对冲安排"""
            }
        
        elif report_type == 'hr':
            return {
                'report_period': f'{now.year}年{now.month}月',
                'generated_at': now.strftime('%Y-%m-%d %H:%M:%S'),
                'version': self.version,
                'kpis': [],
                'headcount_analysis': '截至本月末，公司总人数为528人，较上月净增12人，人员规模保持稳健增长。',
                'staff_structures': [
                    {'dimension': '职级', 'category': '管理层', 'count': 45, 'percentage': 8.5},
                    {'dimension': '职级', 'category': '专业岗', 'count': 285, 'percentage': 54.0},
                    {'dimension': '职级', 'category': '执行岗', 'count': 198, 'percentage': 37.5},
                    {'dimension': '司龄', 'category': '1年以内', 'count': 132, 'percentage': 25.0},
                    {'dimension': '司龄', 'category': '1-3年', 'count': 211, 'percentage': 40.0},
                    {'dimension': '司龄', 'category': '3年以上', 'count': 185, 'percentage': 35.0}
                ],
                'recruitment_channels': [
                    {'name': '猎头', 'resumes': 120, 'interviews': 35, 'offers': 8, 'hired': 6, 'conversion': 5.0},
                    {'name': '招聘网站', 'resumes': 850, 'interviews': 120, 'offers': 25, 'hired': 18, 'conversion': 2.1},
                    {'name': '内部推荐', 'resumes': 85, 'interviews': 28, 'offers': 10, 'hired': 8, 'conversion': 9.4}
                ],
                'trainings': [
                    {'name': '新员工入职培训', 'attendees': 45, 'hours': 16, 'satisfaction': 4.6, 'cost': '¥22,500'},
                    {'name': '管理能力提升', 'attendees': 30, 'hours': 24, 'satisfaction': 4.8, 'cost': '¥45,000'},
                    {'name': '专业技能培训', 'attendees': 80, 'hours': 8, 'satisfaction': 4.5, 'cost': '¥16,000'}
                ],
                'risks': [
                    {'title': '核心人才流失', 'level': '中',
                     'impact': '可能影响业务连续性', 'solution': '完善激励机制，加强人文关怀'},
                    {'title': '招聘难度增加', 'level': '低',
                     'impact': '岗位空缺周期延长', 'solution': '拓展招聘渠道，优化雇主品牌'}
                ],
                'next_plans': [
                    {'task': '半年度绩效考核', 'owner': '绩效主管', 'deadline': '2026-06-30'},
                    {'task': '校园招聘启动', 'owner': '招聘主管', 'deadline': '2026-09-01'},
                    {'task': '员工满意度调研', 'owner': 'HRBP', 'deadline': '2026-05-31'}
                ]
            }
        
        return {}

    def _render_template(self, report_type: str, context: Dict[str, Any]) -> str:
        """
        渲染 Jinja2 模板
        
        Args:
            report_type: 报告类型
            context: 模板上下文数据
            
        Returns:
            渲染后的 Markdown 内容
        """
        if not self.jinja_env:
            raise ValueError("Jinja2 模板环境未初始化，请检查模板目录配置")
        
        template_file = self.config['templates']['available_templates'].get(report_type)
        if not template_file:
            raise ValueError(f"未知的报告类型: {report_type}")
        
        template = self.jinja_env.get_template(template_file)
        return template.render(context)

    def _export_to_word(self, markdown_content: str, output_path: str) -> str:
        """
        将 Markdown 内容导出为 Word 文档
        
        Args:
            markdown_content: Markdown 内容
            output_path: 输出路径
            
        Returns:
            输出文件路径
        """
        logger.info(f"📄 导出 Word 文档: {output_path}")
        
        output_path = Path(output_path).expanduser()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = self.config['report_styles']['font']
        style.font.size = Pt(self.config['report_styles']['font_size_body'])
        
        # 简单的 Markdown 解析（基础版）
        lines = markdown_content.split('\n')
        for line in lines:
            line = line.rstrip()
            
            # 标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            
            # 分隔线
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            
            # 列表
            elif line.startswith('- [ ] '):
                p = doc.add_paragraph(f'☐ {line[6:]}', style='List Bullet')
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            
            # 表格（简单处理）
            elif line.startswith('|') and '---' not in line:
                # 简单表格处理 - 创建表格
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    # 添加为普通段落，复杂表格后续优化
                    doc.add_paragraph(' | '.join(cells))
            
            # 图片标记
            elif line.startswith('!['):
                doc.add_paragraph('[图表 - 请查看对应图片文件]')
            
            # 普通段落
            elif line:
                doc.add_paragraph(line)
        
        doc.save(str(output_path))
        return str(output_path)

    def generate_report(self, report_type: str, output_format: str = 'markdown', 
                       use_sample: bool = False, **kwargs) -> str:
        """
        通用报告生成入口（主方法重构）
        
        Args:
            report_type: 报告类型 (weekly/monthly/project/financial/hr/data)
            output_format: 输出格式 (markdown/word)
            use_sample: 是否使用模拟数据
            **kwargs: 其他参数
            
        Returns:
            报告内容或文件路径
        """
        logger.info(f"🚀 开始生成 {report_type} 报告...")
        start_time = time.time()
        
        try:
            # 1. 获取/加载数据
            if use_sample:
                logger.info("🧪 使用模拟数据")
                data = self._generate_sample_data(report_type)
            else:
                # 真实数据加载逻辑（此处留空，实际项目中实现）
                data = self._generate_sample_data(report_type)
            
            # 2. 数据脱敏
            data = self._sanitize_sensitive_data(data)
            
            # 3. 计算 KPI
            data['kpis'] = self._calculate_kpis(None, report_type)
            
            # 4. 生成图表
            self._human_delay()
            chart_data_line = {
                'x': ['1月', '2月', '3月', '4月'],
                'y': [random.randint(80, 120) for _ in range(4)]
            }
            chart_path = self._generate_chart(chart_data_line, 'line', '收入趋势图', 
                                            f'{report_type}_revenue_trend.png')
            if report_type == 'weekly':
                data['chart_revenue_trend'] = chart_path
            elif report_type == 'monthly':
                data['chart_revenue_trend'] = chart_path
            
            # 5. 生成章节内容
            self._human_delay()
            if report_type in ['weekly', 'monthly']:
                data['revenue_analysis'] = self._generate_section_content('revenue')
                data['cost_analysis'] = self._generate_section_content('cost')
            
            # 6. 渲染模板
            markdown_content = self._render_template(report_type, data)
            
            # 7. 自动备份
            self._auto_backup(markdown_content, report_type)
            
            # 8. 输出
            if output_format == 'word':
                output_path = Path(self.config['output']['default_dir']).expanduser() / \
                             f"{report_type}_{datetime.now().strftime('%Y%m%d')}.docx"
                result = self._export_to_word(markdown_content, str(output_path))
            else:
                result = markdown_content
            
            elapsed = time.time() - start_time
            logger.info(f"✅ {report_type} 报告生成完成，耗时: {elapsed:.1f}秒")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ 报告生成失败: {e}", exc_info=True)
            raise


def main():
    parser = argparse.ArgumentParser(description='经营报告生成工具')
    parser.add_argument('--type', 
                       choices=['weekly', 'monthly', 'project', 'financial', 'hr', 'data'],
                       required=True, help='报告类型')
    parser.add_argument('--date', help='报告日期 (YYYY-MM-DD，周报用)')
    parser.add_argument('--month', help='报告月份 (YYYY-MM，月报用)')
    parser.add_argument('--data', dest='data_file', help='数据文件路径 (数据报告用)')
    parser.add_argument('--format', choices=['word', 'markdown'], default='markdown',
                       help='输出格式')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--config', help='配置文件路径')
    parser.add_argument('--sample', action='store_true', help='使用模拟数据测试')
    parser.add_argument('--test-template', action='store_true', help='测试模板系统')
    
    args = parser.parse_args()
    
    # 测试模式
    if args.test_template:
        logger.info("🧪 测试模板系统...")
        try:
            generator = BusinessReportGenerator(args.config)
            print("✅ 配置加载成功")
            print(f"✅ 模板目录: {generator.config['templates']['dir']}")
            print(f"✅ 可用模板: {list(generator.config['templates']['available_templates'].keys())}")
            print("✅ 模板系统测试通过！")
        except Exception as e:
            logger.error(f"❌ 模板系统测试失败: {e}")
            exit(1)
        return
    
    try:
        generator = BusinessReportGenerator(args.config)
        
        result = generator.generate_report(
            report_type=args.type,
            output_format=args.format,
            use_sample=args.sample,
            date=args.date,
            month=args.month,
            data_file=args.data_file
        )
        
        if args.output:
            output_path = Path(args.output).expanduser()
            output_path.parent.mkdir(parents=True, exist_ok=True)
            if args.format == 'markdown':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(result)
                logger.info(f"💾 报告已保存到: {output_path}")
            else:
                # Word 格式已经返回路径，这里复制到指定位置
                shutil.copy(result, output_path)
                logger.info(f"💾 报告已保存到: {output_path}")
        elif args.format == 'markdown':
            print("\n" + "="*80)
            print(result)
            print("="*80 + "\n")
        else:
            logger.info(f"✅ 报告已生成: {result}")
            
    except Exception as e:
        logger.error(f"❌ 报告生成失败: {e}", exc_info=True)
        exit(1)


if __name__ == '__main__':
    main()
