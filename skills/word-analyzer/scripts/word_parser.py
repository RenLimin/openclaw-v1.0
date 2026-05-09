#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档结构解析模块
=====================
完整解析 Word 文档结构树、标题层级、段落属性、样式信息

功能:
  - 文档结构树构建（标题/段落/表格层级关系）
  - 按标题提取段落
  - 段落属性提取（字体、字号、颜色、加粗、斜体、段落间距）
  - 批量导出为 JSON

作者: Word-Analyzer Team
版本: v1.0
日期: 2026-04-24
"""

import os
import re
import json
import argparse
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
from dataclasses import dataclass, asdict

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
except ImportError:
    print("❌ 错误: 请安装 python-docx: pip install python-docx")
    exit(1)

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False


@dataclass
class ParagraphStyle:
    """段落样式数据结构"""
    font_name: str = ""
    font_size: float = 0.0
    font_color: str = ""
    bold: bool = False
    italic: bool = False
    underline: bool = False
    alignment: str = ""
    line_spacing: float = 0.0
    space_before: float = 0.0
    space_after: float = 0.0
    left_indent: float = 0.0
    right_indent: float = 0.0
    first_line_indent: float = 0.0
    style_name: str = ""


@dataclass
class HeadingInfo:
    """标题信息数据结构"""
    level: int
    text: str
    index: int
    style: ParagraphStyle
    children: List[Dict] = None


class WordParser:
    """Word 文档解析器"""

    def __init__(self, file_path: str, config_path: str = None):
        """
        初始化解析器

        Args:
            file_path: Word 文档路径 (.docx)
            config_path: 配置文件路径（可选）
        """
        self.file_path = Path(file_path)
        self.config = self._load_config(config_path)
        self.doc: Optional[DocumentType] = None
        self.heading_patterns = self._compile_heading_patterns()

    def _load_config(self, config_path: str = None) -> Dict:
        """加载配置文件"""
        default_config = {
            'heading_patterns': {
                'level_1': [r'^第[一二三四五六七八九十百千\d]+章\s', r'^第[一二三四五六七八九十百千\d]+编\s'],
                'level_2': [r'^第[一二三四五六七八九十百千\d]+条\s', r'^第[一二三四五六七八九十百千\d]+节\s', r'^\d+\.\s'],
                'level_3': [r'^\d+\.\d+\s', r'^[（(][一二三四五六七八九十\d]+[）)]\s*'],
                'level_4': [r'^\d+\.\d+\.\d+\s', r'^[一二三四五六七八九十]{1,2}[、．.]\s*'],
                'level_5': [r'^\d+\.\d+\.\d+\.\d+\s'],
            }
        }

        if config_path and HAS_YAML:
            cfg_path = Path(config_path)
            if cfg_path.exists():
                with open(cfg_path, 'r', encoding='utf-8') as f:
                    user_config = yaml.safe_load(f)
                    default_config.update(user_config)

        return default_config

    def _compile_heading_patterns(self) -> Dict[int, re.Pattern]:
        """编译标题匹配正则"""
        patterns = {}
        for level_name, patterns_list in self.config.get('heading_patterns', {}).items():
            level_num = int(level_name.replace('level_', ''))
            combined = '|'.join(patterns_list)
            patterns[level_num] = re.compile(combined)
        return patterns

    def load_document(self) -> bool:
        """加载 Word 文档"""
        try:
            self.doc = Document(str(self.file_path))
            return True
        except Exception as e:
            print(f"❌ 加载文档失败: {e}")
            return False

    def extract_paragraph_style(self, para: Paragraph) -> ParagraphStyle:
        """提取段落样式属性"""
        style = ParagraphStyle()
        style.style_name = para.style.name if para.style else ""

        # 对齐方式
        if para.alignment is not None:
            style.alignment = str(para.alignment).split('.')[-1]

        # 段落间距
        if para.paragraph_format:
            pf = para.paragraph_format
            style.line_spacing = pf.line_spacing or 0.0
            style.space_before = pf.space_before.pt if pf.space_before else 0.0
            style.space_after = pf.space_after.pt if pf.space_after else 0.0
            style.left_indent = pf.left_indent.pt if pf.left_indent else 0.0
            style.right_indent = pf.right_indent.pt if pf.right_indent else 0.0
            style.first_line_indent = pf.first_line_indent.pt if pf.first_line_indent else 0.0

        # 字体属性（取第一个 run 的属性作为段落属性）
        if para.runs:
            run = para.runs[0]
            style.font_name = run.font.name or ""
            style.font_size = run.font.size.pt if run.font.size else 0.0
            style.bold = bool(run.bold)
            style.italic = bool(run.italic)
            style.underline = bool(run.underline)
            if run.font.color and run.font.color.rgb:
                style.font_color = str(run.font.color.rgb)

        return style

    def detect_heading_level(self, para: Paragraph) -> Optional[int]:
        """检测段落的标题级别"""
        text = para.text.strip()
        if not text:
            return None

        # 先检查 Word 内置标题样式
        if para.style and para.style.name.startswith('Heading'):
            try:
                return int(para.style.name.replace('Heading ', '').replace('标题', ''))
            except ValueError:
                pass

        # 再用正则匹配
        for level in sorted(self.heading_patterns.keys()):
            pattern = self.heading_patterns[level]
            if pattern.match(text):
                return level

        return None

    def extract_all_paragraphs(self) -> List[Dict[str, Any]]:
        """提取所有段落及其属性"""
        if not self.doc:
            self.load_document()

        paragraphs = []
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style = self.extract_paragraph_style(para)
            heading_level = self.detect_heading_level(para)

            paragraphs.append({
                'index': idx,
                'text': text,
                'length': len(text),
                'heading_level': heading_level,
                'style': asdict(style),
                'is_empty': len(text) == 0
            })

        return paragraphs

    def extract_headings(self) -> List[Dict[str, Any]]:
        """提取所有标题（带层级）"""
        paragraphs = self.extract_all_paragraphs()
        headings = []

        for para in paragraphs:
            if para['heading_level'] is not None:
                headings.append({
                    'level': para['heading_level'],
                    'text': para['text'],
                    'index': para['index'],
                    'style': para['style']
                })

        return headings

    def build_structure_tree(self) -> List[Dict[str, Any]]:
        """构建文档结构树"""
        headings = self.extract_headings()
        paragraphs = self.extract_all_paragraphs()

        if not headings:
            return []

        # 构建树结构
        tree = []
        stack = []  # 用于追踪当前路径

        for heading in headings:
            level = heading['level']
            node = {
                'level': level,
                'text': heading['text'],
                'heading_index': heading['index'],
                'style': heading['style'],
                'content_paragraphs': [],
                'tables': [],
                'children': []
            }

            # 弹出层级大于等于当前级别的节点
            while stack and stack[-1]['level'] >= level:
                stack.pop()

            if stack:
                # 添加为父节点的子节点
                stack[-1]['children'].append(node)
            else:
                # 顶级节点
                tree.append(node)

            stack.append(node)

        # 填充每个标题下的内容段落
        self._fill_content_to_tree(tree, paragraphs)

        return tree

    def _fill_content_to_tree(self, tree: List[Dict], paragraphs: List[Dict]):
        """将内容段落填充到结构树中"""
        # 收集所有标题的索引
        all_heading_indices = set()
        self._collect_heading_indices(tree, all_heading_indices)

        # 按标题分段
        heading_order = sorted(all_heading_indices)

        for i, heading_idx in enumerate(heading_order):
            next_heading_idx = heading_order[i + 1] if i + 1 < len(heading_order) else len(paragraphs)

            # 找到对应的标题节点并填充内容
            content = []
            for para in paragraphs:
                if heading_idx < para['index'] < next_heading_idx:
                    if para['heading_level'] is None and not para['is_empty']:
                        content.append(para)

            self._add_content_to_node(tree, heading_idx, content)

    def _collect_heading_indices(self, nodes: List[Dict], indices: set):
        """递归收集所有标题的索引"""
        for node in nodes:
            indices.add(node['heading_index'])
            self._collect_heading_indices(node['children'], indices)

    def _add_content_to_node(self, nodes: List[Dict], heading_idx: int, content: List[Dict]):
        """递归为指定节点添加内容"""
        for node in nodes:
            if node['heading_index'] == heading_idx:
                node['content_paragraphs'] = content
                return
            self._add_content_to_node(node['children'], heading_idx, content)

    def extract_by_heading(self, heading_text: str) -> Optional[Dict[str, Any]]:
        """按标题文本提取段落"""
        tree = self.build_structure_tree()
        return self._find_node_by_text(tree, heading_text)

    def _find_node_by_text(self, nodes: List[Dict], text: str) -> Optional[Dict]:
        """递归按文本查找节点"""
        for node in nodes:
            if text in node['text']:
                return node
            found = self._find_node_by_text(node['children'], text)
            if found:
                return found
        return None

    def extract_metadata(self) -> Dict[str, Any]:
        """提取文档元数据"""
        if not self.doc:
            self.load_document()

        core_props = self.doc.core_properties
        return {
            'title': core_props.title or "",
            'author': core_props.author or "",
            'subject': core_props.subject or "",
            'keywords': core_props.keywords or "",
            'created': str(core_props.created) if core_props.created else "",
            'modified': str(core_props.modified) if core_props.modified else "",
            'last_modified_by': core_props.last_modified_by or "",
            'revision': core_props.revision or 0,
            'paragraph_count': len(self.doc.paragraphs),
            'table_count': len(self.doc.tables),
            'section_count': len(self.doc.sections)
        }

    def parse(self) -> Dict[str, Any]:
        """完整解析文档"""
        print(f"📄 解析文档: {self.file_path.name}")

        if not self.load_document():
            return {'success': False, 'error': '加载文档失败'}

        result = {
            'success': True,
            'file_path': str(self.file_path),
            'metadata': self.extract_metadata(),
            'headings': self.extract_headings(),
            'paragraphs': self.extract_all_paragraphs(),
            'structure_tree': self.build_structure_tree(),
        }

        print(f"  ✅ 标题数量: {len(result['headings'])}")
        print(f"  ✅ 段落数量: {len(result['paragraphs'])}")
        print(f"  ✅ 结构树层级: {max([h['level'] for h in result['headings']], default=0)} 层")

        return result

    def export_json(self, output_path: str = None):
        """导出解析结果为 JSON"""
        result = self.parse()

        if not output_path:
            output_path = self.file_path.with_suffix('.parsed.json')

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        print(f"💾 结果已导出: {output_path}")
        return output_path


def main():
    parser = argparse.ArgumentParser(description='Word 文档结构解析工具')
    parser.add_argument('--input', '-i', required=True, help='输入 Word 文档路径')
    parser.add_argument('--output', '-o', help='输出 JSON 文件路径')
    parser.add_argument('--config', '-c', help='配置文件路径')
    parser.add_argument('--extract-by-heading', '-e', help='按标题提取段落')

    args = parser.parse_args()

    word_parser = WordParser(args.input, args.config)

    if args.extract_by_heading:
        content = word_parser.extract_by_heading(args.extract_by_heading)
        if content:
            print(f"\n📑 标题: {content['text']}")
            print(f"📝 内容段落数: {len(content['content_paragraphs'])}")
            for para in content['content_paragraphs']:
                print(f"  {para['text'][:100]}...")
        else:
            print(f"❌ 未找到包含 '{args.extract_by_heading}' 的标题")
    else:
        word_parser.export_json(args.output)


if __name__ == "__main__":
    main()
